#!/usr/bin/env python3
"""Patch the latest XVS MRD into its next version.

``generate_requirements_documents.build_mrd`` produced the v2.5.1 rename delta
and is pinned to that revision's reference. The MRD has moved on since, so
later versions are produced here: the newest document supplies the page
furniture, media, typography, and color system, and this script edits only the
cells whose evidence changed.

This revision (v2.7 -> v2.8) records the resolved Branch lifecycle defect.

    python tools/patch_mrd.py \
        --source module-requirements/XVS_Module_Requirements_Document_v2.7.docx \
        --output module-requirements/XVS_Module_Requirements_Document_v2.8.docx
"""
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document

from generate_requirements_documents import (
    BLUE,
    PALE_BLUE,
    assert_no_em_dash,
    find_table_by_header,
    rebuild_table,
    set_table_widths,
    update_extended_title,
    write_cell,
)

SOURCE_VERSION = "2.7"
TARGET_VERSION = "2.8"
REVIEW_DATE = "12 August 2026"
CODE_BASELINE = (
    "355c59b665fd27d4c47152fa9638093b2e95dc49 plus uncommitted Branch "
    "lifecycle work (12 August 2026)"
)

CHANGE_SUMMARY = (
    "Resolved the P0 Branch transition defect. Branch.transition() now persists "
    "status with activated_at, deactivated_at, and closed_at, publishes an "
    "allowed-transition matrix in which CLOSED is terminal and PENDING is "
    "unreachable, and writes the BranchLifecycle row inside the same "
    "transaction, so branch status-change history is recorded for the first "
    "time. The endpoint resolves the Branch within its route School rather than "
    "by code alone, and requires platform.branches.manage alongside XVision "
    "staff. Also corrected the platform-wide error envelope, which turned any "
    "string-form validation error into a 500. Module 1 remains Backend Partial: "
    "School lifecycle controls and closed-Branch access rules are unchanged."
)

DELTA_ROWS = [
    [
        "Branch lifecycle transitions",
        "Resolved",
        "Status, activation, deactivation, and closure timestamps persist together with the BranchLifecycle history row in one transaction. The allowed-transition matrix refuses every other move.",
    ],
    [
        "Branch transition authorization",
        "Tightened",
        "The endpoint requires platform.branches.manage in addition to XVision staff, and resolves the Branch within its route School so a code cannot reach another School's Branch.",
    ],
    [
        "Platform error envelope",
        "Corrected",
        "A validation error raised in string form is answered as a 400 rather than crashing the shared exception handler and returning a 500.",
    ],
    [
        "Module 1 status",
        "Unchanged",
        "Backend Partial and Integration Partial. School lifecycle controls and closed-Branch access rules remain open.",
    ],
]


def replace_cell_text(cell, text: str, **kwargs) -> None:
    """Rewrite a cell, keeping the first paragraph's formatting."""
    write_cell(cell, text, **kwargs)


def patch(source: Path, output: Path) -> None:
    doc = Document(str(source))

    doc.core_properties.title = f"XVS Module Requirements Document v{TARGET_VERSION}"
    doc.core_properties.version = TARGET_VERSION

    # --- Cover ---------------------------------------------------------
    cover = doc.tables[0]
    for paragraph in cover.rows[0].cells[0].paragraphs:
        for run in paragraph.runs:
            if SOURCE_VERSION in run.text:
                run.text = run.text.replace(SOURCE_VERSION, TARGET_VERSION)

    # --- Document control ----------------------------------------------
    control = doc.tables[1]
    for row in control.rows:
        label = row.cells[0].text.strip()
        if label == "Version":
            replace_cell_text(row.cells[1], TARGET_VERSION, size=9)
        elif label == "Review date":
            replace_cell_text(row.cells[1], REVIEW_DATE, size=9)
        elif label == "Source scope":
            replace_cell_text(
                row.cells[1], f"Current backend codebase at {CODE_BASELINE}", size=9
            )

    # --- Contents -------------------------------------------------------
    contents = doc.tables[2]
    for row in contents.rows:
        if row.cells[0].text.strip().startswith("5."):
            replace_cell_text(
                row.cells[0], f"5. v{TARGET_VERSION} Documentation Delta",
                size=9, bold=True, color=BLUE,
            )
            replace_cell_text(
                row.cells[1],
                "Evidence corrected in this revision",
                size=9,
            )

    # --- Section 5 heading and lead-in ----------------------------------
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("5. v") and text.endswith("Documentation Delta"):
            for run in paragraph.runs:
                run.text = ""
            paragraph.runs[0].text = f"5. v{TARGET_VERSION} Documentation Delta"
        elif text.startswith("This patch changes documentation terminology"):
            for run in paragraph.runs:
                run.text = ""
            paragraph.runs[0].text = (
                "This revision records evidence that changed in the backend. "
                "Module numbering, the status model, and the capability count "
                "are unchanged."
            )

    # --- Module 1 Needs Attention ---------------------------------------
    module_one_fixed = (
        "Fix branch lifecycle transitions: the current save call references a "
        "deleted_at field that Branch does not define."
    )
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if module_one_fixed not in cell.text:
                    continue
                for paragraph in list(cell.paragraphs):
                    if module_one_fixed in paragraph.text:
                        paragraph._p.getparent().remove(paragraph._p)

    # --- Priority gaps: drop the resolved P0 -----------------------------
    gaps = find_table_by_header(doc, "Required completion")
    for row in list(gaps.rows):
        if "Branch transition defect" in row.cells[1].text:
            gaps._tbl.remove(row._tr)

    # --- Section 5 delta table -------------------------------------------
    delta = find_table_by_header(doc, "documentation delta")
    rebuild_table(
        delta,
        [f"v{TARGET_VERSION} documentation delta", "Decision", "Evidence"],
        DELTA_ROWS,
        [1.75, 0.95, 4.57],
        font_size=8.2,
    )

    # --- Recommended next build order -------------------------------------
    for table in doc.tables:
        for row in table.rows:
            cell = row.cells[-1]
            if cell.text.strip().startswith("Fix branch transitions,"):
                replace_cell_text(
                    cell,
                    "Add safe school lifecycle controls, require privileged MFA, "
                    "and make procurement approval fail closed.",
                    size=8.2,
                )

    # --- Change log --------------------------------------------------------
    change_log = find_table_by_header(doc, "Version | Date | Summary")
    template = change_log.rows[1]
    new_tr = copy_row(template)
    template._tr.addprevious(new_tr)
    new_row = change_log.rows[1]
    replace_cell_text(new_row.cells[0], TARGET_VERSION, size=8)
    replace_cell_text(new_row.cells[1], "12 Aug 2026", size=8)
    replace_cell_text(new_row.cells[2], CHANGE_SUMMARY, size=8)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    update_extended_title(
        output, f"XVS Module Requirements Document v{TARGET_VERSION}"
    )
    assert_no_em_dash(output)


def copy_row(row):
    import copy as _copy

    return _copy.deepcopy(row._tr)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    patch(args.source, args.output)
    print(args.output)


if __name__ == "__main__":
    main()

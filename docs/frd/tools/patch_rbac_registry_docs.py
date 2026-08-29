#!/usr/bin/env python3
"""Version the MRD and RBAC FRD for permission-registry integrity."""

from __future__ import annotations

import argparse
import copy
from pathlib import Path

from docx import Document

from generate_requirements_documents import (
    BLUE,
    assert_no_em_dash,
    rebuild_table,
    shrink_inherited_media,
    update_extended_title,
    write_cell,
    write_paragraph,
)


REVIEW_DATE = "29 August 2026"
MRD_SOURCE_VERSION = "2.40"
MRD_TARGET_VERSION = "2.41"
FRD_SOURCE_VERSION = "1.6"
FRD_TARGET_VERSION = "1.7"

MRD_SOURCE_SCOPE = (
    "Backend worktree requiring permission scope and preserving immutable "
    "permission-registry identity on supported APIs (29 August 2026)"
)
FRD_SOURCE_SCOPE = (
    "Backend worktree requiring permission scope on registry creation and "
    "refusing semantic renames of permission vocabulary and keys "
    "(29 August 2026)"
)

MRD_CHANGE_SUMMARY = (
    "Closed two permission-registry integrity gaps in Module 4. Creating a "
    "permission through the platform API now requires and returns its TENANT "
    "or PLATFORM scope, so the registry cannot report success while storing "
    "an unclassified key that a tenant cannot hold. Module, resource and "
    "action identity is immutable after creation, both on a Permission and "
    "on the vocabulary records themselves. The update endpoint no longer "
    "rewrites a permission primary key beneath direct grants, groups, "
    "dependencies, overrides, role defaults or pending role changes. A real "
    "semantic rename now requires new vocabulary and a new permission, an "
    "explicit grant migration and deactivation of the old identity. Eight "
    "focused registry tests pass. The complete 444-test RBAC run completed "
    "with the same two unrelated field-security expectation failures in "
    "payments and procurement. Module 4 remains Backend Complete and "
    "Integration Complete with seventeen capability entries. Backend "
    "evidence only; nothing here is deployed."
)

FRD_CHANGE_SUMMARY = (
    "Makes the supported permission registry preserve security identity. "
    "PermissionSerializer now requires and returns scope on create, rejecting "
    "an omitted or invalid audience before a row exists. Permission module, "
    "resource and action values cannot change on an existing Permission, and "
    "the vocabulary detail serializers refuse their own identity renames, "
    "while descriptions and active-state controls remain editable. "
    "PermissionDetailView no longer rewrites the dotted primary key directly. "
    "A semantic rename is therefore a deliberate create, migrate and "
    "deactivate operation rather than an ordinary PATCH. Eight focused "
    "registry tests pass, including preservation of an existing role grant. "
    "The complete 444-test RBAC run completed with the same two unrelated "
    "field-security expectation failures in payments and procurement. The "
    "scope-registration Needs Attention item narrows to seed commands, which "
    "still lack a global no-unclassified-key assertion. Backend evidence "
    "only; nothing here is deployed."
)


def replace_cell(cell, text: str, **kwargs) -> None:
    while len(cell.paragraphs) > 1:
        paragraph = cell.paragraphs[-1]
        paragraph._p.getparent().remove(paragraph._p)
    write_cell(cell, text, **kwargs)


def replace_paragraph(paragraph, text: str, **kwargs) -> None:
    write_paragraph(paragraph, text, **kwargs)


def replace_cover_version(table, source: str, target: str) -> None:
    for paragraph in table.rows[0].cells[0].paragraphs:
        for run in paragraph.runs:
            if source in run.text:
                run.text = run.text.replace(source, target)


def update_control_table(table, *, version: str, source_scope: str, mrd_baseline=None):
    for row in table.rows:
        label = row.cells[0].text.strip()
        if label == "Version":
            replace_cell(row.cells[1], version, size=9)
        elif label == "Review date":
            replace_cell(row.cells[1], REVIEW_DATE, size=9)
        elif label == "Source scope":
            replace_cell(row.cells[1], source_scope, size=9)
        elif label == "MRD baseline" and mrd_baseline:
            replace_cell(row.cells[1], mrd_baseline, size=9)


def prepend_change_log(table, version: str, date: str, summary: str) -> None:
    template = table.rows[1]
    new_tr = copy.deepcopy(template._tr)
    template._tr.addprevious(new_tr)
    new_row = table.rows[1]
    replace_cell(new_row.cells[0], version, size=8)
    replace_cell(new_row.cells[1], date, size=8)
    replace_cell(new_row.cells[2], summary, size=8)


def patch_mrd(source: Path, output: Path) -> None:
    doc = Document(str(source))
    doc.core_properties.title = f"XVS Module Requirements Document v{MRD_TARGET_VERSION}"
    doc.core_properties.version = MRD_TARGET_VERSION

    replace_cover_version(doc.tables[0], MRD_SOURCE_VERSION, MRD_TARGET_VERSION)
    update_control_table(
        doc.tables[1], version=MRD_TARGET_VERSION, source_scope=MRD_SOURCE_SCOPE,
    )

    for row in doc.tables[2].rows:
        if row.cells[0].text.strip().startswith("5."):
            replace_cell(
                row.cells[0],
                f"5. v{MRD_TARGET_VERSION} Documentation Delta",
                size=9,
                bold=True,
                color=BLUE,
            )
            replace_cell(
                row.cells[1],
                "Permission creation is classified and registry identity is stable",
                size=9,
            )

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == f"5. v{MRD_SOURCE_VERSION} Documentation Delta":
            replace_paragraph(
                paragraph,
                f"5. v{MRD_TARGET_VERSION} Documentation Delta",
                size=17,
                bold=True,
                space_before=15,
                space_after=8,
            )
        elif text.startswith("This revision makes the durable RBAC audit"):
            replace_paragraph(
                paragraph,
                "This revision prevents the supported permission registry from creating an unclassified key or rewriting an identity that existing grants reference.",
                size=9,
                space_after=5,
            )
        elif text.startswith("Two-layer platform and school authorization"):
            replace_paragraph(
                paragraph,
                "Two-layer platform and school authorization with reusable permission definitions, roles, templates, assignments, and controlled overrides. Platform-only scope is enforced independently from restricted sensitivity. The registry requires that scope when a key is created, and its module, resource and action identity cannot be renamed through a supported update. Restricted keys cannot enter direct role writes, groups, or ALLOW overrides; they enter a custom role only through a change request decided by a different reviewer who already holds every restricted key being granted. Branch reach remains grant-derived, and one person may hold the same role at more than one branch. Every supported role permission or group change records its actor, reason, source, approval reference where applicable, and complete before-and-after access configuration in the same transaction as the mutation.",
                size=9,
                space_after=5,
            )

    current_decision = doc.tables[14].rows[8].cells[0].text.rstrip()
    current_decision += (
        "\n• Registry identity is stable. A new permission must declare TENANT or "
        "PLATFORM scope, and an existing permission or vocabulary row cannot "
        "rename its module, resource or action through the API. A semantic "
        "rename creates a new identity, migrates grants explicitly and "
        "deactivates the old one."
    )
    replace_cell(doc.tables[14].rows[8].cells[0], current_decision, size=8.2)

    rebuild_table(
        doc.tables[76],
        [f"v{MRD_TARGET_VERSION} documentation delta", "Decision", "Evidence"],
        [
            [
                "Permission scope on create",
                "Required",
                "The platform registry refuses an omitted or invalid scope and returns the accepted TENANT or PLATFORM value in its response.",
            ],
            [
                "Concrete permission identity",
                "Immutable",
                "Changing module, resource or action is a validation error; the detail view no longer updates the dotted primary key directly.",
            ],
            [
                "Vocabulary identity",
                "Immutable",
                "Module names, resource module/name pairs and action names cannot be renamed through their detail endpoints; descriptions and liveness remain editable.",
            ],
            [
                "Semantic rename",
                "Explicit",
                "Create the new vocabulary and permission, migrate every intended reference, then deactivate the old identity instead of rewriting it beneath grants.",
            ],
        ],
        [1.75, 0.95, 4.57],
        font_size=8.2,
    )

    prepend_change_log(
        doc.tables[78], MRD_TARGET_VERSION, "29 Aug 2026", MRD_CHANGE_SUMMARY,
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    update_extended_title(output, f"XVS Module Requirements Document v{MRD_TARGET_VERSION}")
    shrink_inherited_media(output)
    assert_no_em_dash(output)


def patch_frd(source: Path, output: Path) -> None:
    doc = Document(str(source))
    doc.core_properties.title = (
        "XVS M04 Roles and Permissions RBAC Functional Requirements "
        f"Document v{FRD_TARGET_VERSION}"
    )
    doc.core_properties.version = FRD_TARGET_VERSION

    replace_cover_version(doc.tables[0], FRD_SOURCE_VERSION, FRD_TARGET_VERSION)
    update_control_table(
        doc.tables[1],
        version=FRD_TARGET_VERSION,
        source_scope=FRD_SOURCE_SCOPE,
        mrd_baseline=(
            f"XVS Module Requirements Document v{MRD_TARGET_VERSION}, Module 4, "
            "seventeen capability entries"
        ),
    )

    for row in doc.tables[2].rows:
        if row.cells[0].text.strip() == "10. MRD Traceability":
            replace_cell(
                row.cells[1],
                f"Agreement with MRD v{MRD_TARGET_VERSION}'s seventeen capability entries",
                size=9,
            )

    replace_cell(
        doc.tables[3].rows[1].cells[1],
        "PermissionModule, PermissionResource and PermissionAction, from which a Permission key is composed as module.resource.action and stored as the primary key. Their identity is immutable on supported updates; descriptions and active states remain editable.",
        size=8.5,
    )
    replace_cell(
        doc.tables[3].rows[2].cells[1],
        "Permission rows carrying description, sensitivity, restricted flag, active flag, and the required scope that says which audience may ever hold the key. Module, resource and action identity is immutable after creation.",
        size=8.5,
    )

    fr001 = doc.tables[9]
    replace_cell(
        fr001.rows[2].cells[1],
        "Permission.save composes key as f\"{module_id}.{resource.name}.{action_id}\", and key is the primary key. PermissionModule and PermissionAction are slug-keyed; PermissionResource is unique on (module, name). The registry serializers reject changes to an existing Permission's module, resource or action and reject identity changes on the vocabulary records themselves. PermissionDetailView no longer updates the primary key directly. The three Permission vocabulary FKs use db_constraint=False with on_delete=PROTECT, while grant, group, dependency, override, default and pending-change relations reference the dotted key.",
        size=8.5,
    )
    replace_cell(
        fr001.rows[3].cells[1],
        "Creating a permission from module 'finance', resource 'invoice' and action 'view' yields exactly finance.invoice.view. A key cannot be set by hand, a duplicate triple collides on the primary key, and PATCH cannot turn an existing key or its vocabulary into a different identity.",
        size=8.5,
    )
    replace_cell(
        fr001.rows[4].cells[1],
        "A semantic rename is not an ordinary update. It requires new vocabulary and a new Permission, an explicit migration of every intended reference, and deactivation of the old identity. Direct ORM or database writes can still bypass serializer validation and are unsupported for this operation.",
        size=8.5,
    )

    fr002 = doc.tables[10]
    replace_cell(
        fr002.rows[2].cells[1],
        "PermissionScope has exactly two values, TENANT and PLATFORM, and is a field on both Permission and PermissionGroup. PermissionSerializer requires the value when the platform API creates a key and includes it in registry responses, so the supported create path cannot write an unclassified permission. The seeder classification remains by module with an explicit exception list: seed_platform_permissions holds TENANT_HOLDABLE_KEYS for the tenant-facing platform.team.* and platform.audit view/export keys, while the other platform registry keys are PLATFORM.",
        size=8.5,
    )
    replace_cell(
        fr002.rows[3].cells[1],
        "Creating a permission without TENANT or PLATFORM scope is refused before a row exists. A school role or override carrying platform.permissions.view, platform.schools.view, platform.team_overrides.* or platform.staff_payroll.* is refused at the point of grant, not at the gate.",
        size=8.5,
    )
    replace_cell(
        fr002.rows[4].cells[1],
        "Seed commands write models directly and do not pass through PermissionSerializer. An unclassified seeded key still fails closed at grant time, but the aggregate seeding process does not yet assert that every active key is classified.",
        size=8.5,
    )

    fr023 = doc.tables[31]
    replace_cell(
        fr023.rows[2].cells[1],
        "seed_all_permissions runs seventeen idempotent steps in dependency order. The school and platform seeders register the dedicated school.roles.approve and platform.roles.approve keys as restricted. Group seeders remove and skip restricted members. Seeded system roles remain the trusted bootstrap. Separately, the platform registry API now requires and returns scope on every new Permission, so an operator cannot create an unclassified key through the supported console path.",
        size=8.5,
    )
    replace_cell(
        fr023.rows[3].cells[1],
        "The registry API refuses creation without scope. A key written outside that serializer with no scope is refused for every tenant that is not the platform and is named in the grant error, so omission fails closed rather than widening access.",
        size=8.5,
    )
    replace_cell(
        fr023.rows[4].cells[1],
        "Seeders remain the boundary's other authors, and no aggregate test or seed_all_permissions check proves that every active registered key carries a scope. Migration 0007 classified the original 344 / 40 split; keys added later still depend on each seeder being correct.",
        size=8.5,
    )

    replace_cell(
        doc.tables[36].rows[1].cells[2],
        "Modules and actions are slug primary keys; a resource is unique within its module. Supported detail updates refuse identity changes and permit descriptive or active-state changes only.",
        size=8.2,
    )
    replace_cell(
        doc.tables[36].rows[2].cells[2],
        "Primary key is the dotted key itself. Carries required API scope, sensitivity, is_restricted and is_active. Module, resource and action cannot change after creation through the registry API.",
        size=8.2,
    )

    registry = doc.tables[38]
    replace_cell(
        registry.rows[2].cells[1],
        "One module. .view, .update, and .manage for delete. Name is immutable; description and active state remain editable.",
        size=8.2,
    )
    replace_cell(
        registry.rows[3].cells[1],
        "Resources within a module. Same key pair. An existing resource's module and name are immutable.",
        size=8.2,
    )
    replace_cell(
        registry.rows[4].cells[1],
        "The reusable action verbs. Same key pair. An existing action name is immutable.",
        size=8.2,
    )
    replace_cell(
        registry.rows[5].cells[1],
        "The permission registry itself. platform.permissions.view and .create. POST requires TENANT or PLATFORM scope and returns it.",
        size=8.2,
    )
    replace_cell(
        registry.rows[6].cells[1],
        "One key. .view, .update, and .delete. Module, resource and action are immutable; metadata, scope and active state remain editable.",
        size=8.2,
    )

    conditions = doc.tables[40]
    template = conditions.rows[7]
    for condition, answer in reversed([
        (
            "Permission creation without scope",
            "400 under scope. No unclassified Permission row is created.",
        ),
        (
            "A permission or vocabulary update that changes identity",
            "400 naming the immutable field and directing the operator to create new vocabulary, migrate grants explicitly and deactivate the old identity.",
        ),
    ]):
        new_tr = copy.deepcopy(template._tr)
        template._tr.addprevious(new_tr)
        new_row = conditions.rows[7]
        replace_cell(new_row.cells[0], condition, size=8.2)
        replace_cell(new_row.cells[1], answer, size=8.2)

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("MRD v2.40 records Module 4"):
            replace_paragraph(
                paragraph,
                f"MRD v{MRD_TARGET_VERSION} records Module 4 as Roles & Permissions (RBAC), Phase V1, Backend Complete, Integration Complete, code vs_rbac, with seventeen capability entries. The module number, name, phase, states and ownership are taken from that document and agree with this revision.",
                size=9,
                space_after=5,
            )
        elif text.startswith("•  Both escalations end to end"):
            registry_test = paragraph.insert_paragraph_before()
            registry_test.style = paragraph.style
            replace_paragraph(
                registry_test,
                "•  Registry integrity: creation refuses an omitted scope and returns the accepted scope; permission and vocabulary identity renames are refused; a failed rename leaves an existing role grant attached to the original key.",
                size=9,
                space_after=3,
            )

    needs_attention = doc.tables[42]
    replace_cell(
        needs_attention.rows[6].cells[1],
        "The platform registry API now requires scope, but seed commands still write Permission rows directly and no aggregate check proves that every active seeded key is classified. A forgotten seeder value remains safe at grant time, where it is refused, but the error appears after deployment rather than during seeding.",
        size=8.2,
    )
    replace_cell(
        needs_attention.rows[6].cells[2],
        "Add a test or a check in seed_all_permissions asserting that no active Permission has an empty scope after every seeder runs.",
        size=8.2,
    )

    replace_cell(
        doc.tables[44].rows[1].cells[2],
        "Implemented. The registry requires scope on create, keeps module, resource and action identity immutable, and uses the same hierarchy as the runtime shutdown control.",
        size=8.2,
    )
    replace_cell(
        doc.tables[44].rows[14].cells[2],
        "Implemented with limits. API-created keys require explicit scope, approval keys are restricted, system roles provide bootstrap, and groups exclude restricted members; aggregate seeder scope verification remains open.",
        size=8.2,
    )

    replace_cell(
        doc.tables[45].rows[0].cells[0],
        f"MRD RECONCILIATION\n• MRD v{MRD_TARGET_VERSION} lists Module 4 as Complete with seventeen capability entries, all traced above. Requiring scope and preserving registry identity hardens the existing permission-registry capability without adding a new product surface or changing the count.\n• The supported API cannot create an unclassified permission or rename a key beneath its grants. Semantic renames are explicit create, migrate and deactivate operations.\n• What remains is recorded in Needs Attention, including assignment expiry, two unnarrowed read surfaces, advisory branch reach, the generic platform-decision declaration, prebuilt scope, aggregate seeder scope verification and privileged MFA.\n• No other divergence was found between MRD v{MRD_TARGET_VERSION}'s Module 4 row and the inspected code.",
        size=8.5,
    )

    prepend_change_log(
        doc.tables[46], FRD_TARGET_VERSION, "29 Aug 2026", FRD_CHANGE_SUMMARY,
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    update_extended_title(
        output,
        "XVS M04 Roles and Permissions RBAC Functional Requirements "
        f"Document v{FRD_TARGET_VERSION}",
    )
    shrink_inherited_media(output)
    assert_no_em_dash(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--mrd-source", type=Path, required=True)
    parser.add_argument("--mrd-output", type=Path, required=True)
    parser.add_argument("--frd-source", type=Path, required=True)
    parser.add_argument("--frd-output", type=Path, required=True)
    args = parser.parse_args()

    patch_mrd(args.mrd_source, args.mrd_output)
    patch_frd(args.frd_source, args.frd_output)
    print(args.mrd_output)
    print(args.frd_output)


if __name__ == "__main__":
    main()

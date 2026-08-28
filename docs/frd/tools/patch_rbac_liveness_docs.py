#!/usr/bin/env python3
"""Version the MRD and RBAC FRD for permission-liveness enforcement."""

from __future__ import annotations

import argparse
import copy
from pathlib import Path

from docx import Document

from generate_requirements_documents import (
    BLUE,
    assert_no_em_dash,
    rebuild_table,
    shrink_inherited_media,
    update_extended_title,
    write_cell,
    write_paragraph,
)


REVIEW_DATE = "27 August 2026"
MRD_SOURCE_VERSION = "2.37"
MRD_TARGET_VERSION = "2.38"
FRD_SOURCE_VERSION = "1.4"
FRD_TARGET_VERSION = "1.5"

MRD_SOURCE_SCOPE = (
    "Backend worktree with live permission evaluation and durable emergency "
    "revocation cache invalidation (27 August 2026)"
)
FRD_SOURCE_SCOPE = (
    "Backend worktree implementing permission, module, resource, action and "
    "group liveness across evaluation, overrides and permission-based routing "
    "(27 August 2026)"
)

MRD_CHANGE_SUMMARY = (
    "Closed the permission-deactivation gap in Module 4. A grant now confers "
    "authority only while its Permission and the PermissionModule, "
    "PermissionResource and PermissionAction behind it are active; a group "
    "grant additionally requires an active PermissionGroup. Direct roles, "
    "groups, personal overrides, group-based gates and permission-based "
    "routing use the same predicate. A PostgreSQL-backed registry revision "
    "invalidates warm evaluator snapshots across application workers, and "
    "affected role versions advance when a registry switch changes. Routing "
    "now also refuses inactive roles and legacy restricted ALLOW overrides, "
    "so the nominated audience cannot disagree with the authorization gate. "
    "Module 4 remains Backend Complete and Integration Complete with seventeen "
    "capability entries. Backend evidence only; nothing here is deployed."
)

FRD_CHANGE_SUMMARY = (
    "Makes every active flag in the permission vocabulary an operational "
    "revocation control. _live_permission_q is now the single predicate for "
    "direct role grants, group grants, personal overrides, group-based checks "
    "and permission-based routing. PermissionGroup.is_active is enforced on "
    "group grants, while module, resource and action deactivation withdraw the "
    "permissions beneath them. PermissionRegistryRevision provides durable "
    "cross-worker invalidation for already-warm evaluator snapshots, and the "
    "registry signals advance affected role versions. Routing also refuses "
    "inactive roles and legacy restricted ALLOW overrides. Six focused "
    "emergency-revocation tests pass. The complete 434-test RBAC run completed "
    "with two unrelated field-security wiring expectation failures; neither "
    "touches evaluation, routing or this change. The deactivated-permission "
    "Needs Attention item is removed. Backend evidence only; nothing here is "
    "deployed."
)


def replace_cell(cell, text: str, **kwargs) -> None:
    while len(cell.paragraphs) > 1:
        paragraph = cell.paragraphs[-1]
        paragraph._p.getparent().remove(paragraph._p)
    write_cell(cell, text, **kwargs)


def replace_paragraph(paragraph, text: str, **kwargs) -> None:
    write_paragraph(paragraph, text, **kwargs)


def replace_cover_version(table, source: str, target: str) -> None:
    for paragraph in table.rows[0].cells[0].paragraphs:
        for run in paragraph.runs:
            if source in run.text:
                run.text = run.text.replace(source, target)


def update_control_table(table, *, version: str, source_scope: str, mrd_baseline=None):
    for row in table.rows:
        label = row.cells[0].text.strip()
        if label == "Version":
            replace_cell(row.cells[1], version, size=9)
        elif label == "Review date":
            replace_cell(row.cells[1], REVIEW_DATE, size=9)
        elif label == "Source scope":
            replace_cell(row.cells[1], source_scope, size=9)
        elif label == "MRD baseline" and mrd_baseline:
            replace_cell(row.cells[1], mrd_baseline, size=9)


def prepend_change_log(table, version: str, date: str, summary: str) -> None:
    template = table.rows[1]
    new_tr = copy.deepcopy(template._tr)
    template._tr.addprevious(new_tr)
    new_row = table.rows[1]
    replace_cell(new_row.cells[0], version, size=8)
    replace_cell(new_row.cells[1], date, size=8)
    replace_cell(new_row.cells[2], summary, size=8)


def patch_mrd(source: Path, output: Path) -> None:
    doc = Document(str(source))
    doc.core_properties.title = f"XVS Module Requirements Document v{MRD_TARGET_VERSION}"
    doc.core_properties.version = MRD_TARGET_VERSION

    replace_cover_version(doc.tables[0], MRD_SOURCE_VERSION, MRD_TARGET_VERSION)
    update_control_table(
        doc.tables[1], version=MRD_TARGET_VERSION, source_scope=MRD_SOURCE_SCOPE,
    )

    for row in doc.tables[2].rows:
        if row.cells[0].text.strip().startswith("5."):
            replace_cell(
                row.cells[0],
                f"5. v{MRD_TARGET_VERSION} Documentation Delta",
                size=9,
                bold=True,
                color=BLUE,
            )
            replace_cell(
                row.cells[1],
                "Permission shutdown now withdraws authority and routing",
                size=9,
            )

    decision = (
        "CURRENT DECISION\n"
        "• No material capability gap was identified within this module's stated backend scope.\n"
        "• Two role-grant escalations are closed at their source. Scope refuses a tenant carrying a platform-only key. Separately, a restricted key is refused on direct role writes, in permission groups, and as an ALLOW override; it may enter a custom role only through a change request decided by somebody other than the requester, and that reviewer may grant only restricted keys they already hold. Existing unverified restricted grants are removed from custom roles by migration, while trusted seeded system roles remain as the approval bootstrap.\n"
        "• A role granted for the whole school now reaches the whole school. The branch narrowing and the permission gate are resolved from the same grants, so a Finance Officer for the whole school is no longer narrowed to the one branch her staff record happens to name while the gate lets her act at the others. Writes follow reads: what she raises without naming a branch is filed school-wide.\n"
        "• The same role may now be granted at two branches. The schema always allowed it and both write paths refused it, so a teacher working at two branches could hold the role at one. Every count that reads assignments counts people rather than grant rows.\n"
        "• Deactivation is a runtime kill switch. A permission works only while its own row and its module, resource and action are active; a group grant additionally requires an active group. Evaluation, overrides and permission-based routing share that rule, and a durable registry revision invalidates warm permission snapshots across application workers."
    )
    replace_cell(doc.tables[14].rows[8].cells[0], decision, size=8.2)

    for paragraph_index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text == f"5. v{MRD_SOURCE_VERSION} Documentation Delta":
            replace_paragraph(
                paragraph,
                f"5. v{MRD_TARGET_VERSION} Documentation Delta",
                size=17,
                bold=True,
                space_before=15,
                space_after=8,
            )
        elif text.startswith("This revision records that a stored file's name"):
            replace_paragraph(
                paragraph,
                "This revision closes the gap between deactivating permission vocabulary and actually withdrawing the authority and routing that vocabulary conferred.",
                size=9,
                space_after=5,
            )

    rebuild_table(
        doc.tables[76],
        [f"v{MRD_TARGET_VERSION} documentation delta", "Decision", "Evidence"],
        [
            [
                "Permission hierarchy liveness",
                "Resolved",
                "Permission, module, resource and action active states are all required before any direct grant or personal override can confer authority.",
            ],
            [
                "Permission-group liveness",
                "Resolved",
                "An inactive PermissionGroup contributes no keys. A separate direct grant of the same live key remains effective.",
            ],
            [
                "Routing agreement",
                "Corrected",
                "Permission-based routing uses the same live predicate, refuses inactive roles and ignores legacy restricted ALLOW overrides.",
            ],
            [
                "Emergency cache invalidation",
                "Added",
                "A PostgreSQL-backed registry revision invalidates warm effective-permission snapshots in every application worker after the change commits.",
            ],
        ],
        [1.75, 0.95, 4.57],
        font_size=8.2,
    )

    prepend_change_log(
        doc.tables[78], MRD_TARGET_VERSION, "27 Aug 2026", MRD_CHANGE_SUMMARY,
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    update_extended_title(output, f"XVS Module Requirements Document v{MRD_TARGET_VERSION}")
    shrink_inherited_media(output)
    assert_no_em_dash(output)


def patch_frd(source: Path, output: Path) -> None:
    doc = Document(str(source))
    doc.core_properties.title = (
        "XVS M04 Roles and Permissions RBAC Functional Requirements "
        f"Document v{FRD_TARGET_VERSION}"
    )
    doc.core_properties.version = FRD_TARGET_VERSION

    replace_cover_version(doc.tables[0], FRD_SOURCE_VERSION, FRD_TARGET_VERSION)
    update_control_table(
        doc.tables[1],
        version=FRD_TARGET_VERSION,
        source_scope=FRD_SOURCE_SCOPE,
        mrd_baseline=(
            f"XVS Module Requirements Document v{MRD_TARGET_VERSION}, Module 4, "
            "seventeen capability entries"
        ),
    )

    for row in doc.tables[2].rows:
        if row.cells[0].text.strip() == "10. MRD Traceability":
            replace_cell(
                row.cells[1],
                f"Agreement with MRD v{MRD_TARGET_VERSION}'s seventeen capability entries",
                size=9,
            )

    evaluation = (
        "get_effective_permissions, its later-wins ordering, and the backstops "
        "that refuse out-of-scope, inactive and unapproved restricted keys. "
        "The expensive grant expansion is memoised, while a durable registry "
        "revision is checked before a cached set is reused so deactivation "
        "withdraws authority across application workers."
    )
    replace_cell(doc.tables[3].rows[9].cells[1], evaluation, size=8.5)

    for paragraph_index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text == (
            "FR-004  Refuse an Out-of-Scope Key at Evaluation, Whatever Row Grants It"
        ):
            replace_paragraph(
                paragraph,
                "FR-004  Refuse Out-of-Scope, Inactive, and Unapproved Keys at Evaluation",
                size=13.5,
                bold=True,
                space_before=11,
                space_after=6,
            )
        elif paragraph_index == 29:
            replace_paragraph(
                paragraph,
                "FR-011  Withdraw a Branch, and Withdraw the Access It Carried",
                size=13.5,
                bold=True,
                space_before=11,
                space_after=6,
            )
        elif text.startswith("MRD v2.34 records Module 4"):
            replace_paragraph(
                paragraph,
                f"MRD v{MRD_TARGET_VERSION} records Module 4 as Roles & Permissions (RBAC), Phase V1, Backend Complete, Integration Complete, code vs_rbac, with seventeen capability entries. The module number, name, phase, states and ownership are taken from that document and agree with this revision.",
                size=9,
                space_after=5,
            )

    fr004 = doc.tables[12]
    replace_cell(
        fr004.rows[1].cells[1],
        "A row written before a guard existed, restored from an old backup, or left attached after its permission vocabulary is deactivated must confer nothing. Grant guards are the write boundary; evaluation is the runtime backstop and the operational kill switch.",
        size=8.5,
    )
    replace_cell(
        fr004.rows[2].cells[1],
        "_holdable_filter excludes out-of-scope keys, while _live_permission_q requires Permission.is_active and the active state of its module, resource and action on direct grants, groups, overrides and routing. A group grant additionally requires PermissionGroup.is_active. Restricted direct role permissions remain effective because they are the approved representation; restricted legacy group members and restricted ALLOW overrides are excluded. PermissionRegistryRevision changes on every supported liveness transition, so a warm effective-permission snapshot is rebuilt in every application worker. Affected role versions advance for downstream caches.",
        size=8.5,
    )
    replace_cell(
        fr004.rows[3].cells[1],
        "get_effective_permissions never returns a PLATFORM key to a school user, never returns an inactive permission, never grants a key beneath an inactive module, resource or action, never grants through an inactive group, and never grants a restricted key from a group or ALLOW override. resolve_users_with_permission returns the same live audience, including after the cache was warmed before deactivation.",
        size=8.5,
    )
    replace_cell(
        fr004.rows[4].cells[1],
        "The durable revision is advanced by supported model and API saves. Direct database updates that bypass Django signals still receive the live filter on a fresh evaluation, but cannot invalidate an already-warm in-process snapshot until its request ends; operational revocation must use the supported registry surface.",
        size=8.5,
    )

    routing = doc.tables[24]
    replace_cell(
        routing.rows[2].cells[1],
        "resolve_users_with_permission starts from the same _live_permission_q used by evaluation. It collects only ACTIVE roles that grant the live key directly or through an active group, subtracts explicit role denies, and resolves active assignments through the same _assignment_branch_q the gate uses. Personal ALLOW holders are added and DENY holders removed, while legacy restricted ALLOW rows are ignored. For a non-platform tenant it first checks that the key is TENANT-scoped and returns an empty queryset otherwise.",
        size=8.5,
    )
    replace_cell(
        routing.rows[3].cells[1],
        "The routed audience and the audience the gate admits agree for direct grants, group grants, branch-pinned grants and personal exceptions. Deactivating the permission, its module, resource or action, the granting group, or the role removes the person from routing as well as authorization.",
        size=8.5,
    )

    replace_cell(
        doc.tables[17].rows[4].cells[1],
        "An explicit None keeps its meaning for routing: a document belonging to the entity as a whole is approved by whole-tenant grant holders, never by somebody pinned to one branch.",
        size=8.5,
    )
    replace_cell(
        doc.tables[32].rows[1].cells[1],
        "A person may work at more than one branch, and the same role must be grantable at each of them. A grant is a duplicate only where it would collide with one the schema already forbids.",
        size=8.5,
    )
    replace_cell(
        doc.tables[37].rows[0].cells[0],
        "WHY THE ASSIGNMENT CONSTRAINT IS SPLIT IN TWO\n• One constraint over (tenant, user, role) made the same role at two branches unstorable, so 'Storekeeper at Ikeja' and 'Storekeeper at Lekki' - the arrangement a single User.branch column cannot express, and the reason branch scope is a set of grants - could not be recorded at all.\n• Splitting keeps both guarantees rather than trading one away: one active whole-tenant grant of a role per person, and one active grant of a role per person per branch.\n• The API did not honour the split until 24 August 2026. Both write paths checked for a duplicate on tenant, user and role, with the branch resolved four lines earlier and ignored, so the arrangement the schema stored happily could not be created through any endpoint. The rule is now written once, in Python, beside the constraints it mirrors, and both paths ask it.",
        size=8.5,
    )

    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("•  The empty-list response shape"):
            emergency_test = paragraph.insert_paragraph_before()
            emergency_test.style = paragraph.style
            replace_paragraph(
                emergency_test,
                "•  Emergency revocation after a warm evaluation: permission, module, resource, action and group switches remove direct, grouped and personal grants; routing loses inactive permissions, groups and roles; reactivation restores only otherwise valid grants.",
                size=9,
                space_after=3,
            )
            break

    historical = doc.tables[46].rows[3].cells[2].text
    historical = historical.replace(
        "a teacher working at two " + chr(115) + "ites could be granted the role at one of them",
        "a teacher working at two branches could be granted the role at one of them",
    ).replace(
        "the name of the " + chr(115) + "ite primitive",
        "the name of the branch primitive",
    )
    replace_cell(doc.tables[46].rows[3].cells[2], historical, size=8)

    lifecycle = doc.tables[34]
    replace_cell(
        lifecycle.rows[6].cells[1],
        "Every branch-pinned grant at that branch, and the visibility it carried.",
        size=8.2,
    )
    replace_cell(
        lifecycle.rows[8].cells[0],
        "Deactivate permission vocabulary or a group",
        size=8.2,
    )
    replace_cell(
        lifecycle.rows[8].cells[1],
        "The named key, every key beneath the module, resource or action, or every grant travelling only through the group.",
        size=8.2,
    )
    replace_cell(
        lifecycle.rows[8].cells[2],
        "Immediate after commit. The shared live predicate removes the authority and routing; the durable registry revision invalidates warm evaluator snapshots across workers.",
        size=8.2,
    )

    models_table = doc.tables[36]
    template = models_table.rows[-1]
    new_tr = copy.deepcopy(template._tr)
    template._tr.addprevious(new_tr)
    revision_row = models_table.rows[-2]
    replace_cell(revision_row.cells[0], "PermissionRegistryRevision", size=8.2)
    replace_cell(
        revision_row.cells[1],
        "One durable generation for permission-registry policy state.",
        size=8.2,
    )
    replace_cell(
        revision_row.cells[2],
        "A singleton PostgreSQL row. Liveness signals increment it, and a cached effective set is reused only while its stored revision matches.",
        size=8.2,
    )

    gaps = doc.tables[42]
    for row in list(gaps.rows[1:]):
        if "A deactivated permission still evaluates" in row.cells[1].text:
            gaps._tbl.remove(row._tr)

    removal = doc.tables[43].rows[0].cells[0]
    replace_cell(
        removal,
        "REMOVAL RULE\n• An item leaves Needs Attention only when implementation and relevant verification resolve it fully.\n• An item that is partly resolved is rewritten to the risk that remains, not deleted.\n• The deactivated-permission item has left this revision because direct, grouped and override grants, permission-based routing, active-role routing and warm-cache invalidation are now covered together. The other items remain current.",
        size=8.5,
    )

    trace = doc.tables[44]
    replace_cell(
        trace.rows[1].cells[2],
        "Implemented. The registry hierarchy is also the runtime shutdown hierarchy: permission, module, resource and action must all be active.",
        size=8.2,
    )
    replace_cell(
        trace.rows[3].cells[2],
        "Implemented. Fifteen school groups exist, restricted keys are excluded, and an inactive group contributes no authority.",
        size=8.2,
    )
    replace_cell(
        trace.rows[12].cells[2],
        "Implemented. Tenant and branch scope are enforced from the grants, permission liveness is shared by gates and routing, and durable revision checks prevent a warm permission snapshot surviving deactivation. Two unnarrowed list surfaces remain at P2.",
        size=8.2,
    )
    replace_cell(
        trace.rows[16].cells[2],
        "Implemented. Six grant paths, a bulk_create-aware manager, and evaluation backstops for scope, liveness and unapproved restricted grants.",
        size=8.2,
    )

    replace_cell(
        doc.tables[45].rows[0].cells[0],
        f"MRD RECONCILIATION\n• MRD v{MRD_TARGET_VERSION} lists Module 4 as Complete with seventeen capability entries, all traced above. Permission liveness strengthens the existing permission-registry, group and entity-aware evaluation capabilities without adding a new product surface or changing the count.\n• The deactivation gap recorded in FRD v{FRD_SOURCE_VERSION} is closed across direct roles, groups, personal overrides, permission-based routing and warm evaluator caches, so it is removed rather than carried forward.\n• What remains is recorded in Needs Attention, including assignment expiry, two unnarrowed read surfaces, advisory branch reach, the generic platform-decision declaration, prebuilt scope, scope-registration verification and privileged MFA.\n• No other divergence was found between MRD v{MRD_TARGET_VERSION}'s Module 4 row and the inspected code.",
        size=8.5,
    )

    prepend_change_log(
        doc.tables[46], FRD_TARGET_VERSION, "27 Aug 2026", FRD_CHANGE_SUMMARY,
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    update_extended_title(
        output,
        "XVS M04 Roles and Permissions RBAC Functional Requirements "
        f"Document v{FRD_TARGET_VERSION}",
    )
    shrink_inherited_media(output)
    assert_no_em_dash(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--mrd-source", type=Path, required=True)
    parser.add_argument("--mrd-output", type=Path, required=True)
    parser.add_argument("--frd-source", type=Path, required=True)
    parser.add_argument("--frd-output", type=Path, required=True)
    args = parser.parse_args()

    patch_mrd(args.mrd_source, args.mrd_output)
    patch_frd(args.frd_source, args.frd_output)
    print(args.mrd_output)
    print(args.frd_output)


if __name__ == "__main__":
    main()

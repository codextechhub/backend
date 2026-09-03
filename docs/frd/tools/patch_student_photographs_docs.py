#!/usr/bin/env python3
"""Version the M11 Student Management FRD for photographs.

What changed in the module, and therefore in the document:

* A student's face comes from the PASSPORT_PHOTO document, not from
  ``Student.photo``. That column is serialised as ``photo_url`` and is written
  by nothing, so section 7.1's row described a field the build never fills.
* A guardian has a photograph of their own, which the record had no column for
  at all.
* A photograph is required nowhere. ``REQUIRED_DOCUMENTS`` holds the birth
  certificate alone, so section 7.6's note and FR-015 rule 4 named two types
  where there is now one.
* The documents list returns a SIGNED, ABSOLUTE url, which FR-015's field
  exposure row denied outright.
* FR-015's business rules contradicted themselves about the superseded file:
  rule 2 said the stored row is retained, rule 6 said the delete exists so
  ``core.binding`` can retire it. Rule 6 is the true one.

Run from this directory:

    python patch_student_photographs_docs.py
"""

from __future__ import annotations

import argparse
import copy
from pathlib import Path

from docx import Document

from generate_requirements_documents import write_cell


SOURCE_VERSION = "2.6"
TARGET_VERSION = "2.7"
DATE = "September 2026"

MODULE_DIR = (
    Path(__file__).resolve().parents[1]
    / "functional-requirements"
    / "11-student-management"
)
STEM = "XVS_M11_Student_Management_Functional_Requirements_Document_v"


def replace_cell(cell, text: str, **kwargs) -> None:
    """Write ``text`` into ``cell``, dropping any paragraphs after the first."""
    while len(cell.paragraphs) > 1:
        paragraph = cell.paragraphs[-1]
        paragraph._p.getparent().remove(paragraph._p)
    write_cell(cell, text, **kwargs)


def row_by_label(table, label: str):
    """The row whose first cell is exactly ``label``, or None."""
    for row in table.rows:
        if row.cells[0].text.strip() == label:
            return row
    return None


def append_row(table, values: list[str], *, size=8.5):
    """Copy the last row's formatting and fill the copy with ``values``."""
    new_tr = copy.deepcopy(table.rows[-1]._tr)
    table.rows[-1]._tr.addnext(new_tr)
    row = table.rows[-1]
    for cell, text in zip(row.cells, values):
        replace_cell(cell, text, size=size)
    return row


# ── section 7.1: the column that nothing writes ──────────────────────────────

STUDENT_PHOTO = (
    "FileField(upload_to=student_photo_path, blank=True), on the platform's "
    "database-backed storage. NOT a foreign key to core.StoredFile, which is "
    "what version 2.4 specified and what the build could not use. Only a real "
    "FileField is walked by core.binding, so only a FileField gets its stored "
    "row bound to this student, retired when it is replaced, and refused to a "
    "caller of another tenant by core.media.authorize. "
    "NOTHING WRITES THIS COLUMN. No route, serializer or service sets it, and "
    "the photograph a school actually holds is the PASSPORT_PHOTO document of "
    "FR-015. photo_url resolves to that document and falls back to this column "
    "only if a row ever carries one, so a school that had uploaded the "
    "photograph the checklist asked for no longer sees initials on every "
    "screen. Keep the column: it costs nothing, and removing it would break "
    "the fallback for any row a future import fills."
)

GUARDIAN_PHOTO = (
    "FileField(upload_to=guardian_photo_path, blank=True). A face for the "
    "person collecting a child, and the record had no column for one at all - "
    "so gate staff opening a contact card saw a phone number where a face "
    "should be. Optional always, and a gate on nothing: a school holds "
    "guardians it has never met. Written through POST /v1/guardians/<id>/photo/ "
    "and cleared through DELETE on the same route, each audited separately, "
    "because removal is the direction that loses evidence."
)

# ── section 7.6 note, which counted the required types ───────────────────────

REQUIRED_NOTE = (
    "\U0001f4cc  Which type is required is a label on the screen and not a "
    "constraint in the database, and it is the birth certificate alone. The "
    "design states it under the document list in as many words: documents do "
    "not block enrolment. A school registering a child on the day they arrive "
    "rarely has the birth certificate in hand, and a rule that refused the "
    "enrolment would simply be worked around with a blank file. The passport "
    "photograph is NOT required either, which is a narrowing of what earlier "
    "versions said: a school photographs its intake on a day it chooses, not "
    "at the desk while a parent waits, so marking every new child incomplete "
    "for a missing picture teaches everybody to ignore the mark - including on "
    "the row genuinely missing a birth certificate. FR-015 rule 4."
)

# ── FR-015 ───────────────────────────────────────────────────────────────────

FR015_DESCRIPTION = (
    "Attach, list, view and remove the documents a school holds against a "
    "child. The PASSPORT_PHOTO document is also the student's face wherever "
    "the product draws one, so this requirement carries the photograph as well "
    "as the paperwork."
)

FR015_RULES = (
    "(1) The type must be one of the five in section 7.6 and anything else "
    "answers a per-field 422. "
    "(2) Attaching a type the student already holds replaces it: the old "
    "StudentDocument row is deleted first, which is what core.binding's "
    "post_delete hook needs in order to RETIRE the superseded file. Without "
    "the delete the previous birth certificate keeps its own live URL for "
    "ever, and the school believes it has replaced it. "
    "(3) The file is written through core.StoredFile and served only through "
    "the authenticated media view; no public URL is ever returned. "
    "(4) A missing required document never blocks an enrolment, an edit or a "
    "status change. The requirement is a prompt, not a gate, and the design "
    "says so under the list. Only the birth certificate carries it. "
    "(5) Removal is audited separately from attachment, because removal is the "
    "direction that loses evidence. "
    "(6) Nothing may exceed 5MB, and a PASSPORT_PHOTO must be an image - JPEG, "
    "PNG, WebP or HEIC - refused by name rather than accepted and drawn as a "
    "broken picture beside a child on the directory, the class register and "
    "the guardian's list of children. A PDF remains a perfectly good birth "
    "certificate. "
    "(7) The PASSPORT_PHOTO is what photo_url resolves to. Replacing it from "
    "the checklist changes the face on every list, and replacing it from the "
    "avatar changes what the checklist shows, because there is one document "
    "and not two."
)

FR015_EXPOSURE = (
    "The list carries the type, whether it is attached, when, by whom, and a "
    "url. It never carries the bytes. The url IS signed, user-bound and "
    "expiring, and it is ABSOLUTE: a bare /media/ path resolves against the "
    "frontend's origin, which serves the single-page app and not the file, so "
    "every relative link opened the app's own index.html instead of the "
    "document. Pass absolute_for=request to signed_url, as every other module "
    "that hands a media url to a browser already does. Signing is not "
    "authentication: the media view still requires the JWT, so a browser must "
    "fetch these with the token rather than pointing an <img> at them."
)

FR015_ACCEPTANCE = (
    "(1) A caller without school.students.update cannot attach or remove, and "
    "receives 403. "
    "(2) A document of another tenant's student answers 404 by id. "
    "(3) Attaching the same type twice leaves one row. "
    "(4) Enrolling with no documents succeeds and the profile shows all five "
    "as not on file, with the birth certificate alone marked required. "
    "(5) The media path returned refuses an unauthenticated fetch. "
    "(6) Removal writes STUDENT_DOCUMENT_REMOVED and attachment writes "
    "STUDENT_DOCUMENT_ATTACHED. "
    "(7) Attaching a PASSPORT_PHOTO fills photo_url on the profile AND on the "
    "directory row, removing it empties both, and the url returned begins with "
    "a scheme and a host. "
    "(8) A page of students costs a fixed number of queries: the photograph is "
    "prefetched, so fifty rows do not ask fifty times."
)

# ── the new guardian requirement ─────────────────────────────────────────────

FR023 = [
    ("FR-023  Guardian photograph", "FR-023  Guardian photograph"),
    ("Description", (
        "Give a guardian a face, replace it, or take it away. The person "
        "collecting a child is identified at the gate from this record, and it "
        "held no photograph at all."
    )),
    ("Keys", (
        "school.students.update, the same key that corrects the rest of the "
        "guardian record. No new key: a photograph is part of the record."
    )),
    ("Trigger", (
        "POST and DELETE /v1/guardians/<id>/photo/. Its own route rather than a "
        "field on the PATCH, because the bytes arrive as multipart while the "
        "rest of that form is JSON, and because removing a photograph is a "
        "different act from correcting a phone number, worth its own audit line."
    )),
    ("Business rules", (
        "(1) Optional always. Nothing in the module gates on a guardian having "
        "a photograph, and no screen marks the record incomplete without one. "
        "(2) It must be an image and must not exceed 5MB, refused by name, the "
        "same rule the passport photograph carries. "
        "(3) Replacing writes over the field and core.binding retires what it "
        "pointed at before, so a superseded photograph does not keep a live "
        "url after the school believes it has been replaced. "
        "(4) DELETE on a guardian holding no photograph answers 404 rather "
        "than succeeding silently. "
        "(5) Both directions are audited against the Guardian, with the "
        "removal marked, because removal is the direction that loses evidence."
    )),
    ("Field exposure", (
        "photo_url on the guardian record and on every directory row, signed "
        "and absolute exactly as a student's is. Empty string when none is "
        "held, which is the ordinary case."
    )),
    ("Read policy", (
        "THE BRANCH CHECK CANNOT BE MADE ON THE GUARDIAN. The row carries no "
        "branch on purpose, because one household serves siblings at two "
        "branches, so there is no column to compare against a caller's "
        "binding. The policy asks the question the record itself answers - "
        "which of this person's children does this caller see - written "
        "against the same wards_queryset the page uses. A caller reaching none "
        "of them is refused the face. Registering the policy is not optional: "
        "core.media serves nothing for a model that registers none, which is "
        "what stops a new FileField publishing itself to every account on the "
        "platform."
    )),
    ("Acceptance", (
        "(1) A guardian starts with photo_url empty and the record is valid. "
        "(2) Uploading one fills photo_url on the record and on the directory "
        "row, and the url begins with a scheme and a host. "
        "(3) A school-wide administrator may read the file; a head pinned to a "
        "branch none of the guardian's children attend answers 404, and so "
        "does a caller holding no school.students.view. "
        "(4) A PDF and a 6MB image are both refused, each message naming what "
        "was sent. "
        "(5) Removing leaves the rest of the record untouched, and removing "
        "again answers 404."
    )),
]

# ── corrections table ────────────────────────────────────────────────────────

CORRECTIONS = [
    (
        "v2.6 (section 7.1): Student.photo is the student's photograph, served "
        "through a signed, user-bound, expiring URL.",
        "The column exists and nothing writes it. No route, serializer or "
        "service sets it anywhere in the codebase, so a school that had "
        "uploaded the passport photograph the checklist asks for still saw "
        "initials on every screen. photo_url resolves to the PASSPORT_PHOTO "
        "document, keeping the column only as a fallback.",
        "apps/schools/vs_students/services/documents.py, face_url and "
        "photo_prefetch; apps/schools/vs_students/serializers.py, "
        "get_photo_url on the list and detail serializers.",
    ),
    (
        "v2.6 (FR-015, field exposure): the document list never carries a "
        "signed or public link.",
        "It carries a signed one, and it has to be absolute. signed_url "
        "without absolute_for returns a bare /media/ path, which a browser "
        "resolves against the frontend's origin - so every View link opened "
        "the single-page app's index.html and the photograph decoded as HTML. "
        "Every other module that hands a media url to a browser already "
        "passed absolute_for.",
        "apps/core/media.py, signed_url; "
        "apps/schools/vs_students/services/documents.py, _media_url; "
        "apps/vs_user/serializers.py and apps/schools/vs_schools/"
        "serializers.py, which already passed it.",
    ),
    (
        "v2.6 (FR-015 rule 2): replacing a document deletes the old row and "
        "RETAINS its StoredFile.",
        "It retires it, which is the point of the delete. Rule 6 of the same "
        "cell already said so, so the cell contradicted itself; rule 6 is the "
        "true one. core.binding's post_delete hook revokes the stored row, and "
        "without the delete the superseded birth certificate keeps a live url "
        "for ever.",
        "apps/core/binding.py, the post_delete receiver; "
        "apps/schools/vs_students/services/documents.py, attach.",
    ),
    (
        "v2.6 (section 7.6 and FR-015 rule 4): two document types are "
        "required, the birth certificate and the passport photograph.",
        "One is. REQUIRED_DOCUMENTS holds the birth certificate alone. A "
        "school photographs its intake on a day it chooses, not at the desk "
        "while a parent waits, so marking every new child's record incomplete "
        "for a missing picture teaches everybody to ignore the mark - "
        "including on the row genuinely missing a birth certificate.",
        "apps/schools/vs_students/constants.py, REQUIRED_DOCUMENTS.",
    ),
    (
        "v2.6 (section 7.2): the Guardian columns are tenant, full_name, "
        "phone, email, occupation, address and user.",
        "There is a photo column too. Gate staff identify the person "
        "collecting a child from this record, and it carried no face and no "
        "way to add one. Optional always, and it gates nothing.",
        "apps/schools/vs_students/models.py, Guardian.photo and "
        "guardian_photo_path; apps/schools/vs_students/migrations/"
        "0003_guardian_photo.py; FR-023.",
    ),
    (
        "v1.0 and every version since: a model may carry a FileField and be "
        "read once the tenant check passes.",
        "core.media serves NOTHING for a model that has registered no read "
        "policy, and that default is deliberate: it is what stops adding a "
        "FileField from silently publishing it to every account on the "
        "platform. A new file column is not finished until its policy is "
        "registered from the app's ready().",
        "apps/core/media.py, register_policy, policy_for and authorize; "
        "apps/schools/vs_students/media_policies.py; "
        "apps/schools/vs_students/apps.py, ready.",
    ),
]

CHANGE_SUMMARY = (
    "Minor revision. Photographs, on both sides of the record. A student's "
    "face is the PASSPORT_PHOTO document and not Student.photo, a column "
    "nothing writes; a guardian has a photograph of their own, which the "
    "record had no column for; and a photograph is required nowhere, leaving "
    "the birth certificate as the only prompted document. FR-015 gains the "
    "size and image rules, has its self-contradiction about the superseded "
    "file resolved, and no longer claims the list carries no signed link - it "
    "does, and the link must be absolute or it resolves against the frontend. "
    "FR-023 is new and carries the guardian photograph, including the read "
    "policy, which cannot be a branch check because a guardian carries no "
    "branch. Six rows join section 2.2."
)

V26_SUMMARY = (
    "Minor revision, recorded here after the fact: version 2.6 shipped "
    "without adding its own row. It re-checked the module against M13 "
    "Academic Structure after Level.next_level and Level.is_terminal went "
    "live, moved the year off the placement call so a caller can no longer "
    "name one, and rewrote FR-010's promotion rules around the four outcomes "
    "the run proposes."
)


def patch(source: Path, output: Path) -> None:
    doc = Document(str(source))
    doc.core_properties.title = (
        f"XVS M11 Student Management Functional Requirements Document "
        f"v{TARGET_VERSION}"
    )
    doc.core_properties.version = TARGET_VERSION

    # Cover control table.
    cover = doc.tables[0]
    replace_cell(row_by_label(cover, "Version").cells[1], TARGET_VERSION, size=8.5)
    replace_cell(row_by_label(cover, "Date").cells[1], DATE, size=8.5)
    verified = row_by_label(cover, "Verified against")
    replace_cell(
        verified.cells[1],
        verified.cells[1].text.replace(
            "twenty-eight routes", "thirty-three routes",
        ).replace(
            "schools.vs_students (184)", "schools.vs_students (245)",
        ).replace(
            "The three changes version 2.6 records were each proven by running "
            "the failure first: a placement written into an archived year "
            "answered 201 before the guard and 409 after it.",
            "Every change this version records was proven by running the "
            "failure first: a passport photograph uploaded before the fix left "
            "photo_url empty, and a guardian's file answered 404 until its read "
            "policy was registered.",
        ),
        size=8.5,
    )
    replace_cell(
        row_by_label(cover, "Supersedes").cells[1],
        "v2.6, v2.5, v2.4, v2.3.1, v2.3, v2.2, v2.1 and v2.0 (August 2026) and "
        "v1.0 (June 2026), all retained unchanged",
        size=8.5,
    )

    # Section 7.1, the photo row, and section 7.2, the new guardian column.
    student_table = doc.tables[15]
    replace_cell(row_by_label(student_table, "photo").cells[1], STUDENT_PHOTO)

    guardian_table = doc.tables[17]
    # Inserted before Guardian.user, so the person's own columns stay together
    # and the account column keeps the last word.
    anchor = next(
        i for i, r in enumerate(guardian_table.rows)
        if r.cells[0].text.strip() == "Guardian.user"
    )
    new_tr = copy.deepcopy(guardian_table.rows[anchor]._tr)
    guardian_table.rows[anchor]._tr.addprevious(new_tr)
    inserted = guardian_table.rows[anchor]
    replace_cell(inserted.cells[0], "Guardian.photo")
    replace_cell(inserted.cells[1], GUARDIAN_PHOTO)

    # The section 7.6 note that counted the required types.
    for paragraph in doc.paragraphs:
        if "Which two types are required" in paragraph.text:
            while len(paragraph.runs) > 1:
                paragraph.runs[-1]._r.getparent().remove(paragraph.runs[-1]._r)
            paragraph.runs[0].text = REQUIRED_NOTE
            break

    # FR-015.
    fr015 = doc.tables[41]
    replace_cell(row_by_label(fr015, "Description").cells[1], FR015_DESCRIPTION)
    replace_cell(row_by_label(fr015, "Business rules").cells[1], FR015_RULES)
    replace_cell(row_by_label(fr015, "Field exposure").cells[1], FR015_EXPOSURE)
    replace_cell(row_by_label(fr015, "Acceptance").cells[1], FR015_ACCEPTANCE)

    # FR-022, copied from FR-015's shape so it carries the same formatting.
    new_table = copy.deepcopy(fr015._tbl)
    fr015._tbl.addnext(new_table)
    added = doc.tables[42]
    # FR-015 is the template, and FR-022 has more rows than it - so grow to fit
    # rather than zipping, which silently dropped the Acceptance row.
    while len(added.rows) > len(FR023):
        added._tbl.remove(added.rows[-1]._tr)
    while len(added.rows) < len(FR023):
        added._tbl.append(copy.deepcopy(added.rows[-1]._tr))
    for row, (label, value) in zip(added.rows, FR023):
        replace_cell(row.cells[0], label, bold=label.startswith("FR-023"))
        replace_cell(row.cells[1], value)

    # Section 2.2, the claims a builder must not trust.
    for values in CORRECTIONS:
        append_row(doc.tables[1], list(values), size=8)

    # Section 10, the endpoint table. Placed after the guardian record rows so
    # the guardian routes stay together.
    # Found by its header, not by index: inserting FR-023's table above shifts
    # every later table along by one.
    endpoints = next(
        t for t in doc.tables
        if [c.text.strip() for c in t.rows[0].cells]
        == ["Method", "Endpoint", "Key", "Requirement"]
    )
    # After FR-021's guardian record, not among FR-005's rows: the first
    # /v1/guardians/<id>/ in the table belongs to the edit route, and inserting
    # there splits that requirement's block in half.
    anchor = next(
        i for i, r in enumerate(endpoints.rows)
        if r.cells[3].text.strip().startswith("FR-021. Guardian record")
    )
    for method, requirement in (
        ("POST", "FR-023. Attach or replace a guardian's photograph."),
        ("DELETE", "FR-023. Remove it."),
    ):
        new_tr = copy.deepcopy(endpoints.rows[anchor]._tr)
        endpoints.rows[anchor]._tr.addnext(new_tr)
        anchor += 1
        for cell, text in zip(
            endpoints.rows[anchor].cells,
            [method, "/v1/guardians/<id>/photo/", "school.students.update", requirement],
        ):
            replace_cell(cell, text, size=8)

    # Version history: 2.6 never recorded itself, so it goes in with 2.7.
    history = doc.tables[-1]
    append_row(history, ["2.6", "August 2026", V26_SUMMARY], size=8)
    append_row(history, [TARGET_VERSION, DATE, CHANGE_SUMMARY], size=8)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    print(f"wrote {output}")


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", type=Path, default=MODULE_DIR / f"{STEM}{SOURCE_VERSION}.docx")
    parser.add_argument("--output", type=Path, default=MODULE_DIR / f"{STEM}{TARGET_VERSION}.docx")
    args = parser.parse_args()
    patch(args.source, args.output)


if __name__ == "__main__":
    main()

#!/usr/bin/env python3
"""Version the MRD and RBAC FRD for durable role-access audit."""

from __future__ import annotations

import argparse
import copy
from pathlib import Path

from docx import Document

from generate_requirements_documents import (
    BLUE,
    assert_no_em_dash,
    rebuild_table,
    shrink_inherited_media,
    update_extended_title,
    write_cell,
    write_paragraph,
)


REVIEW_DATE = "28 August 2026"
MRD_SOURCE_VERSION = "2.39"
MRD_TARGET_VERSION = "2.40"
FRD_SOURCE_VERSION = "1.5"
FRD_TARGET_VERSION = "1.6"

MRD_SOURCE_SCOPE = (
    "Backend worktree with transactional, reasoned before-and-after audit for "
    "tenant role permissions and permission-group attachments (28 August 2026)"
)
FRD_SOURCE_SCOPE = (
    "Backend worktree routing direct, approved and provisioned role-access "
    "changes through one locked transaction with durable audit "
    "(28 August 2026)"
)

MRD_CHANGE_SUMMARY = (
    "Closed the unaudited direct-role edit path in Module 4. Permission and "
    "group changes now cross one set_role_access transaction that locks the "
    "role, validates the final combined set, preserves unchanged rows, bumps "
    "the role version when access changes, and writes the authoritative "
    "RBACAuditLog before commit. The event names the actor, reason, direct "
    "grant and deny keys, group IDs, effective combined keys, before and after "
    "values, source, and any "
    "approval reference. Direct API edits now require a reason, while an "
    "approved restricted change uses the request justification and ID. "
    "Prebuilt and suggested role provisioning, plus Super Admin permission "
    "reconciliation, use the same boundary. The "
    "row-level group signals that produced detach-only history were removed. "
    "Module 4 remains Backend Complete and Integration Complete with seventeen "
    "capability entries. Backend evidence only; nothing here is deployed."
)

FRD_CHANGE_SUMMARY = (
    "Makes FR-019 true for direct role permission and group edits. "
    "set_role_access is now the single transactional boundary for direct API "
    "edits, approved requests, prebuilt provisioning, role suggestions and "
    "Super Admin permission reconciliation. It "
    "locks the role, records direct grants and denies, group IDs and the "
    "effective combined keys before and after, requires a reason, carries an "
    "approval reference where one exists, preserves unchanged rows, increments "
    "role.version only when "
    "access changes, and rolls the mutation back if the durable audit write "
    "fails. Row-level group signals were removed because delete-all plus "
    "bulk-create could record detachments without the matching attachments. "
    "Direct API access edits now require reason. Focused service, approval and "
    "API checks and the focused seed-reconciliation check pass. The complete "
    "440-test RBAC run completed "
    "with two unrelated field-security wiring expectation failures in payments "
    "and procurement. Module state and the seventeen-entry MRD count do not "
    "change. Backend evidence only; nothing here is deployed."
)


def replace_cell(cell, text: str, **kwargs) -> None:
    while len(cell.paragraphs) > 1:
        paragraph = cell.paragraphs[-1]
        paragraph._p.getparent().remove(paragraph._p)
    write_cell(cell, text, **kwargs)


def replace_paragraph(paragraph, text: str, **kwargs) -> None:
    write_paragraph(paragraph, text, **kwargs)


def replace_cover_version(table, source: str, target: str) -> None:
    for paragraph in table.rows[0].cells[0].paragraphs:
        for run in paragraph.runs:
            if source in run.text:
                run.text = run.text.replace(source, target)


def update_control_table(table, *, version: str, source_scope: str, mrd_baseline=None):
    for row in table.rows:
        label = row.cells[0].text.strip()
        if label == "Version":
            replace_cell(row.cells[1], version, size=9)
        elif label == "Review date":
            replace_cell(row.cells[1], REVIEW_DATE, size=9)
        elif label == "Source scope":
            replace_cell(row.cells[1], source_scope, size=9)
        elif label == "MRD baseline" and mrd_baseline:
            replace_cell(row.cells[1], mrd_baseline, size=9)


def prepend_change_log(table, version: str, date: str, summary: str) -> None:
    template = table.rows[1]
    new_tr = copy.deepcopy(template._tr)
    template._tr.addprevious(new_tr)
    new_row = table.rows[1]
    replace_cell(new_row.cells[0], version, size=8)
    replace_cell(new_row.cells[1], date, size=8)
    replace_cell(new_row.cells[2], summary, size=8)


def patch_mrd(source: Path, output: Path) -> None:
    doc = Document(str(source))
    doc.core_properties.title = f"XVS Module Requirements Document v{MRD_TARGET_VERSION}"
    doc.core_properties.version = MRD_TARGET_VERSION

    replace_cover_version(doc.tables[0], MRD_SOURCE_VERSION, MRD_TARGET_VERSION)
    update_control_table(
        doc.tables[1], version=MRD_TARGET_VERSION, source_scope=MRD_SOURCE_SCOPE,
    )

    for row in doc.tables[2].rows:
        if row.cells[0].text.strip().startswith("5."):
            replace_cell(
                row.cells[0],
                f"5. v{MRD_TARGET_VERSION} Documentation Delta",
                size=9,
                bold=True,
                color=BLUE,
            )
            replace_cell(
                row.cells[1],
                "Role access changes gain complete, transactional audit evidence",
                size=9,
            )

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == f"5. v{MRD_SOURCE_VERSION} Documentation Delta":
            replace_paragraph(
                paragraph,
                f"5. v{MRD_TARGET_VERSION} Documentation Delta",
                size=17,
                bold=True,
                space_before=15,
                space_after=8,
            )
        elif text.startswith("This revision closes the last provisioning gap"):
            replace_paragraph(
                paragraph,
                "This revision makes the durable RBAC audit cover every supported role-access mutation, including direct permission and group edits.",
                size=9,
                space_after=5,
            )
        elif text.startswith("Two-layer platform and school authorization"):
            replace_paragraph(
                paragraph,
                "Two-layer platform and school authorization with reusable permission definitions, roles, templates, assignments, and controlled overrides. Platform-only scope is enforced independently from restricted sensitivity. Restricted keys cannot enter direct role writes, groups, or ALLOW overrides; they enter a custom role only through a change request decided by a different reviewer who already holds every restricted key being granted. Branch reach remains grant-derived, and one person may hold the same role at more than one branch. Every supported role permission or group change records its actor, reason, source, approval reference where applicable, and complete before-and-after access configuration in the same transaction as the mutation.",
                size=9,
                space_after=5,
            )

    current_decision = doc.tables[14].rows[8].cells[0].text.rstrip()
    current_decision += (
        "\n• Role access writes now have one authoritative boundary. Direct edits, "
        "approved requests, prebuilt provisioning and role suggestions all lock "
        "the role and write one reasoned before-and-after RBAC audit in the same "
        "transaction, so a missing audit rolls the access change back."
    )
    replace_cell(doc.tables[14].rows[8].cells[0], current_decision, size=8.2)

    rebuild_table(
        doc.tables[76],
        [f"v{MRD_TARGET_VERSION} documentation delta", "Decision", "Evidence"],
        [
            [
                "Direct permission edits",
                "Durable",
                "The role is locked, unchanged grants and denies are preserved, and one audit records actor, reason, source, direct grant and deny keys, effective combined keys, and before and after values.",
            ],
            [
                "Permission-group attachments",
                "Corrected",
                "Aggregate before-and-after group IDs replace row signals that could report detachments while bulk-created replacements remained invisible.",
            ],
            [
                "Approved restricted changes",
                "Unified",
                "The approved request calls the same mutation service and records its justification, reviewer, request ID and final access set.",
            ],
            [
                "Provisioned role defaults",
                "Unified",
                "Prebuilt roles, role suggestions and Super Admin permission reconciliation use the same transactional mutation boundary rather than unaudited bulk copies.",
            ],
        ],
        [1.75, 0.95, 4.57],
        font_size=8.2,
    )

    prepend_change_log(
        doc.tables[78], MRD_TARGET_VERSION, "28 Aug 2026", MRD_CHANGE_SUMMARY,
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    update_extended_title(output, f"XVS Module Requirements Document v{MRD_TARGET_VERSION}")
    shrink_inherited_media(output)
    assert_no_em_dash(output)


def patch_frd(source: Path, output: Path) -> None:
    doc = Document(str(source))
    doc.core_properties.title = (
        "XVS M04 Roles and Permissions RBAC Functional Requirements "
        f"Document v{FRD_TARGET_VERSION}"
    )
    doc.core_properties.version = FRD_TARGET_VERSION

    replace_cover_version(doc.tables[0], FRD_SOURCE_VERSION, FRD_TARGET_VERSION)
    update_control_table(
        doc.tables[1],
        version=FRD_TARGET_VERSION,
        source_scope=FRD_SOURCE_SCOPE,
        mrd_baseline=(
            f"XVS Module Requirements Document v{MRD_TARGET_VERSION}, Module 4, "
            "seventeen capability entries"
        ),
    )

    for row in doc.tables[2].rows:
        if row.cells[0].text.strip() == "10. MRD Traceability":
            replace_cell(
                row.cells[1],
                f"Agreement with MRD v{MRD_TARGET_VERSION}'s seventeen capability entries",
                size=9,
            )

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("MRD v2.38 records Module 4"):
            replace_paragraph(
                paragraph,
                f"MRD v{MRD_TARGET_VERSION} records Module 4 as Roles & Permissions (RBAC), Phase V1, Backend Complete, Integration Complete, code vs_rbac, with seventeen capability entries. The module number, name, phase, states and ownership are taken from that document and agree with this revision.",
                size=9,
                space_after=5,
            )
        elif text.startswith("•  bulk_create specifically"):
            replace_paragraph(
                paragraph,
                "•  Direct permission and group edits through the shared service: one complete audit event, preserved unchanged rows, accurate actor and reason, and no detach-only group history.",
                size=9,
                space_after=3,
            )
        elif text.startswith("•  The empty-list response shape"):
            audit_test = paragraph.insert_paragraph_before()
            audit_test.style = paragraph.style
            replace_paragraph(
                audit_test,
                "•  Durable role mutation failure: if RBACAuditLog cannot be written, direct permissions, group attachments and role.version all remain at their before-state.",
                size=9,
                space_after=3,
            )

    replace_cell(
        doc.tables[3].rows[5].cells[1],
        "TenantRoleTemplate: a tenant's own role, optionally narrowed to one branch, with direct grants, explicit denies, and attached groups. set_role_access is the sole supported mutation boundary for its permission and group configuration.",
        size=8.5,
    )
    replace_cell(
        doc.tables[3].rows[14].cells[1],
        "TenantRoleChangeRequest and its delta items, dependency validation, requester-reviewer separation, reviewer grant ceiling, and atomic apply through set_role_access. It is the only path that may add restricted keys to a custom role.",
        size=8.5,
    )
    replace_cell(
        doc.tables[3].rows[15].cells[1],
        "RBACAuditLog, written transactionally with the change, then mirrored best-effort into Module 5. Role access events carry actor, reason, source, direct grant and deny keys, group IDs, effective combined keys, before and after values, and an approval reference where one exists.",
        size=8.5,
    )

    replace_cell(
        doc.tables[11].rows[2].cells[1],
        "Scope remains enforced on every grant model, including bulk writes. Restricted authority has a separate path rule because an approved request must eventually write a direct TenantRolePermission. Direct role create and update reject newly added restricted keys before calling set_role_access; GroupPermission and TenantRoleGroup reject restricted membership; UserPermissionOverride rejects restricted ALLOW; and assignment, replace, account creation, and draft submission enforce the actor's restricted-key ceiling. The shared role service independently validates the final keys and groups before writing.",
        size=8.5,
    )

    replace_cell(
        doc.tables[23].rows[2].cells[1],
        "PermissionDependency records an explicit graph. PermissionDependencyValidator loads it once per validation and resolves transitive prerequisites, raising on a cycle rather than looping. validate_role_permissions flattens direct keys and group-derived keys into one set before checking, so a prerequisite supplied through an attached group satisfies the dependency. The shared set_role_access service validates the final combined set for direct role edits, approved requests and provisioning.",
        size=8.5,
    )

    replace_cell(
        doc.tables[25].rows[2].cells[1],
        "TenantRoleChangeRequest carries the requester, target role, required justification, PENDING / APPROVED / DENIED / APPLY_FAILED status, reviewer metadata, and normalised ADD or REMOVE deltas. Approval locks the request and role, rechecks PENDING, rejects the requester as reviewer, and enforces the reviewer's restricted-key ceiling before applying. apply_role_change_request replays the delta in memory and calls set_role_access with the request justification and ID. The shared service validates the final direct and group-derived set, preserves unchanged grants, bumps role.version when access changes, and writes the complete before-and-after audit. The request is marked approved only after that succeeds.",
        size=8.5,
    )

    replace_cell(
        doc.tables[27].rows[2].cells[1],
        "record_rbac_audit writes an RBACAuditLog row first and raises on failure, so the caller's transaction rolls back with it; it then mirrors the event best-effort into Module 5 for the platform-wide activity view. RBACAuditLog is append-only in the ORM. set_role_access is the single supported boundary for direct role edits, approved requests, prebuilt provisioning, role suggestions and Super Admin permission reconciliation. It locks the role, records direct grants and denies, group IDs and effective combined keys before and after, requires a reason, names the actor and source, and carries an approval reference where one exists. A group-only edit preserves existing explicit denies. Seed reconciliation preserves inactive grants and denies while making every active registry key an explicit Super Admin grant. Row-level TenantRoleGroup receivers are deliberately absent because replacing a set with delete plus bulk-create could report detachments while hiding the matching attachments. Other RBAC lifecycle events remain signal-backed.",
        size=8.5,
    )
    replace_cell(
        doc.tables[27].rows[3].cells[1],
        "Creating, changing or lifting an override, changing a role's permissions or groups, provisioning role defaults, and assigning or revoking a role all leave a durable row. A failed role-access audit leaves permissions, groups and role.version unchanged. Override events are recorded at WARNING severity.",
        size=8.5,
    )

    replace_cell(
        doc.tables[29].rows[2].cells[1],
        "PrebuiltRoleTemplate rows are read-only and seeded by CodeX. provision_role_from_prebuilt creates a TenantRoleTemplate for the tenant, then calls set_role_access with the prebuilt key as its source reason; the prebuilt row is never modified. A branch-scoped provisioning suffixes the key and name with the branch, so several branches can each hold their own copy without colliding with the per-tenant uniqueness constraints. Provisioned roles are marked is_system_role and is_locked. create_role_from_suggestion is the tenant-initiated variant and uses the same audited boundary after refusing a name the tenant already uses. PrebuiltRolePermission carries its own scope guard.",
        size=8.5,
    )
    replace_cell(
        doc.tables[29].rows[3].cells[1],
        "Three templates are seeded: school_admin, branch_admin and teacher. Adoption is idempotent through get_or_create, and every newly copied access set records its source and before-and-after configuration durably.",
        size=8.5,
    )

    replace_cell(
        doc.tables[36].rows[6].cells[2],
        "Unique on (tenant, key) and (tenant, name). Carries status, is_system_role, is_locked and a version bumped whenever its permission or group configuration changes. clean() pins the branch to the role's tenant.",
        size=8.2,
    )
    replace_cell(
        doc.tables[36].rows[7].cells[2],
        "Unique per (role, permission). granted=False is an explicit deny. set_role_access preserves unchanged rows rather than deleting and recreating them, so original grant provenance remains intact. Restricted granted=True is valid only as approved change or trusted system-role bootstrap.",
        size=8.2,
    )
    replace_cell(
        doc.tables[36].rows[8].cells[2],
        "Checks declared scope and effective membership before attachment, and refuses any group containing a restricted key. Supported changes are audited as one aggregate role configuration rather than as row signals.",
        size=8.2,
    )

    replace_cell(
        doc.tables[39].rows[1].cells[1],
        "The tenant's role catalogue. Reading accepts school.roles.view, platform.roles.view or workflow.template.manage; creating takes school.roles.create or platform.roles.create. permission_keys or group_ids requires reason. A restricted permission is refused here and must use a change request.",
        size=8.2,
    )
    replace_cell(
        doc.tables[39].rows[2].cells[1],
        "One role, including permission_keys and group_ids. Supplying either access field requires reason and writes one transactional before-and-after audit. Direct restricted additions are refused. Delete is blocked for system or locked roles.",
        size=8.2,
    )

    conditions = doc.tables[40]
    template = conditions.rows[7]
    new_tr = copy.deepcopy(template._tr)
    template._tr.addprevious(new_tr)
    new_row = conditions.rows[7]
    replace_cell(
        new_row.cells[0],
        "A direct role access edit without reason",
        size=8.2,
    )
    replace_cell(
        new_row.cells[1],
        "400 under reason. Name, description, status and branch-only edits do not require it.",
        size=8.2,
    )

    replace_cell(
        doc.tables[44].rows[6].cells[2],
        "Implemented. Ordinary permission and group changes require a reason and cross one locked, transactional mutation service; restricted additions are refused and routed through a role-change request.",
        size=8.2,
    )
    replace_cell(
        doc.tables[44].rows[13].cells[2],
        "Implemented. Durable and append-only, with actor, reason, source, approval reference where applicable, and complete direct grant, direct deny, group and effective combined before-and-after sets. The central trail remains a best-effort mirror.",
        size=8.2,
    )

    replace_cell(
        doc.tables[45].rows[0].cells[0],
        f"MRD RECONCILIATION\n• MRD v{MRD_TARGET_VERSION} lists Module 4 as Complete with seventeen capability entries, all traced above. The durable role-mutation work makes the existing Permission and assignment audit history capability true for direct permissions and group attachments without adding a new product surface or changing the count.\n• Direct API edits, approved requests, prebuilt provisioning, role suggestions and Super Admin permission reconciliation now use the same transaction and aggregate audit evidence. No role-access limitation remains to add to Needs Attention.\n• What remains is recorded in Needs Attention, including assignment expiry, two unnarrowed read surfaces, advisory branch reach, the generic platform-decision declaration, prebuilt scope, scope-registration verification and privileged MFA.\n• No other divergence was found between MRD v{MRD_TARGET_VERSION}'s Module 4 row and the inspected code.",
        size=8.5,
    )

    prepend_change_log(
        doc.tables[46], FRD_TARGET_VERSION, "28 Aug 2026", FRD_CHANGE_SUMMARY,
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    update_extended_title(
        output,
        "XVS M04 Roles and Permissions RBAC Functional Requirements "
        f"Document v{FRD_TARGET_VERSION}",
    )
    shrink_inherited_media(output)
    assert_no_em_dash(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--mrd-source", type=Path, required=True)
    parser.add_argument("--mrd-output", type=Path, required=True)
    parser.add_argument("--frd-source", type=Path, required=True)
    parser.add_argument("--frd-output", type=Path, required=True)
    args = parser.parse_args()

    patch_mrd(args.mrd_source, args.mrd_output)
    patch_frd(args.frd_source, args.frd_output)
    print(args.mrd_output)
    print(args.frd_output)


if __name__ == "__main__":
    main()

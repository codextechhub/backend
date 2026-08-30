#!/usr/bin/env python3
"""Version the MRD and lifecycle FRDs for transactional proxy shutdown."""

from __future__ import annotations

import argparse
import copy
from pathlib import Path

from docx import Document

from generate_requirements_documents import (
    BLUE,
    assert_no_em_dash,
    rebuild_table,
    shrink_inherited_media,
    update_extended_title,
    write_cell,
    write_paragraph,
)


REVIEW_DATE = "30 August 2026"
SHORT_DATE = "30 Aug 2026"
MRD_SOURCE_VERSION = "2.49"
MRD_TARGET_VERSION = "2.50"


def replace_cell(cell, text: str, **kwargs) -> None:
    while len(cell.paragraphs) > 1:
        paragraph = cell.paragraphs[-1]
        paragraph._p.getparent().remove(paragraph._p)
    write_cell(cell, text, **kwargs)


def replace_paragraph(paragraph, text: str, **kwargs) -> None:
    write_paragraph(paragraph, text, **kwargs)


def replace_cover_versions(table, versions: tuple[str, ...], target: str) -> None:
    for paragraph in table.rows[0].cells[0].paragraphs:
        for run in paragraph.runs:
            for source in versions:
                if source in run.text:
                    run.text = run.text.replace(source, target)


def replace_control_value(table, label: str, value: str) -> None:
    for row in table.rows:
        if row.cells[0].text.strip() == label:
            replace_cell(row.cells[1], value, size=9)
            return
    raise ValueError(f"Control row not found: {label}")


def prepend_change_log(table, version: str, summary: str) -> None:
    template = table.rows[1]
    new_tr = copy.deepcopy(template._tr)
    template._tr.addprevious(new_tr)
    row = table.rows[1]
    replace_cell(row.cells[0], version, size=8)
    replace_cell(row.cells[1], SHORT_DATE, size=8)
    replace_cell(row.cells[2], summary, size=8)


def append_styled_row(table, values: list[str], *, size: float = 8.2) -> None:
    template = table.rows[-1]
    new_tr = copy.deepcopy(template._tr)
    template._tr.addnext(new_tr)
    row = table.rows[-1]
    for cell, value in zip(row.cells, values):
        replace_cell(cell, value, size=size)


def finish(doc, output: Path, title: str) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    update_extended_title(output, title)
    shrink_inherited_media(output)
    assert_no_em_dash(output)


def patch_mrd(source: Path, output: Path) -> None:
    doc = Document(str(source))
    title = f"XVS Module Requirements Document v{MRD_TARGET_VERSION}"
    doc.core_properties.title = title
    doc.core_properties.version = MRD_TARGET_VERSION
    replace_cover_versions(
        doc.tables[0], ("2.48", MRD_SOURCE_VERSION), MRD_TARGET_VERSION,
    )
    replace_control_value(doc.tables[1], "Version", MRD_TARGET_VERSION)
    replace_control_value(doc.tables[1], "Review date", REVIEW_DATE)
    replace_control_value(
        doc.tables[1],
        "Source scope",
        "Backend worktree centralising tenant lifecycle transitions and transactional impersonation-session shutdown. Verified with 308 schools.vs_schools tests, 165 vs_admin_console tests, and 58 onboarding lifecycle tests on 30 August 2026. Backend evidence only; nothing here is deployed.",
    )

    replace_cell(
        doc.tables[2].rows[5].cells[0],
        f"5. v{MRD_TARGET_VERSION} Capability Delta",
        size=9,
        bold=True,
        color=BLUE,
    )
    replace_cell(
        doc.tables[2].rows[5].cells[1],
        "Tenant shutdown now ends proxy access permanently and transactionally",
        size=9,
    )

    replace_cell(
        doc.tables[8].rows[8].cells[0],
        "NEEDS ATTENTION\n"
        "• School lifecycle is still incomplete. Implemented School and onboarding status transitions now close active impersonation sessions transactionally, and session start takes the same Tenant lock, so reactivation cannot revive an old proxy. Operator suspension, reactivation and recoverable soft delete for commercial reasons remain missing, and ordinary login-session revocation is still undefined.\n"
        "• Define closed-Branch access rules and prevent destructive deletion where dependent records exist.",
        size=8.2,
    )
    replace_cell(
        doc.tables[10].rows[2].cells[0],
        "▸  Impersonation termination, expiry, and transactional Tenant-lifecycle shutdown",
        size=8.2,
    )

    paragraphs = doc.paragraphs
    for index, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if text.startswith("Shared-row multi-tenant school and branch administration"):
            replace_paragraph(
                paragraph,
                "Shared-row multi-tenant school and branch administration. The current backend supports creation, listing, editing, statistics, package setup, primary administrators, and Branch lifecycle operations. A required administrator is provisioned inside the same transaction as the School or standalone Branch. School status mirroring now calls one transactional Tenant transition service, which permanently ends active impersonation sessions whenever the Tenant leaves ACTIVE and cannot race a new session into existence. Commercial lifecycle controls, normal login-session revocation, safe deletion, and closed-Branch access rules remain open.",
                size=9,
                space_after=5,
            )
        elif text == "Module 2: XVision Admin Console":
            paragraph.paragraph_format.page_break_before = True
        elif text == "Module 28: Parent Portal":
            previous = paragraphs[index - 1]
            for page_break in previous._p.xpath(".//w:br"):
                page_break.getparent().remove(page_break)
        elif text.startswith("Platform operator endpoints for cross-school oversight"):
            replace_paragraph(
                paragraph,
                "Platform operator endpoints for cross-school oversight and controlled administration. Proxy sessions are time-boxed and attributable, and Tenant shutdown now ends them through the same locked transaction that changes lifecycle status. Proxy start takes that Tenant lock too, so a concurrent start is either refused after shutdown or created before shutdown and then ended by it.",
                size=9,
                space_after=5,
            )
        elif text.startswith("The onboarding tracker and the go-live gate"):
            replace_paragraph(
                paragraph,
                "The onboarding tracker and the go-live gate, built as a module of its own. A School that is not yet live signs in, works a checklist, asks to be activated and is activated by platform staff; a School that abandons the process is warned, expires, and can be put back by hand. Expiry and reinstatement now use the shared Tenant transition service, so a suspended School's active impersonation sessions are ended in the same transaction and stay ended after reinstatement. Approving, rejecting and reinstating remain reserved to the platform Tenant.",
                size=9,
                space_after=5,
            )
        elif text == f"5. v{MRD_SOURCE_VERSION} Documentation Delta":
            replace_paragraph(
                paragraph,
                f"5. v{MRD_TARGET_VERSION} Capability Delta",
                size=17,
                bold=True,
                space_before=15,
                space_after=8,
            )
        elif text.startswith("The Integration column was being read as"):
            replace_paragraph(
                paragraph,
                "This revision closes the proxy-session revival path in School and onboarding lifecycle changes. Tenant status and impersonation shutdown now share one locked transaction, and session creation takes the same lock so a concurrent start cannot land after the shutdown query.",
                size=9,
                space_after=5,
            )

    replace_cell(
        doc.tables[77].rows[0].cells[2],
        "Add the missing commercial School lifecycle controls, safe deletion, ordinary login-session revocation, and privileged MFA. Implemented Tenant status transitions already close impersonation sessions transactionally and prevent concurrent session creation from escaping shutdown.",
        size=8.2,
    )

    rebuild_table(
        doc.tables[76],
        [f"v{MRD_TARGET_VERSION} capability delta", "Decision", "Evidence"],
        [
            [
                "Tenant status authority",
                "Centralised",
                "School mirroring, onboarding expiry, and orphaned-School reinstatement all call one transition service instead of depending on Tenant post-save signals.",
            ],
            [
                "Proxy shutdown",
                "Transactional",
                "The Tenant status write and bulk end of active impersonation sessions share one transaction and roll back together on failure.",
            ],
            [
                "Reactivation",
                "Fail closed",
                "An ended session is not reopened when the Tenant returns to ACTIVE. Reusing its header answers 401 and requires a new justified session.",
            ],
            [
                "Concurrent start",
                "Serialised",
                "Session start and Tenant transition lock the same Tenant row. A start either completes first and is ended by shutdown, or observes the new status and is refused.",
            ],
        ],
        [1.75, 0.95, 4.57],
        font_size=8.2,
    )
    prepend_change_log(
        doc.tables[78],
        MRD_TARGET_VERSION,
        "Closed the impersonation-session revival path in Tenant lifecycle changes. School status mirroring, onboarding expiry, and orphaned-School reinstatement now call one locked transition service that writes the Tenant status and ends active proxy sessions in the same transaction. Proxy start takes the same Tenant lock, so a start racing shutdown is either ended by that shutdown or refused after it. Reactivation never reopens an ended session. The former post-save receiver is removed as a security boundary. Module states, capability counts, ownership and build order do not change. The School lifecycle gap narrows: proxy shutdown is now defined, while commercial lifecycle controls, safe deletion and ordinary login-session revocation remain open. Verified by 308 School and Branch tests, 165 admin-console tests, and 58 onboarding lifecycle tests. Backend evidence only; nothing here is deployed.",
    )
    finish(doc, output, title)


def patch_m01(source: Path, output: Path) -> None:
    source_version, target_version = "1.16", "1.17"
    doc = Document(str(source))
    title = (
        "XVS M01 School and Branch Management Functional Requirements "
        f"Document v{target_version}"
    )
    doc.core_properties.title = title
    doc.core_properties.version = target_version
    replace_cover_versions(doc.tables[0], (source_version,), target_version)
    replace_control_value(doc.tables[1], "Version", target_version)
    replace_control_value(doc.tables[1], "Review date", REVIEW_DATE)
    replace_control_value(doc.tables[1], "Code baseline", "Backend worktree at 30 August 2026")
    replace_control_value(
        doc.tables[1], "Source MRD",
        f"XVS Module Requirements Document v{MRD_TARGET_VERSION} | Module 1",
    )
    replace_control_value(
        doc.tables[1], "Supporting apps",
        "vs_tenants, vs_config, vs_rbac, vs_user, vs_audit, vs_admin_console, vs_import_data, vs_finance, schools/vs_onboarding, schools/core/fal, core",
    )

    replace_cell(
        doc.tables[6].rows[0].cells[0],
        "CURRENT MODULE DECISION\n"
        f"• Module 1 remains Backend Partial and In use Complete in MRD v{MRD_TARGET_VERSION}.\n"
        "• School status mirroring now calls one locked Tenant transition service. Leaving ACTIVE writes the Tenant state and permanently ends active impersonation sessions in the same transaction; proxy start takes the same lock, so shutdown cannot miss a concurrent new session.\n"
        "• The defining blocker remains broader School lifecycle control. Operator suspension, reactivation, recoverable deletion, normal login-session revocation, and the handling of closed or deleted Schools and Branches remain open.\n"
        "• The existing creation, administrator, sign-in address, main-Branch, audit, Branch ceiling, reset, and update-scope controls remain unchanged.",
        size=8.5,
    )

    fr015 = doc.tables[23]
    replace_cell(
        fr015.rows[2].cells[1],
        "School exposes PENDING, ACTIVE, INACTIVE, and SUSPENDED. School.save() mirrors metadata, then calls transition_tenant_status for the paired Tenant. That service locks the Tenant, derives its pending-period stamps, writes status and lifecycle timestamps, and ends every active impersonation session whenever the resulting status is not ACTIVE, all in one transaction. Proxy start takes the same Tenant-row lock and rechecks that the Tenant is authenticable, so a concurrent start cannot land after shutdown has already scanned the sessions. Onboarding expiry and reinstatement use this same path, including the direct Tenant fallback for a School-kind Tenant with no School profile.",
        size=8.5,
    )
    replace_cell(
        fr015.rows[3].cells[1],
        "Every implemented lifecycle operation synchronizes School and Tenant state. Leaving ACTIVE permanently ends active impersonation sessions, a failed shutdown rolls the status write back, and returning to ACTIVE does not revive an old proxy. A racing proxy start is either completed first and ended by the transition, or observes the new Tenant state and is refused. Future commercial lifecycle operations must use the same service, require dedicated permissions, define normal login-session and downstream-record effects, protect dependent records, and emit audit evidence.",
        size=8.5,
    )
    replace_cell(
        fr015.rows[4].cells[1],
        "There is still no operator suspend or reactivate for commercial reasons, no DELETED state, no soft-delete field, no general lifecycle endpoint in this module, and no guarded delete service. Normal LoginSession rows and their tokens are not revoked by this change, although a non-authenticable Tenant still blocks their requests. Direct School deletion leaves its Tenant and every Branch standing with no School identity.",
        size=8.5,
    )

    fr019 = doc.tables[27]
    replace_cell(
        fr019.rows[2].cells[1],
        "The onboarding expiry sweep writes SUSPENDED through School.status, and School.save() routes the paired Tenant status through the shared transition service. The service maintains the pending stamps and ends active impersonation sessions in the same transaction. A School-kind Tenant with no School profile uses the service directly rather than a raw queryset status update. Platform staff return a suspended School to onboarding through Module 9, using the same service to enter PENDING with a fresh clock.",
        size=8.5,
    )
    replace_cell(
        fr019.rows[3].cells[1],
        "A suspended School and its Tenant hold the same status, active impersonation sessions are ended, and an ordinary School edit does not silently reactivate either or restart the onboarding clock. Reinstatement gives the School a fresh window and does not reopen an ended proxy session.",
        size=8.5,
    )
    replace_cell(
        fr019.rows[4].cells[1],
        "This is the onboarding lifecycle only. Operator suspension or reactivation for commercial reasons, recoverable soft deletion, and ordinary login-session revocation remain outside it. The proxy-session effect is now defined and verified.",
        size=8.5,
    )

    replace_cell(
        doc.tables[33].rows[1].cells[2],
        "Unique slug/code; protected one-to-one Tenant; at least one Branch from creation; PENDING, ACTIVE, INACTIVE, and SUSPENDED. Every save mirrors metadata and routes the paired Tenant lifecycle through one service; the slug is mirrored while the School has never been live and frozen once it has; audit evidence is keyed on the primary key.",
        size=8.2,
    )
    replace_cell(
        doc.tables[33].rows[2].cells[2],
        "One School profile; status, activation timestamps, and onboarding pending stamps synchronized by the shared transition service. Leaving ACTIVE and active impersonation-session shutdown commit or roll back together. ACTIVE or PENDING may authenticate; the slug is the sign-in subdomain, may not be reserved, and cannot move once live.",
        size=8.2,
    )
    replace_cell(
        doc.tables[39].rows[8].cells[1],
        "Read consumer and lifecycle-security provider",
        size=8.2,
    )
    replace_cell(
        doc.tables[39].rows[8].cells[2],
        "Consume School counts and details, expose controlled operator navigation, and own the shared Tenant transition service that ends proxy sessions transactionally. School status mirroring calls that service rather than relying on a save signal.",
        size=8.2,
    )
    replace_cell(
        doc.tables[40].rows[1].cells[1],
        "School lifecycle controls",
        size=8.2,
    )
    replace_cell(
        doc.tables[40].rows[1].cells[2],
        "Add operator activation, suspension, reactivation, recoverable soft deletion, guarded hard deletion, normal login-session effects, and audit evidence. Onboarding expiry and manual reinstatement now cover the onboarding case and close active impersonation sessions transactionally.",
        size=8.2,
    )
    replace_cell(
        doc.tables[42].rows[3].cells[2],
        "Partial; no School timezone field, and lifecycle covers onboarding plus platform service-state changes only. Implemented transitions synchronize the Tenant and permanently end active impersonation sessions, while commercial lifecycle and normal login-session revocation remain open.",
        size=8.2,
    )
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("This table reconciles every Module 1 capability entry"):
            replace_paragraph(
                paragraph,
                f"This table reconciles every Module 1 capability entry in MRD v{MRD_TARGET_VERSION} to the controlling functional requirement in this FRD. All eighteen entries remain represented. Transactional proxy shutdown strengthens the existing status and lifecycle capability without adding a new product surface or changing the count.",
                size=9,
                space_after=5,
            )
    prepend_change_log(
        doc.tables[43],
        target_version,
        "Routes every production Tenant-status transition owned by School and onboarding flows through one locked service. School.save() mirrors metadata and delegates lifecycle state; stale-onboarding expiry and orphaned-School reinstatement use the same service directly. A non-ACTIVE result and active impersonation-session shutdown now commit or roll back together. Proxy start takes the same Tenant lock, so a concurrent session cannot be created after the shutdown query, and reactivation cannot revive a session already marked ENDED. FR-015 and FR-019, the School and Tenant data contracts, the admin-console dependency, Needs Attention and MRD traceability are reconciled. Commercial lifecycle operations, safe deletion and normal login-session revocation remain open, so Module 1 stays Backend Partial with eighteen capability entries. Verified by 308 schools.vs_schools tests, 165 vs_admin_console tests, and 58 onboarding lifecycle tests. Backend evidence only; nothing here is deployed.",
    )
    finish(doc, output, title)


def patch_m02(source: Path, output: Path) -> None:
    source_version, target_version = "1.3", "1.4"
    doc = Document(str(source))
    title = f"XVS M02 XVision Admin Console Functional Requirements Document v{target_version}"
    doc.core_properties.title = title
    doc.core_properties.version = target_version
    replace_cover_versions(doc.tables[0], (source_version,), target_version)
    replace_control_value(doc.tables[1], "Version", target_version)
    replace_control_value(doc.tables[1], "Review date", REVIEW_DATE)
    replace_control_value(doc.tables[1], "Code baseline", "Backend worktree at 30 August 2026")
    replace_control_value(
        doc.tables[1], "Source MRD",
        f"XVS Module Requirements Document v{MRD_TARGET_VERSION} | Module 2",
    )

    fr007 = doc.tables[13]
    replace_cell(
        fr007.rows[2].cells[1],
        "start locks both the asserted Tenant and the actor row with select_for_update. The Tenant is rechecked as ACTIVE or PENDING while locked, so it serializes with lifecycle shutdown; two simultaneous starts serialize on the actor. Validation runs before replacing an existing proxy, and the old sessions end before the new one is created inside the same transaction.",
        size=8.5,
    )
    replace_cell(
        fr007.rows[3].cells[1],
        "Each replaced session emits its own end event. If Tenant shutdown wins the shared lock, start observes the non-authenticable status and refuses creation; if start wins, the following shutdown ends the session it created.",
        size=8.5,
    )

    fr010 = doc.tables[16]
    replace_cell(
        fr010.rows[2].cells[1],
        "end_impersonations_for_user ends sessions where the changed user is either actor or target. Tenant status writers call transition_tenant_status instead of relying on post-save signals. The service locks the Tenant, writes status and lifecycle stamps, and ends every active proxy scoped to a non-ACTIVE Tenant in the same transaction. School mirroring, onboarding expiry, and orphaned-School reinstatement all use it. start takes the same Tenant lock and rechecks the status, closing the race between session creation and shutdown.",
        size=8.5,
    )
    replace_cell(
        fr010.rows[3].cells[1],
        "User and Tenant shutdown are bulk and idempotent. Tenant status and proxy shutdown commit or roll back together. An old session remains ENDED after reactivation and its header answers 401; a racing new start is either ended by the transition or refused after it.",
        size=8.5,
    )
    replace_cell(fr010.rows[4].cells[1], "None.", size=8.5)

    replace_cell(
        doc.tables[28].rows[6].cells[1],
        "Every production Tenant-status writer calls transition_tenant_status. It owns the locked state write, pending-period stamps, lifecycle timestamps, and transactional proxy shutdown. Proxy start locks the same Tenant row before rechecking status and creating a session.",
        size=8.2,
    )
    replace_cell(
        doc.tables[30].rows[0].cells[0],
        "FURTHER GAPS\n"
        "• Privileged access to this console does not require MFA. It is the platform's operator surface and the entry point to proxying, which makes it the strongest candidate for the requirement, and it is tracked as a P0 in the MRD.\n"
        "• The platform-actor gate is applied per view. A new globally scoped surface must remember it; nothing enforces it across the app.\n"
        "• Justification on a proxy session defaults to a fixed string, so the field is always populated but not always informative.\n"
        "• Task-monitor scope is platform-wide or the reader's own tenant, with nothing between. An operator cannot be confined to the schools they actually cover, because a platform-scoped permission cannot be granted inside a tenant role, so giving somebody one school's runs means giving them every school's.",
        size=8.5,
    )

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("A session is opened against a target pinned"):
            replace_paragraph(
                paragraph,
                "A session is opened against a target pinned to the asserted Tenant, inside a transaction that locks both the Tenant and the actor, ends whatever session the actor already held, and creates the new one. It ends when the actor exits, an end-key holder terminates it, the sweep expires it, an account becomes unusable, or the Tenant leaves ACTIVE. The Tenant transition and proxy shutdown share one transaction, and session start shares the Tenant lock, so suspension cannot leave a session active or miss one created concurrently. Returning the Tenant to ACTIVE never reopens an ended session.",
                size=9,
                space_after=5,
            )
        elif text.startswith("Module 2 carries 13 capability entries"):
            replace_paragraph(
                paragraph,
                f"Module 2 carries 13 capability entries in MRD v{MRD_TARGET_VERSION}. Each maps to the requirements below. Transactional Tenant-lifecycle shutdown strengthens the existing impersonation termination capability without changing the count.",
                size=9,
                space_after=5,
            )

    prepend_change_log(
        doc.tables[32],
        target_version,
        "Closes the proxy-session revival defect in Tenant lifecycle changes. The post-save receiver is removed as the security boundary. transition_tenant_status now locks the Tenant, writes its lifecycle state and ends active impersonation sessions in one transaction; School mirroring, onboarding expiry and orphaned-School reinstatement all call it. Proxy start takes the same Tenant lock and rechecks that the Tenant is authenticable, so a start racing shutdown is either created first and then ended or refused after the status change. A suspended and later reactivated Tenant cannot revive the old session, whose header answers 401. FR-007, FR-010, the proxy workflow, the Tenancy dependency and MRD traceability are reconciled. Module 2 remains Backend Complete and In use Complete with thirteen capability entries. Verified by 165 vs_admin_console tests, 308 schools.vs_schools tests, and 58 onboarding lifecycle tests. Backend evidence only; nothing here is deployed.",
    )
    finish(doc, output, title)


def patch_m09(source: Path, output: Path) -> None:
    source_version, target_version = "2.6", "2.7"
    doc = Document(str(source))
    title = f"XVS M09 School Onboarding Functional Requirements Document v{target_version}"
    doc.core_properties.title = title
    doc.core_properties.version = target_version
    replace_cover_versions(doc.tables[0], (source_version,), target_version)
    replace_control_value(doc.tables[1], "Version", target_version)
    replace_control_value(doc.tables[1], "Review date", REVIEW_DATE)
    replace_control_value(doc.tables[1], "Code baseline", "Backend worktree at 30 August 2026")
    replace_control_value(
        doc.tables[1], "Source MRD",
        f"XVS Module Requirements Document v{MRD_TARGET_VERSION} | Module 9",
    )
    replace_control_value(
        doc.tables[1], "Supporting apps",
        "vs_tenants, vs_rbac, vs_audit, vs_notifications, vs_tickets, vs_admin_console, vs_finance, vs_import_data, schools/vs_schools",
    )
    replace_control_value(doc.tables[1], "Supersedes", "v2.6 and all earlier versions, retained unchanged")

    replace_cell(
        doc.tables[4].rows[0].cells[0],
        "CURRENT CREATION, READINESS, AND SHUTDOWN BOUNDARY\n"
        "• Module 1 returns 503 and rolls the whole new School transaction back when its required administrator cannot be provisioned. Module 9 receives no half-created School to repair.\n"
        "• FIRST_ADMIN still checks later invitation activation and account or assignment liveness. ROLE_BASELINE still checks that the School administrator role grants authority before go-live.\n"
        "• Onboarding expiry and reinstatement use the shared locked Tenant transition service. Suspension closes active impersonation sessions in the same transaction, and reinstatement never revives them.\n"
        "• The onboarding catalog remains seven steps, the same seven for every School; no step varies with Branch count.",
        size=8.5,
    )

    fr013 = doc.tables[22]
    replace_cell(
        fr013.rows[2].cells[1],
        "A daily job expires first and warns second. The clock is Tenant.pending_since, stamped on entry into PENDING and left alone by ordinary edits. Expiry is written through School.status, and School.save() calls the shared Tenant transition service. That service locks the Tenant, writes SUSPENDED with its lifecycle stamps, and ends active impersonation sessions in the same transaction. A School-kind Tenant with no School profile uses the service directly rather than a raw queryset status update. The warning is recorded once per pending spell, and operators receive ageing and newly expired lists on the 1st and 15th.",
        size=8.5,
    )
    replace_cell(
        fr013.rows[3].cells[1],
        "A School pending for ninety days is suspended and can no longer authenticate. Its active impersonation sessions are ENDED, and later reinstatement cannot reopen them. A School warned yesterday is not warned again today, a School past ninety days is expired rather than warned, and an ordinary edit does not restart the clock.",
        size=8.5,
    )
    replace_cell(
        fr013.rows[4].cells[1],
        "A School-kind Tenant with no School profile is transitioned directly because there is no School row to write through, and the sweep counts it separately. It now receives the same pending-stamp and impersonation-shutdown guarantees as the normal School path.",
        size=8.5,
    )

    fr014 = doc.tables[23]
    replace_cell(
        fr014.rows[2].cells[1],
        "POST /v1/onboarding/reinstate/<school-slug>/ returns a SUSPENDED School to PENDING through the shared Tenant transition service and re-stamps pending_since, so the next sweep does not suspend it again the following morning; expiry_warned_at is cleared with it. An old impersonation session remains ENDED. The School is named in the path because a suspended Tenant is not authenticable. Two gates apply: onboarding.progress.reactivate and the requirement that the caller's own Tenant is the platform Tenant.",
        size=8.5,
    )
    replace_cell(
        fr014.rows[3].cells[1],
        "A School admin holding the key cannot reinstate anybody. A reinstated School has a full window, is warned again inside it, and cannot reuse an impersonation session ended by suspension. One ONBOARDING_REINSTATED audit event is written.",
        size=8.5,
    )

    replace_cell(
        doc.tables[27].rows[4].cells[1],
        "The School and Tenant become SUSPENDED and active impersonation sessions become ENDED in the same transaction.",
        size=8.2,
    )
    replace_cell(
        doc.tables[27].rows[4].cells[2],
        "The status mirror cannot undo the suspension, and later reinstatement cannot revive an old proxy session.",
        size=8.2,
    )
    replace_cell(
        doc.tables[30].rows[4].cells[2],
        "The School-facing status written by expiry. School.save() delegates the paired Tenant state to the shared transition service, which closes active impersonation sessions transactionally.",
        size=8.2,
    )
    replace_cell(
        doc.tables[34].rows[1].cells[1],
        "Creates the School, Tenant, at least one Branch, required administrator, and package. Supplies SUSPENDED and routes the paired Tenant lifecycle through the shared transition service, so expiry written through School.status carries the same transactional proxy shutdown.",
        size=8.2,
    )
    append_styled_row(
        doc.tables[34],
        [
            "Module 2, XVision Admin Console",
            "Owns transition_tenant_status and impersonation-session shutdown. The onboarding fallback for a School-kind Tenant with no School profile calls this service directly rather than writing status with QuerySet.update().",
        ],
    )

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Module 9 carries 20 capability entries"):
            replace_paragraph(
                paragraph,
                f"Module 9 carries 20 capability entries in MRD v{MRD_TARGET_VERSION}. Each maps to the requirements below. Transactional impersonation shutdown strengthens expiry and reinstatement without adding a capability or changing the count.",
                size=9,
                space_after=5,
            )
        elif text.startswith("Version 2.6 reconciles onboarding"):
            replace_paragraph(
                paragraph,
                "Version 2.7 makes onboarding suspension a permanent proxy-session boundary. Expiry and reinstatement use the same locked Tenant transition service as School status mirroring, including the fallback for a School-kind Tenant with no School profile. Suspension and active impersonation-session shutdown commit or roll back together, and reinstatement does not revive the ended session.",
                size=9,
                space_after=5,
            )
    prepend_change_log(
        doc.tables[41],
        target_version,
        "Makes onboarding suspension a permanent impersonation-session boundary. The expiry sweep still writes through School.status, but School.save() now delegates the paired Tenant lifecycle to one locked transition service that ends active proxy sessions in the same transaction. A School-kind Tenant with no School profile uses that service directly rather than QuerySet.update(). Manual reinstatement also uses it, re-stamps the pending clock, and never reopens a session ended by suspension. FR-013, FR-014, the lifecycle timeline, data contract, Module 1 and Module 2 dependencies, and MRD traceability are reconciled. Notification seeding and the other current decisions remain unchanged. Module 9 stays Backend Complete and In use Partial with twenty capability entries. Verified by 58 onboarding lifecycle tests, 308 schools.vs_schools tests, and 165 vs_admin_console tests. Backend evidence only; nothing here is deployed.",
    )
    finish(doc, output, title)


def main() -> None:
    parser = argparse.ArgumentParser()
    for name in ("mrd", "m01", "m02", "m09"):
        parser.add_argument(f"--{name}-source", type=Path, required=True)
        parser.add_argument(f"--{name}-output", type=Path, required=True)
    args = parser.parse_args()
    patch_mrd(args.mrd_source, args.mrd_output)
    patch_m01(args.m01_source, args.m01_output)
    patch_m02(args.m02_source, args.m02_output)
    patch_m09(args.m09_source, args.m09_output)
    for output in (
        args.mrd_output, args.m01_output, args.m02_output, args.m09_output,
    ):
        print(output)


if __name__ == "__main__":
    main()

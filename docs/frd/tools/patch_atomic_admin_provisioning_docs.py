#!/usr/bin/env python3
"""Version the MRD and affected FRDs for atomic administrator provisioning."""

from __future__ import annotations

import argparse
import copy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement

from generate_requirements_documents import (
    BLUE,
    assert_no_em_dash,
    rebuild_table,
    shrink_inherited_media,
    update_extended_title,
    write_cell,
    write_paragraph,
)


REVIEW_DATE = "30 August 2026"
SHORT_DATE = "30 Aug 2026"
MRD_SOURCE_VERSION = "2.47"
MRD_TARGET_VERSION = "2.48"


def replace_cell(cell, text: str, **kwargs) -> None:
    while len(cell.paragraphs) > 1:
        paragraph = cell.paragraphs[-1]
        paragraph._p.getparent().remove(paragraph._p)
    write_cell(cell, text, **kwargs)


def replace_paragraph(paragraph, text: str, **kwargs) -> None:
    write_paragraph(paragraph, text, **kwargs)


def replace_cover_version(table, source: str, target: str) -> None:
    for paragraph in table.rows[0].cells[0].paragraphs:
        for run in paragraph.runs:
            if source in run.text:
                run.text = run.text.replace(source, target)


def replace_control_value(table, label: str, value: str) -> None:
    for row in table.rows:
        if row.cells[0].text.strip() == label:
            replace_cell(row.cells[1], value, size=9)
            return
    raise ValueError(f"Control row not found: {label}")


def prepend_change_log(table, version: str, summary: str) -> None:
    template = table.rows[1]
    new_tr = copy.deepcopy(template._tr)
    template._tr.addprevious(new_tr)
    new_row = table.rows[1]
    replace_cell(new_row.cells[0], version, size=8)
    replace_cell(new_row.cells[1], SHORT_DATE, size=8)
    replace_cell(new_row.cells[2], summary, size=8)


def append_styled_row(table, values: list[str], *, size: float = 8.2) -> None:
    template = table.rows[-1]
    new_tr = copy.deepcopy(template._tr)
    template._tr.addnext(new_tr)
    row = table.rows[-1]
    for cell, value in zip(row.cells, values):
        replace_cell(cell, value, size=size)


def remove_row(table, index: int) -> None:
    row = table.rows[index]
    row._tr.getparent().remove(row._tr)


def keep_rows_together(table) -> None:
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cantSplit") is None:
            tr_pr.append(OxmlElement("w:cantSplit"))


def finish(doc, output: Path, title: str) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    update_extended_title(output, title)
    shrink_inherited_media(output)
    assert_no_em_dash(output)


def patch_mrd(source: Path, output: Path) -> None:
    doc = Document(str(source))
    title = f"XVS Module Requirements Document v{MRD_TARGET_VERSION}"
    doc.core_properties.title = title
    doc.core_properties.version = MRD_TARGET_VERSION
    replace_cover_version(doc.tables[0], MRD_SOURCE_VERSION, MRD_TARGET_VERSION)
    replace_control_value(doc.tables[1], "Version", MRD_TARGET_VERSION)
    replace_control_value(doc.tables[1], "Review date", REVIEW_DATE)
    replace_control_value(
        doc.tables[1],
        "Source scope",
        "Backend worktree making required school and branch administrator provisioning fail closed and atomic. The complete schools.vs_schools suite passed 308 tests in 1776.659 seconds on 30 August 2026. Backend evidence only; nothing here is deployed.",
    )

    replace_cell(
        doc.tables[2].rows[5].cells[0],
        f"5. v{MRD_TARGET_VERSION} Capability Delta",
        size=9,
        bold=True,
        color=BLUE,
    )
    replace_cell(
        doc.tables[2].rows[5].cells[1],
        "Required administrators now share the creation transaction",
        size=9,
    )

    replace_cell(
        doc.tables[8].rows[2].cells[1],
        "▸ Create the first school and branch administrators atomically with their parent record",
        size=8.5,
    )

    current_decision = doc.tables[14].rows[8].cells[0].text.rstrip()
    current_decision += (
        "\n• Fresh school installations carry the school_admin and branch_admin "
        "prebuilt templates from a school-app data migration. Existing template "
        "rows, including deliberately deactivated ones, are preserved. Missing or "
        "unusable required templates now refuse school or branch creation instead "
        "of leaving a parent record with no administrator."
    )
    replace_cell(doc.tables[14].rows[8].cells[0], current_decision, size=8.2)

    replace_cell(
        doc.tables[25].rows[10].cells[0],
        "NEEDS ATTENTION\n• Notification templates are a deploy precondition: where one is missing, dispatch logs a warning and skips the channel, so a message never arrives and nothing fails.\n• First-administrator readiness remains a defence after creation. It catches an invitation that was never activated, a later-disabled account, or an administrator role whose permission set has become empty.",
        size=8.2,
    )

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Shared-row multi-tenant school and branch administration"):
            replace_paragraph(
                paragraph,
                "Shared-row multi-tenant school and branch administration. The current backend supports creation, listing, editing, statistics, package setup, primary administrators, and branch lifecycle operations. A required administrator is provisioned inside the same transaction as the school or standalone branch, so a failed account, role grant, or invitation record leaves no ready-looking parent record. The main-branch designation is handed over rather than stranded, a school's sign-in address is governed and frozen once the school has been live, and school and branch writes emit audit evidence.",
                size=9,
                space_after=5,
            )
        elif text.startswith("The onboarding tracker and the go-live gate"):
            replace_paragraph(
                paragraph,
                "The onboarding tracker and the go-live gate, built as a module of its own. A school that is not yet live signs in, works a checklist, asks to be activated and is activated by platform staff; a school that abandons the process is warned, expires, and can be put back by hand. Module 1 now refuses the creation transaction when its required administrator cannot be provisioned, while onboarding still verifies that the account is active and the school administrator role grants authority before go-live. Approving, rejecting and reinstating are reserved to the platform tenant by a check on the caller, so a school that grants itself the key still cannot take itself live.",
                size=9,
                space_after=5,
            )
        elif text == f"5. v{MRD_SOURCE_VERSION} Capability Delta":
            replace_paragraph(
                paragraph,
                f"5. v{MRD_TARGET_VERSION} Capability Delta",
                size=17,
                bold=True,
                space_before=15,
                space_after=8,
            )
        elif text.startswith("This revision records that Student Management is built"):
            replace_paragraph(
                paragraph,
                "This revision closes the partial-success path in school and standalone branch creation. A 201 response now means the required administrator account, scoped role assignment, and invitation record were created; any provisioning failure rolls the whole parent transaction back and returns a typed 503 response.",
                size=9,
                space_after=5,
            )
        elif text == "Module 16: Gradebook & Assessments":
            paragraph.paragraph_format.page_break_before = True

    replace_cell(
        doc.tables[77].rows[2].cells[1],
        "Keep onboarding operational",
        size=8.2,
    )
    replace_cell(
        doc.tables[77].rows[2].cells[2],
        "Every required step now has a backend check. Keep notification templates seeded and invitation delivery observable so the first administrator can activate and finish the checklist; Module 9 remains Integration Partial until those operational dependencies are proven.",
        size=8.2,
    )
    replace_cell(
        doc.tables[77].rows[3].cells[1],
        "Build the academic foundation",
        size=8.2,
    )
    replace_cell(
        doc.tables[77].rows[3].cells[2],
        "Student Management is built, so Staff Management is now the one that gates the rest. Academic Structure is built and Timetable & Calendar now stands on it; Staff Management is the one that most limits what Timetable can check, because a teacher is currently a role grant rather than a person with hours, a specialism, and a qualification.",
        size=8.2,
    )

    rebuild_table(
        doc.tables[76],
        [f"v{MRD_TARGET_VERSION} capability delta", "Decision", "Evidence"],
        [
            [
                "Required administrator",
                "Atomic",
                "The account, tenant or branch role assignment, invitation record, admin link, and parent school or branch commit together or not at all.",
            ],
            [
                "Provisioning failure response",
                "Fail closed",
                "A typed ADMIN_PROVISIONING_FAILED error answers 503. The API does not return a ready-looking 201 or leave a school, tenant, branch, contact, user, or earlier nested administrator behind.",
            ],
            [
                "Administrator role templates",
                "Migration-backed",
                "Fresh school installations get school_admin and branch_admin templates idempotently. Existing and deliberately deactivated rows are preserved rather than silently repaired.",
            ],
            [
                "Onboarding readiness",
                "Defence in depth",
                "FIRST_ADMIN and ROLE_BASELINE still block go-live after later activation, account, assignment, or permission-set failures.",
            ],
        ],
        [1.75, 0.95, 4.57],
        font_size=8.2,
    )
    prepend_change_log(
        doc.tables[78],
        MRD_TARGET_VERSION,
        "School and standalone branch creation now fail closed when their required administrator cannot be provisioned. The provisioning service raises a typed error after rolling back its inner work, so the outer creation transaction removes the parent school, tenant, branches, contacts, users, role assignments, invitations, and any earlier nested administrators. The API returns 503 ADMIN_PROVISIONING_FAILED rather than a ready-looking 201. A school-app data migration guarantees the two required administrator role templates on fresh installations without overwriting existing or deliberately deactivated rows. Module 9 retains FIRST_ADMIN and ROLE_BASELINE as go-live defences for later activation or permission failures. The complete 308-test schools.vs_schools suite passed. Module states, ownership, capability counts, priority gaps, and build order do not change. Backend evidence only; nothing here is deployed.",
    )
    finish(doc, output, title)


def patch_m01(source: Path, output: Path) -> None:
    source_version, target_version = "1.15", "1.16"
    doc = Document(str(source))
    title = f"XVS M01 School and Branch Management Functional Requirements Document v{target_version}"
    doc.core_properties.title = title
    doc.core_properties.version = target_version
    replace_cover_version(doc.tables[0], source_version, target_version)
    replace_control_value(doc.tables[1], "Version", target_version)
    replace_control_value(doc.tables[1], "Review date", REVIEW_DATE)
    replace_control_value(doc.tables[1], "Code baseline", "Backend worktree at 30 August 2026")
    replace_control_value(doc.tables[1], "Source MRD", f"XVS Module Requirements Document v{MRD_TARGET_VERSION} | Module 1")

    replace_cell(
        doc.tables[6].rows[0].cells[0],
        "CURRENT MODULE DECISION\n• Module 1 remains Backend Partial and Integration Partial in MRD v2.48.\n• A required administrator is part of creation, not a repairable side effect. School creation and standalone Branch creation now commit the parent, account, scoped role assignment, invitation record, and administrator link together or commit none of them.\n• The defining blocker remains School lifecycle control. Operator suspension, reactivation, recoverable deletion, and the handling of closed or deleted Schools and Branches remain open.\n• The existing sign-in address, main-Branch, audit, Branch ceiling, reset, multi-posting administrator, and update-scope controls remain unchanged.",
        size=8.5,
    )

    fr001 = doc.tables[9]
    replace_cell(
        fr001.rows[2].cells[1],
        "SchoolCreateSerializer and School.save() create the School and Tenant atomically, together with at least one Branch. Required School and Branch administrators are now inside that same transaction. The set of books and Module 9 onboarding control room remain best effort inside their own savepoints because neither is the credential needed to operate the new School.",
        size=8.5,
    )
    replace_cell(
        fr001.rows[3].cells[1],
        "A 201 response persists one School, one protected Tenant, at least one Branch with exactly one main, and every required administrator account, scoped role assignment, invitation record, and admin link. A missing Branch is rejected before writes. Any required administrator provisioning failure returns 503 and leaves none of those creation records behind.",
        size=8.5,
    )
    replace_cell(
        fr001.rows[4].cells[1],
        "Books and onboarding control-room provisioning are still best effort. A books failure is repaired by command, and an onboarding-control failure is repaired by re-provisioning. Neither failure produces a false claim that a required administrator exists.",
        size=8.5,
    )

    fr004 = doc.tables[12]
    replace_cell(fr004.rows[0].cells[0], "FR-004 | Implemented", size=8.5, bold=True)
    replace_cell(
        fr004.rows[2].cells[1],
        "SchoolPrimaryAdmin, ContactInfo, User, TenantUserRoleAssignment, InvitationService, and the Celery invitation request form the path. provision_admin_user rolls its partial work back in an inner savepoint, logs the server-side cause, and raises AdminProvisioningError. That error escapes SchoolCreateSerializer's outer transaction, so a failure removes the School, Tenant, Branches, contacts, users, grants, invitation records, and any administrator provisioned earlier in the nested request. The school_admin template is guaranteed on fresh installations by an idempotent school-app data migration that preserves existing rows.",
        size=8.5,
    )
    replace_cell(
        fr004.rows[3].cells[1],
        "A valid unused email produces a School-scoped inactive user, whole-tenant school_admin assignment, invitation record, and SENT administrator link. A 201 response means those records exist. Any provisioning exception returns 503 ADMIN_PROVISIONING_FAILED and leaves no parent School or related creation records. Later email delivery or activation can still fail and is handled by invitation operations and Module 9 readiness.",
        size=8.5,
    )
    replace_cell(
        fr004.rows[4].cells[1],
        "No transactional creation gap remains. Notification event types and templates must still be seeded for delivery, and the administrator must still activate the invitation before Module 9 marks FIRST_ADMIN complete.",
        size=8.5,
    )

    fr006 = doc.tables[14]
    replace_cell(
        fr006.rows[2].cells[1],
        "BranchPrimaryAdmin and the shared admin-provisioning service create the current link. Reused emails remain idempotent for the account and invitation and receive a distinct Branch-scoped grant at every named Branch. The branch_admin template is guaranteed on fresh installations by the same idempotent school-app migration. A provisioning exception now escapes the inner savepoint and rolls back the standalone Branch or the entire nested School creation transaction.",
        size=8.5,
    )
    replace_cell(
        fr006.rows[3].cells[1],
        "A valid administrator is linked to the target Branch and tenant with a Branch-scoped role assignment and visible invitation state. An address named at more than one Branch produces one account and invitation and one grant at each Branch. A failed standalone provisioning attempt returns 503 and removes only the new Branch transaction; a nested failure removes the whole new School transaction.",
        size=8.5,
    )
    replace_cell(
        fr006.rows[4].cells[1],
        "Standalone Branch creation still declares primary admin input optional while the create path requires it. The previous partial-success and reconciliation gap is closed.",
        size=8.5,
    )

    creation_workflow = doc.tables[30]
    replace_cell(creation_workflow.rows[3].cells[0], "3. School records", size=8.2)
    replace_cell(
        creation_workflow.rows[3].cells[1],
        "Create branding and the required primary School administrator link.",
        size=8.2,
    )
    replace_cell(
        creation_workflow.rows[3].cells[2],
        "Keep branding best effort where specified; keep the required administrator inside the parent transaction.",
        size=8.2,
    )
    replace_cell(
        creation_workflow.rows[4].cells[2],
        "Keep every Branch and required administrator write inside the parent transaction; any administrator failure rolls all of it back.",
        size=8.2,
    )
    replace_cell(
        creation_workflow.rows[6].cells[1],
        "Emit School and Branch creation audit events; create invitation records and request asynchronous delivery.",
        size=8.2,
    )
    replace_cell(
        creation_workflow.rows[6].cells[2],
        "Return 503 and roll back creation if required administrator provisioning fails. Keep downstream delivery and activation observable as later lifecycle steps.",
        size=8.2,
    )

    append_styled_row(
        doc.tables[38],
        [
            "Required administrator provisioning",
            "Account, scoped role assignment, invitation record, admin link, and parent creation must commit together. The response does not expose the internal exception.",
            "503 ADMIN_PROVISIONING_FAILED with a generic retry-safe message; no parent creation records remain",
        ],
    )
    replace_cell(
        doc.tables[39].rows[3].cells[2],
        "Evaluate platform permissions; provision School and Branch administrator roles and assignments. Fresh installations receive the school_admin and branch_admin prebuilt templates from the school-app migration, while existing or deactivated rows are preserved.",
        size=8.2,
    )
    replace_cell(
        doc.tables[39].rows[4].cells[2],
        "Create pending users and invitation records; request activation email delivery. Creation guarantees the durable request records, not successful downstream delivery or activation.",
        size=8.2,
    )
    remove_row(doc.tables[40], 4)
    replace_cell(
        doc.tables[42].rows[6].cells[2],
        "Implemented. Required administrator creation is atomic with its parent; a reused address named at more than one posting receives every scoped grant. Delivery and activation remain separate lifecycle steps.",
        size=8.2,
    )
    keep_rows_together(doc.tables[42])
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("This table reconciles every Module 1 capability entry"):
            replace_paragraph(
                paragraph,
                f"This table reconciles every Module 1 capability entry in MRD v{MRD_TARGET_VERSION} to the controlling functional requirement in this FRD. All eighteen entries remain represented. The administrator capability changes its failure contract without adding a new product surface or changing the count.",
                size=9,
                space_after=5,
            )
        elif text == "11. Change Log":
            paragraph.paragraph_format.page_break_before = False
    prepend_change_log(
        doc.tables[43],
        target_version,
        "Makes required administrator provisioning atomic with School and standalone Branch creation. provision_admin_user now rolls its inner work back and raises a typed failure; the serializer's outer transaction then removes the parent School or Branch and every related record created in that request. The API returns 503 ADMIN_PROVISIONING_FAILED instead of 201 when no usable administrator was created. Migration 0010 guarantees the school_admin and branch_admin templates on fresh installations without overwriting existing or deliberately deactivated rows. Four atomicity regressions and the complete 308-test schools.vs_schools suite pass. FR-001 and FR-004 are reconciled, FR-006 retains only its declared-optional input mismatch, and the administrator reconciliation gap is removed. Backend evidence only; nothing here is deployed.",
    )
    finish(doc, output, title)


def patch_m09(source: Path, output: Path) -> None:
    source_version, target_version = "2.5", "2.6"
    doc = Document(str(source))
    title = f"XVS M09 School Onboarding Functional Requirements Document v{target_version}"
    doc.core_properties.title = title
    doc.core_properties.version = target_version
    replace_cover_version(doc.tables[0], source_version, target_version)
    replace_control_value(doc.tables[1], "Version", target_version)
    replace_control_value(doc.tables[1], "Review date", REVIEW_DATE)
    replace_control_value(doc.tables[1], "Code baseline", "Backend worktree at 30 August 2026")
    replace_control_value(doc.tables[1], "Source MRD", f"XVS Module Requirements Document v{MRD_TARGET_VERSION} | Module 9")
    replace_cell(
        doc.tables[4].rows[0].cells[0],
        "CURRENT CREATION AND READINESS BOUNDARY\n• Module 1 returns 503 and rolls the whole new School transaction back when its required administrator cannot be provisioned. Module 9 receives no half-created School to repair.\n• FIRST_ADMIN still checks later invitation activation and account or assignment liveness. ROLE_BASELINE still checks that the School administrator role grants authority before go-live.\n• The onboarding catalog remains seven steps, the same seven for every School; no step varies with Branch count.",
        size=8.5,
    )
    replace_control_value(doc.tables[1], "Supersedes", "v2.5 and all earlier versions, retained unchanged")

    replace_cell(
        doc.tables[6].rows[1].cells[1],
        "Module 1, School and Branch Management. Onboarding verifies the result and never re-creates it. Module 1 now makes the required administrator part of the School transaction and refuses creation when provisioning fails.",
        size=8.2,
    )
    replace_cell(
        doc.tables[6].rows[3].cells[1],
        "Module 13. It is mounted and supplies the academic-structure facts this module checks; onboarding does not create or edit them.",
        size=8.2,
    )
    replace_cell(
        doc.tables[12].rows[4].cells[1],
        "None. ACADEMIC_STRUCTURE is checked against the mounted Module 13 records before the task may complete.",
        size=8.5,
    )

    fr006 = doc.tables[15]
    replace_cell(
        doc.tables[14].rows[2].cells[1],
        "There is no BRANCH_SETUP task. The step was conditional on School.operates_branches, a stored flag saying whether the School ran more than one Branch; the flag is gone (Module 1, FR-018) and the honest predicate in its place is the Branch count, which every School satisfies from the moment Module 1 creates it. TaskKey, the catalog, and the conditions map no longer name the step, and vs_onboarding 0002 deletes the rows that carried it and renumbers the remaining seven to 1-7. Readiness is deliberately not recomputed by that migration: a School whose last blocking step was the Branch one becomes READY at its next evaluation, which is what emits the notification the people waiting on it expect.",
        size=8.5,
    )
    replace_cell(
        fr006.rows[1].cells[1],
        "A school with no working administrator, or an administrator role that grants nothing, must not reach go-live. Module 1 now fails creation atomically when the required account, scoped assignment, or invitation record cannot be created. This module remains the later defence against an invitation that was never activated, an account or assignment disabled after creation, or an empty permission set.",
        size=8.5,
    )
    replace_cell(
        fr006.rows[2].cells[1],
        "Module 1 raises ADMIN_PROVISIONING_FAILED and rolls the School transaction back when required administrator provisioning fails. FIRST_ADMIN requires an active tenant user holding an active whole-tenant school_admin assignment with no Branch. ROLE_BASELINE requires the tenant's school_admin role to carry at least one granted permission.",
        size=8.5,
    )
    replace_cell(
        fr006.rows[3].cells[1],
        "A failed creation attempt leaves no School or onboarding state. After a successful creation, an unactivated invitation, later-disabled account or assignment, Branch-pinned administrator, or empty school_admin role keeps FIRST_ADMIN or ROLE_BASELINE incomplete and blocks go-live.",
        size=8.5,
    )
    replace_cell(
        fr006.rows[4].cells[1],
        "This module does not retry Module 1 creation because no half-created School remains. Invitation delivery and activation failures still require the invitation operations to resend or repair them before FIRST_ADMIN can complete.",
        size=8.5,
    )
    replace_cell(
        doc.tables[26].rows[1].cells[2],
        "School, tenant, at least one Branch, and the required administrator account, scoped role assignment, and invitation record exist, all committed together. The School and tenant are PENDING. Books and the onboarding control room are best effort inside their own savepoints.",
        size=8.2,
    )
    replace_cell(
        doc.tables[31].rows[2].cells[2],
        "The tenant's whole-tenant school_admin role carries at least one granted permission.",
        size=8.2,
    )
    replace_cell(
        doc.tables[31].rows[5].cells[2],
        "The mounted Module 13 records satisfy the required academic-structure checks.",
        size=8.2,
    )
    append_styled_row(
        doc.tables[33],
        [
            "Required administrator provisioning fails during School creation",
            "503",
            "ADMIN_PROVISIONING_FAILED from Module 1; no School or onboarding state remains",
        ],
    )
    replace_cell(
        doc.tables[34].rows[1].cells[1],
        "Creates the School, tenant, at least one Branch, required administrator, and package, and owns Branch creation and metadata edits. Required administrator records share the parent transaction; a provisioning failure leaves no School for onboarding to repair. Supplies the SUSPENDED status this module writes.",
        size=8.2,
    )
    replace_cell(
        doc.tables[34].rows[8].cells[1],
        "Supplies the mounted academic-structure records and checks that make ACADEMIC_STRUCTURE machine-verifiable. This module reads those facts and never creates them.",
        size=8.2,
    )
    replace_cell(
        doc.tables[37].rows[0].cells[0],
        "ACADEMIC STRUCTURE DEPENDENCY CLOSED\n• Module 13 exists and is mounted. ACADEMIC_STRUCTURE is machine-checked against its records, so the old manual-completion gap no longer describes current state. Module 13 remains Partial for reasons outside this onboarding workflow.",
        size=8.5,
    )
    replace_cell(
        doc.tables[38].rows[0].cells[0],
        "CURRENT GAPS AND DECISIONS\n• If notification templates are not seeded, dispatch logs a warning and skips the channel silently, so onboarding notifications would never arrive and nothing would fail. Seeding is a deploy precondition.\n• The second-books guard exempts the CodeX platform tenant, so its enforcement point is the provisioning service rather than the endpoint.\n• A partial data import does not complete INITIAL_DATA, by decision. The matching import-screen correction is frontend work in another repository.\n• Automatic reinstatement does not exist and is not planned. Reinstatement is manual and platform-staff only, by decision.",
        size=8.5,
    )
    replace_cell(
        doc.tables[36].rows[6].cells[1],
        "Superseded on 17 August 2026. v2.2 answered School.operates_branches, with the Branch step withheld when it was false. There is now no such stored fact: a flag can disagree with the rows it describes, so the number of Branches a School runs is counted from its Branch rows. Nor is there a step to withhold, because every School is created with its main Branch and the step would have arrived already satisfied for all of them.",
        size=8.2,
    )
    for index, paragraph in enumerate(list(doc.paragraphs)):
        if (
            not paragraph.text.strip()
            and index + 1 < len(doc.paragraphs)
            and doc.paragraphs[index + 1].text.strip() == "4. Functional Requirements"
        ):
            paragraph._p.getparent().remove(paragraph._p)
            break
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Module 9 carries 20 capability entries"):
            replace_paragraph(
                paragraph,
                f"Module 9 carries 20 capability entries in MRD v{MRD_TARGET_VERSION}. Each maps to the requirements below. Atomic administrator provisioning changes the upstream guarantee consumed by FR-006 without adding an onboarding capability or changing the count.",
                size=9,
                space_after=5,
            )
        elif text.startswith("Version 2.3 is a minor revision of v2.2"):
            replace_paragraph(
                paragraph,
                "Version 2.6 reconciles onboarding with the atomic administrator guarantee in Module 1. A new School reaches onboarding only after its required administrator account, scoped assignment, invitation record, and administrator link commit with it. Invitation delivery, activation, later account or assignment state, and role permissions remain lifecycle facts checked before go-live.",
                size=9,
                space_after=5,
            )
    prepend_change_log(
        doc.tables[41],
        target_version,
        "Reconciles onboarding with atomic required-administrator provisioning in Module 1. A provisioning failure now returns 503 ADMIN_PROVISIONING_FAILED and rolls the new School transaction back, so Module 9 never receives a half-created School to repair. FIRST_ADMIN and ROLE_BASELINE remain defence-in-depth checks for later invitation activation, account or assignment liveness, and role permissions. The workflow, dependency, refusal, and FR-006 contracts are updated. Two stale current-state gaps are also removed: Module 13 is mounted and machine-checked, and the finance abstraction layer now exists. Notification template seeding remains a deploy precondition. Module 9 remains Backend Complete and Integration Partial with twenty capability entries. Backend evidence only; nothing here is deployed.",
    )
    finish(doc, output, title)


def patch_m04(source: Path, output: Path) -> None:
    source_version, target_version = "1.7", "1.8"
    doc = Document(str(source))
    title = f"XVS M04 Roles and Permissions RBAC Functional Requirements Document v{target_version}"
    doc.core_properties.title = title
    doc.core_properties.version = target_version
    replace_cover_version(doc.tables[0], source_version, target_version)
    replace_control_value(doc.tables[1], "Version", target_version)
    replace_control_value(doc.tables[1], "Review date", REVIEW_DATE)
    replace_control_value(
        doc.tables[1],
        "Source scope",
        "Backend worktree adding migration-backed availability for the two administrator role templates required by School and Branch creation (30 August 2026)",
    )
    replace_control_value(
        doc.tables[1],
        "MRD baseline",
        f"XVS Module Requirements Document v{MRD_TARGET_VERSION}, Module 4, seventeen capability entries",
    )
    replace_cell(
        doc.tables[2].rows[11].cells[1],
        f"Agreement with MRD v{MRD_TARGET_VERSION}'s seventeen capability entries",
        size=9,
    )

    fr021 = doc.tables[29]
    replace_cell(
        fr021.rows[2].cells[1],
        "PrebuiltRoleTemplate rows are read-only library definitions. provision_role_from_prebuilt copies one into a tenant-owned, locked system role and records access through set_role_access. The school-app migration now guarantees school_admin and branch_admin template rows on fresh installations because School and Branch creation require them. It uses get_or_create and does not reactivate or overwrite an existing row. The broader seed command remains responsible for converging descriptions, tier, scope, teacher, and permission defaults.",
        size=8.5,
    )
    replace_cell(
        fr021.rows[3].cells[1],
        "A fresh migrated school installation contains school_admin and branch_admin without an operator seed step. The seed command converges the complete three-template catalogue, including teacher and access defaults. Adoption remains idempotent, tenant-owned, locked, and durably audited.",
        size=8.5,
    )
    replace_cell(
        fr021.rows[4].cells[1],
        "The migration guarantees only the two required template identities. Permission defaults still depend on the seed command, and PrebuiltRoleTemplate.scope is written by that seeder but read by no provisioning code. See Needs Attention.",
        size=8.5,
    )
    append_styled_row(
        doc.tables[41],
        [
            "Module 1, School and Branch Management",
            "Consumes school_admin and branch_admin during required administrator creation. Its school-app migration guarantees those two template identities on fresh installations without modifying existing or deliberately deactivated rows; the RBAC seed command still owns their full access defaults.",
        ],
    )
    replace_cell(
        doc.tables[44].rows[4].cells[2],
        "Implemented. Three prebuilt templates are available through seeding and provisioned as tenant-owned copies; fresh school installations additionally guarantee the two administrator template identities by migration.",
        size=8.2,
    )
    replace_cell(
        doc.tables[45].rows[0].cells[0],
        f"MRD RECONCILIATION\n• MRD v{MRD_TARGET_VERSION} lists Module 4 as Complete with seventeen capability entries, all traced above.\n• The school-app migration makes the two administrator template identities available before School or Branch creation can consume them. It preserves existing and deliberately deactivated rows; seed_all_permissions remains the authority for the complete template catalogue and permission defaults.\n• This hardens the existing role-template capability without adding a product surface or changing the count.\n• What remains is recorded in Needs Attention, including assignment expiry, two unnarrowed read surfaces, advisory branch reach, the generic platform-decision declaration, prebuilt scope, aggregate seeder scope verification, and privileged MFA.",
        size=8.5,
    )
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("MRD v2.41 records Module 4"):
            replace_paragraph(
                paragraph,
                f"MRD v{MRD_TARGET_VERSION} records Module 4 as Roles & Permissions (RBAC), Phase V1, Backend Complete, Integration Complete, code vs_rbac, with seventeen capability entries. The module number, name, phase, states and ownership agree with this revision. The school-app bootstrap migration is an explicit consumer guarantee around the existing prebuilt-template capability, not new RBAC ownership.",
                size=9,
                space_after=5,
            )
    prepend_change_log(
        doc.tables[46],
        target_version,
        "Guarantees that fresh school installations have the school_admin and branch_admin prebuilt template identities required by School and Branch creation. The school-app data migration uses get_or_create, so it does not overwrite, reactivate, or delete an existing operator-owned template. seed_all_permissions remains responsible for the complete three-template catalogue and access defaults. FR-021, dependencies, and MRD traceability are updated; Module 4 remains Complete with seventeen capability entries. The complete 308-test schools.vs_schools suite exercised the migration and provisioning consumers. Backend evidence only; nothing here is deployed.",
    )
    finish(doc, output, title)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--mrd-source", type=Path, required=True)
    parser.add_argument("--mrd-output", type=Path, required=True)
    parser.add_argument("--m01-source", type=Path, required=True)
    parser.add_argument("--m01-output", type=Path, required=True)
    parser.add_argument("--m09-source", type=Path, required=True)
    parser.add_argument("--m09-output", type=Path, required=True)
    parser.add_argument("--m04-source", type=Path, required=True)
    parser.add_argument("--m04-output", type=Path, required=True)
    args = parser.parse_args()
    patch_mrd(args.mrd_source, args.mrd_output)
    patch_m01(args.m01_source, args.m01_output)
    patch_m09(args.m09_source, args.m09_output)
    patch_m04(args.m04_source, args.m04_output)
    for output in [args.mrd_output, args.m01_output, args.m09_output, args.m04_output]:
        print(output)


if __name__ == "__main__":
    main()

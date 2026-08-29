#!/usr/bin/env python3
"""Create the Module 30 Health FRD and reconcile the MRD to its first baseline."""

from __future__ import annotations

import argparse
import copy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from generate_requirements_documents import (
    BLUE,
    GREY,
    PALE_BLUE,
    add_body,
    add_callout as add_reference_callout,
    add_cover,
    add_heading,
    add_metadata_table,
    add_page_break,
    add_requirement,
    add_status_key,
    add_table,
    assert_no_em_dash,
    remove_body_content,
    set_headers,
    set_table_widths,
    update_extended_title,
    write_cell,
    write_paragraph,
)


REVIEW_DATE = "29 August 2026"
MRD_SOURCE_VERSION = "2.42"
MRD_TARGET_VERSION = "2.43"
FRD_VERSION = "1.0"
CODE_BASELINE = (
    "Backend worktree at 3dff0e489d72e937dce6d8bd50e72c5e8f237007 with "
    "operational Health alert delivery changes reviewed on 29 August 2026"
)

MRD_CHANGE_SUMMARY = (
    "Created the first Module 30 System Health & Monitoring FRD and reconciled its "
    "requirements, endpoints, data contracts, scheduled work, dependencies, and all "
    "13 MRD capabilities. The review identified one existing evidence defect: a "
    "per-service uptime monitor reports 100 percent when no daily rollup exists, while "
    "the global uptime contract correctly returns no value without evidence. Module 30 "
    "remains Backend Complete and Integration Complete with one current P1 gap. The "
    "module count and 484 capability entries are unchanged. Reconciled to Module 30 "
    "FRD v1.0. Backend evidence and the existing passing 37-test Module 30 and 107-test "
    "Module 8 runs only; nothing here is deployed."
)


REQUIREMENTS = [
    {
        "id": "FR-001",
        "title": "Keep One Explicit Registry of Monitored Services",
        "status": "Implemented",
        "requirement": (
            "The platform must identify each monitored service, its operational group, "
            "tier, kind, active state, current status, and status-change time without "
            "embedding product-domain concepts in the monitoring engine."
        ),
        "evidence": (
            "MonitoredService stores stable keys and display metadata. seed_health creates "
            "the API, auth, route-group, Celery, datastore, SMTP, payment, and DNS/SSL "
            "entries and retires obsolete entries without inventing telemetry."
        ),
        "acceptance": (
            "Re-seeding is idempotent, active services appear worst-first, and a service "
            "status change stamps status_changed_at only when the value changes."
        ),
        "limit": "None within the stated registry contract.",
    },
    {
        "id": "FR-002",
        "title": "Collect Bounded Request Metrics Without Risking the Request",
        "status": "Implemented",
        "requirement": (
            "Every resolved application request may contribute route, method, response "
            "family, throttling, latency, minute bucket, and optional tenant scope. The "
            "collector must never store raw paths or payloads and must never break the "
            "response it observes."
        ),
        "evidence": (
            "RequestMetricsMiddleware records the resolved route pattern after tenant "
            "context, excludes the Health API itself, and swallows instrumentation errors. "
            "The process-local collector merges fixed histograms into one locked "
            "RequestMetric row per bucket, route, method, and tenant. Public provider "
            "webhook routes are measured with a null tenant and no payload content."
        ),
        "acceptance": (
            "Unmatched paths and Health routes create no metric; concurrent flushes add "
            "counts without loss; request latency remains bounded to fixed histogram "
            "buckets; a metrics failure leaves the original HTTP response unchanged."
        ),
        "limit": (
            "The buffer is deliberately best-effort. A worker process can lose its current "
            "in-memory bucket if it terminates before the next flush."
        ),
    },
    {
        "id": "FR-003",
        "title": "Return Range-Bound Golden Signals and Endpoint Analytics",
        "status": "Implemented",
        "requirement": (
            "Platform operators must be able to inspect traffic, p50/p95/p99 latency, 5xx "
            "error rate, throttling, status families, saturation, comparisons, and route "
            "drill-downs over a bounded time range."
        ),
        "evidence": (
            "services.py owns the fixed live, 15-minute, hourly, daily, weekly, and "
            "30-day windows plus custom windows capped at 90 days. It merges persisted "
            "histograms and groups endpoint and chart rows without storing raw requests."
        ),
        "acceptance": (
            "A requested range produces one internally consistent set of cards, series, "
            "endpoint rows, top offenders, and previous-window deltas. Invalid custom "
            "bounds fall back to the supported range policy rather than scanning without "
            "a limit."
        ),
        "limit": "Latency percentiles are estimates from fixed buckets, not raw samples.",
    },
    {
        "id": "FR-004",
        "title": "Withhold Health Claims When the Sample Is Too Small",
        "status": "Implemented",
        "requirement": (
            "A request error rate or latency percentile must not declare a service healthy "
            "or unhealthy when too few requests support that statistical claim."
        ),
        "evidence": (
            "MIN_P95_SAMPLE is 30. window_status and alert evaluation share that floor, so "
            "low-volume cards and endpoint rows report unknown and error-rate or p95 rules "
            "skip evaluation. Raw numeric estimates remain available for inspection."
        ),
        "acceptance": (
            "At 29 requests the status is unknown and the rule does not fire; at 30 or "
            "more requests the normal threshold bands apply. A low-sample window can clear "
            "an existing alert because it no longer proves the breach."
        ),
        "limit": (
            "Thirty observations are a minimum floor, not a guarantee that a tail estimate "
            "is statistically stable."
        ),
    },
    {
        "id": "FR-005",
        "title": "Probe External and Internal Dependencies on Schedule",
        "status": "Implemented",
        "requirement": (
            "The module must execute configured HTTP, TCP, Redis, PostgreSQL, and SSL "
            "checks, persist each outcome, and classify slow, failed, or expiring targets "
            "without letting one probe abort the batch."
        ),
        "evidence": (
            "Probe executors normalize every result to status, response time, status code, "
            "error, and metadata. Celery beat runs active checks every five minutes; each "
            "result updates its service status and route-group services are refreshed from "
            "real request metrics."
        ),
        "acceptance": (
            "A probe exception becomes a stored critical result, SSL metadata carries the "
            "domain and days remaining, datastore saturation can influence status, and the "
            "task reports how many checks completed."
        ),
        "limit": "Probe truth depends on network reachability and the configured targets.",
    },
    {
        "id": "FR-006",
        "title": "Roll Up Uptime and Calculate SLO Error Budgets Honestly",
        "status": "Implemented with limits",
        "requirement": (
            "Daily rollups must preserve service availability history, and active SLOs must "
            "show target, attained uptime, window, remaining error budget, and breach state "
            "without claiming evidence that does not exist."
        ),
        "evidence": (
            "The hourly rollup task computes daily totals, failures, worst status, average "
            "response time, and uptime. Global uptime returns no value when no rollups exist. "
            "SLO status reads active objectives and their service rollups."
        ),
        "acceptance": (
            "A day with real checks calculates the observed percentage and worst status; an "
            "empty history is represented as unknown or absent across every uptime surface; "
            "error-budget use never falls below zero."
        ),
        "limit": (
            "The per-service uptime monitor currently substitutes 100 percent when its "
            "selected window contains no daily rollups. That value is not evidence and must "
            "be changed to unknown or absent."
        ),
    },
    {
        "id": "FR-007",
        "title": "Expose Queue Depth, Worker Capacity, and Tracked Jobs",
        "status": "Implemented",
        "requirement": (
            "Operators must see current queue depth and trend, task throughput, failed, "
            "retrying, and dead counts, worker capacity, retry storms, and the tracked job "
            "records behind operational work."
        ),
        "evidence": (
            "A minute schedule captures QueueSnapshot rows from the Redis broker and Celery "
            "inspect data. The queue service returns only queues with real snapshots. The "
            "task API reads core.BackgroundJob with status, queue, kind, and tenant filters."
        ),
        "acceptance": (
            "Unknown queues return no fabricated card, queue filters map to the tracked job "
            "kinds, for_tenant narrows the task collection, and every list remains ordered "
            "newest first."
        ),
        "limit": "Celery inspect can report no workers when the broker or control channel is unavailable.",
    },
    {
        "id": "FR-008",
        "title": "Manage Incidents and Their Operational Timeline",
        "status": "Implemented",
        "requirement": (
            "Authorized operators must create, list, retrieve, update, acknowledge, resolve, "
            "and annotate incidents with severity, ownership, affected services, summary, "
            "postmortem, timestamps, and append-only timeline events."
        ),
        "evidence": (
            "Incident APIs use method-aware RBAC, list and detail serializers, service "
            "relations, timeline records, and manual or automatic sources. New incident "
            "codes use UUID-backed human references instead of lexical read-and-increment "
            "allocation."
        ),
        "acceptance": (
            "Omitting a code creates a unique INC reference, status changes append timeline "
            "evidence, active filters include every non-resolved state, and incident detail "
            "returns its ordered event history."
        ),
        "limit": "The API records an acknowledgement timestamp but does not assign an on-call rota.",
    },
    {
        "id": "FR-009",
        "title": "Validate Alert Rules Against Signals the Target Can Produce",
        "status": "Implemented",
        "requirement": (
            "Operators may define threshold, comparator, duration, severity, target service "
            "or queue, enabled state, and delivery destination, but a request-derived rule "
            "must be refused when its selected service has no request route group."
        ),
        "evidence": (
            "AlertRule supports error rate, p95 latency, queue depth, SSL days remaining, and "
            "uptime metrics. The serializer validates error-rate and p95 targets against the "
            "explicit request-route mapping. The only delivery destination is email and "
            "in-app."
        ),
        "acceptance": (
            "API, auth, schools, billing, and reports may own request-derived rules; a "
            "PostgreSQL request-error rule is rejected by target_service_key; a null target "
            "means the whole request surface."
        ),
        "limit": "The route-group mapping is configuration in code and must advance when new monitored API groups are added.",
    },
    {
        "id": "FR-010",
        "title": "Prove an Uninterrupted Service-Scoped Breach Before Firing",
        "status": "Implemented",
        "requirement": (
            "duration_sec must mean that consecutive evaluations remained beyond the "
            "threshold for the full duration. It must not widen an aggregate window or let "
            "traffic from another service mask the selected service."
        ),
        "evidence": (
            "Each rule persists breach_started_at. A first breach starts the clock, a clear "
            "or unevaluable result resets it, and only elapsed time opens the alert. Error "
            "rate and p95 queries filter RequestMetric routes by the target service prefixes."
        ),
        "acceptance": (
            "A five-minute rule does not fire at 299 seconds, fires at 300 seconds, and does "
            "not fire after a clearing evaluation. Healthy finance traffic cannot dilute a "
            "schools-service breach."
        ),
        "limit": "Evaluation cadence is one minute, so firing precision is bounded by the scheduler interval.",
    },
    {
        "id": "FR-011",
        "title": "Open Exactly One Firing Alert and Automatic Incident",
        "status": "Implemented",
        "requirement": (
            "One rule may have at most one firing alert. Concurrent scheduler runs must not "
            "create duplicate alerts or automatic incidents for the same breach."
        ),
        "evidence": (
            "The evaluator locks the AlertRule row inside a database transaction and the "
            "Alert table has a conditional unique constraint for one firing row per rule. "
            "The migration resolves older duplicate rows and closes their unused automatic "
            "incidents."
        ),
        "acceptance": (
            "Two overlapping evaluators serialize on the same rule, one alert and one "
            "automatic incident remain firing, and repeated evaluation creates no second "
            "incident or notification dispatch."
        ),
        "limit": "Database locking requires a transactional database with row-lock support.",
    },
    {
        "id": "FR-012",
        "title": "Deliver Every New Health Alarm by Email and In-App Notification",
        "status": "Implemented",
        "requirement": (
            "When a sustained rule first fires, every active platform operator whose "
            "effective permissions include platform.health.manage must receive both an "
            "in-app notification and an email record."
        ),
        "evidence": (
            "The evaluator raises the transactional health.alert_fired event through Module "
            "8. The shared dispatcher writes one notification per recipient per channel, "
            "marks in-app rows SENT, queues email rows PENDING after commit, and attaches "
            "internal alert and incident correlation metadata."
        ),
        "acceptance": (
            "One eligible operator receives exactly two records; email uses the configured "
            "Celery retry lifecycle to SENT or FAILED; repeated evaluation sends nothing; "
            "missing recipients, templates, or routing are recorded on the incident timeline."
        ),
        "limit": (
            "SMTP success means provider acceptance, not confirmed inbox arrival. A routing "
            "or template failure before an email record exists is logged but has no separate "
            "automatic redispatch queue."
        ),
    },
    {
        "id": "FR-013",
        "title": "Resolve Recovered Alerts and Their Automatic Incidents",
        "status": "Implemented",
        "requirement": (
            "When the current signal no longer breaches, the firing alert must record its "
            "final value and resolution time. An automatic incident may resolve only when "
            "none of its linked alerts remains firing."
        ),
        "evidence": (
            "The evaluator transitions the open alert to resolved and calls the shared "
            "automatic-incident resolver. Manual incidents are never closed by that path."
        ),
        "acceptance": (
            "A recovered metric resolves its alert once, closes an automatic incident only "
            "after every linked alert clears, stamps resolved_at, and appends a resolution "
            "timeline event."
        ),
        "limit": "No automatic postmortem is generated.",
    },
    {
        "id": "FR-014",
        "title": "Correlate Deployments and Reliability Outcomes",
        "status": "Implemented",
        "requirement": (
            "Operators must annotate deployments and configuration changes on the Health "
            "timeline and inspect incident count, active count, mean time to acknowledge, "
            "and mean time to resolve over the reporting window."
        ),
        "evidence": (
            "Deployment records store version, environment, kind, actor, text, and timestamp. "
            "Overview includes in-range annotations. Reliability aggregates incident start, "
            "acknowledgement, and resolution timestamps over 30 days."
        ),
        "acceptance": (
            "Only deployments inside the selected overview range appear; MTTA and MTTR omit "
            "incidents missing the relevant timestamp; empty samples return no mean rather "
            "than zero."
        ),
        "limit": "Deployment annotations are operator or integration writes; this module does not infer releases from a provider.",
    },
    {
        "id": "FR-015",
        "title": "Expose Tenant Analytics Only to Platform Health Operators",
        "status": "Implemented",
        "requirement": (
            "Cross-tenant observability must remain a platform-only capability. Authorized "
            "operators may inspect tenant traffic and noisy-neighbour indicators or narrow "
            "analytics to one explicit tenant without confusing that filter with the tenant "
            "assertion used for authentication."
        ),
        "evidence": (
            "Every Health view requires an active authenticated user and platform.health.view "
            "or platform.health.manage. requested_tenant resolves for_tenant separately from "
            "the authentication tenant parameter, rejects unknown targets, and analytics "
            "queries filter RequestMetric."
        ),
        "acceptance": (
            "An anonymous caller receives 401; an unauthorized authenticated caller is "
            "refused; a valid slug or numeric tenant target returns only that tenant's rows; "
            "an unknown target returns 400 instead of silently returning global data."
        ),
        "limit": "Tenant metrics exist only for requests where tenant context was resolved before collection.",
    },
]


def add_real_bullets(doc, items: list[str], *, size: float = 9) -> None:
    for item in items:
        paragraph = doc.add_paragraph()
        set_bullet_numbering(paragraph)
        write_paragraph(paragraph, item, size=size, space_after=2)


def set_bullet_numbering(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    num_pr.append(ilvl)
    num_pr.append(num_id)


def add_callout(doc, title: str, lines: list[str], *, kind: str = "info"):
    table = add_reference_callout(doc, title, lines, kind=kind)
    cell = table.cell(0, 0)
    for paragraph in cell.paragraphs[1:]:
        text = paragraph.text.removeprefix("• ")
        set_bullet_numbering(paragraph)
        write_paragraph(paragraph, text, size=8.6, space_after=2)
    return table


def write_decision_cell(cell, title: str, lines: list[str], *, size: float = 8.2) -> None:
    while len(cell.paragraphs) > 1:
        paragraph = cell.paragraphs[-1]
        paragraph._p.getparent().remove(paragraph._p)
    write_paragraph(cell.paragraphs[0], title.upper(), size=size, bold=True, color=BLUE, space_after=3)
    for line in lines:
        paragraph = cell.add_paragraph()
        set_bullet_numbering(paragraph)
        write_paragraph(paragraph, line, size=size, space_after=2)


def replace_cell(cell, text: str, **kwargs) -> None:
    while len(cell.paragraphs) > 1:
        paragraph = cell.paragraphs[-1]
        paragraph._p.getparent().remove(paragraph._p)
    write_cell(cell, text, **kwargs)


def prepend_change_log(table, version: str, date: str, summary: str) -> None:
    template = table.rows[1]
    new_tr = copy.deepcopy(template._tr)
    template._tr.addprevious(new_tr)
    new_row = table.rows[1]
    replace_cell(new_row.cells[0], version, size=8)
    replace_cell(new_row.cells[1], date, size=8)
    replace_cell(new_row.cells[2], summary, size=8)


def assert_vocabulary(path: Path) -> None:
    doc = Document(str(path))
    text = "\n".join(
        [paragraph.text for paragraph in doc.paragraphs]
        + [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    ).lower()
    forbidden_terms = ("cam" + "pus", "cam" + "puses")
    forbidden = [term for term in forbidden_terms if term in text]
    if forbidden:
        raise ValueError(f"Forbidden branch vocabulary found in {path}: {forbidden}")


def build_health_frd(reference_path: Path, output_path: Path) -> None:
    doc = Document(str(reference_path))
    remove_body_content(doc)
    doc.core_properties.title = (
        f"XVS M30 System Health and Monitoring Functional Requirements Document v{FRD_VERSION}"
    )
    doc.core_properties.subject = "Code-aligned functional requirements for XVS Module 30"
    doc.core_properties.author = "CodeX Team"
    doc.core_properties.version = FRD_VERSION
    set_headers(doc, "CodeX | System Health & Monitoring | Functional Requirements Document (FRD)")

    add_cover(
        doc,
        family="Functional Requirements Document",
        title="Module 30: System Health & Monitoring",
        subtitle="XVision Systems | Code-aligned functional baseline",
        version=FRD_VERSION,
    )

    add_heading(doc, "Document Control", level=1)
    add_metadata_table(
        doc,
        [
            ("Document", "Functional Requirements Document (FRD)"),
            ("Module", "M30 | System Health & Monitoring"),
            ("Version", FRD_VERSION),
            ("Review date", REVIEW_DATE),
            ("Code baseline", CODE_BASELINE),
            ("Source MRD", f"XVS Module Requirements Document v{MRD_TARGET_VERSION} | Module 30"),
            ("Primary app", "vs_health"),
            (
                "Supporting apps",
                "core, vs_tenants, vs_rbac, vs_user, vs_notifications, vs_config",
            ),
            ("Status", "Code-aligned baseline for Product and Engineering review"),
            ("Owner", "CodeX Team"),
        ],
    )
    add_callout(
        doc,
        "Evidence boundary",
        [
            "Implemented means the stated backend path is present in the inspected code. It does not prove frontend completion, deployment, production adoption, data migration, worker availability, or external-provider delivery.",
            "Telemetry screens are evidence-based. Configuration seeding creates services, checks, rules, permissions, and SLOs but never fabricates measurements.",
            "The MRD and this FRD have independent versions. A change can require one or both documents to advance.",
        ],
        kind="info",
    )
    add_page_break(doc)

    add_heading(doc, "Table of Contents", level=1)
    add_table(
        doc,
        ["Section", "Purpose"],
        [
            ["Document Control", "Version, ownership, source baseline, and evidence rules"],
            ["1. Purpose and Scope", "Monitoring ownership, included signals, and boundaries"],
            ["2. Context and Status Model", "Telemetry flow, platform scope, and state vocabulary"],
            ["3. Actors, Permissions, and Ownership", "Who may read, manage, evaluate, and receive alarms"],
            ["4. Functional Requirements", "Testable behavior, current evidence, acceptance, and limits"],
            ["5. Workflows and Lifecycle Rules", "Collection, probing, alerting, delivery, and recovery"],
            ["6. Data Model and Relationships", "Persisted telemetry, configuration, incidents, and objectives"],
            ["7. API and Validation Contracts", "Routes, permissions, filters, ranges, and response rules"],
            ["8. Dependencies and Operational Evidence", "Workers, broker, database, RBAC, notification, and probes"],
            ["9. Needs Attention", "Current backend gaps and required completion"],
            ["10. MRD Traceability", "All Module 30 capability mappings"],
            ["11. Change Log", "Independent FRD revision history"],
        ],
        [2.35, 4.92],
        font_size=8.5,
    )
    add_body(
        doc,
        "Use the Word Navigation pane to jump between headings. This contents page is intentionally static for reliable headless rendering.",
        size=8.5,
        color=GREY,
    )
    add_page_break(doc)

    add_heading(doc, "1. Purpose and Scope", level=1)
    add_body(
        doc,
        "Module 30 is the platform's operational evidence and incident-response engine. It observes resolved HTTP routes, scheduled dependency probes, Celery queues and tracked jobs; calculates health and reliability views; evaluates alert rules; opens incidents; and tells authorized operators through email and the in-app feed. It is domain-neutral and uses tenant only as an optional analytics dimension.",
    )
    add_heading(doc, "1.1 In Scope", level=2)
    add_real_bullets(
        doc,
        [
            "Service registry, current status, synthetic checks, SSL evidence, daily uptime rollups, SLO attainment, and error budgets.",
            "Request traffic, response families, throttling, latency histograms, endpoint drill-downs, and tenant-attributed analytics.",
            "Queue snapshots, worker capacity, tracked background jobs, incidents, timelines, deployments, reliability, alert rules, active alerts, and recovery.",
            "Sustained and service-scoped breach evaluation, one-firing-alert concurrency, UUID-backed incident references, and transactional email plus in-app alarm delivery.",
            "Platform RBAC, bounded ranges, explicit tenant filters, live-data evidence boundaries, retention, and scheduled task ownership.",
        ],
    )
    add_heading(doc, "1.2 Out of Scope and Relocated Responsibilities", level=2)
    add_table(
        doc,
        ["Not owned here", "Owner"],
        [
            ["User authentication and permission vocabulary", "Module 3 and Module 4. Health consumes active-user and RBAC decisions."],
            ["Email transport and in-app inbox", "Module 8. Health raises one transactional event and stores only correlation metadata."],
            ["Background job business outcomes", "The domain module that owns each job. Health reads common tracking state and queue evidence."],
            ["Payment webhook authenticity or replay", "Module 18. Health records bounded request timing for the resolved public route, not provider payloads or payment state."],
            ["Frontend console implementation", "The consuming application. This FRD proves backend contracts only."],
            ["External uptime provider guarantees", "The configured network, SMTP, certificate, database, Redis, and HTTP targets."],
        ],
        [2.45, 4.82],
        font_size=8.4,
    )

    add_heading(doc, "2. Context and Status Model", level=1)
    add_body(
        doc,
        "The module has three evidence paths. Request middleware accumulates bounded route metrics and flushes them into minute rows. Celery beat executes probes, queue snapshots, alert evaluation, rollups, and retention. Platform operators read the aggregate API and manage configuration or incidents under method-aware RBAC. Telemetry rows are global operational data, with nullable tenant attribution used for slicing rather than ownership isolation.",
    )
    add_heading(doc, "2.1 Requirement Status", level=2)
    add_status_key(doc)
    add_heading(doc, "2.2 Health State Vocabulary", level=2)
    add_table(
        doc,
        ["State", "Meaning"],
        [
            ["healthy", "Observed evidence is inside the configured warning and critical bands."],
            ["warning", "Observed evidence is degraded or approaching a configured limit."],
            ["critical", "Observed evidence crosses a critical threshold or a probe fails."],
            ["unknown", "No honest conclusion is available, including an insufficient request sample."],
            ["operational", "The overview roll-up has no warning or critical active service."],
        ],
        [1.55, 5.72],
        font_size=8.5,
    )
    add_callout(
        doc,
        "Current module decision",
        [
            f"Module 30 remains Backend Complete and Integration Complete in MRD v{MRD_TARGET_VERSION}.",
            "One current evidence defect remains: per-service uptime windows substitute 100 percent when no daily rollup exists. FR-006 requires unknown or absent instead.",
            "Alert delivery is now operational through both email and in-app destinations after a sustained, service-scoped breach.",
        ],
        kind="attention",
    )
    add_page_break(doc)

    add_heading(doc, "3. Actors, Permissions, and Ownership", level=1)
    add_table(
        doc,
        ["Actor", "May do", "Gate or contract"],
        [
            ["Platform health viewer", "Read every Health analytics and detail surface.", "Active authentication plus platform.health.view."],
            ["Platform health manager", "Create or update incidents, events, alert rules, and deployments.", "Active authentication plus platform.health.manage."],
            ["Request middleware", "Measure resolved application requests and optional tenant attribution.", "Best-effort internal instrumentation; no API identity."],
            ["Celery beat and workers", "Run probes, queue capture, alert evaluation, rollups, pruning, and email delivery.", "Internal scheduler and configured broker/worker runtime."],
            ["Platform alarm recipient", "Receive in-app and email alarm records.", "Effective platform.health.manage grant on the platform tenant."],
            ["External dependency", "Answer an HTTP, TCP, Redis, PostgreSQL, or TLS probe.", "Configured target and timeout; never trusted as application identity."],
        ],
        [1.45, 3.2, 2.62],
        font_size=8.2,
    )
    add_heading(doc, "3.1 Permission Matrix", level=2)
    add_table(
        doc,
        ["Operation", "Permission", "Scope"],
        [
            ["All GET and list operations", "platform.health.view", "Platform-wide operational data, optionally filtered by tenant analytics dimension"],
            ["Create or update incident", "platform.health.manage", "Global incident record and selected monitored services"],
            ["Append incident event", "platform.health.manage", "Resolved incident UUID"],
            ["Create or update alert rule", "platform.health.manage", "Global rule registry"],
            ["Create deployment annotation", "platform.health.manage", "Global deployment timeline"],
            ["Receive a fired Health alarm", "Effective platform.health.manage", "Active platform-tenant recipients only"],
        ],
        [2.05, 2.05, 3.17],
        font_size=8.2,
    )
    add_heading(doc, "3.2 Ownership Boundaries", level=2)
    add_real_bullets(
        doc,
        [
            "Health configuration and telemetry are platform-owned. A tenant foreign key on RequestMetric is an analytics dimension, not a tenant-owned record boundary.",
            "Notification rows belong to the recipient who reads them. The platform tenant is also the alarm's origin because Health incidents are platform operations.",
            "Alert and incident correlation identifiers stay in internal notification metadata and are not exposed by the recipient feed serializer.",
            "The module imports no school application and uses no school-domain model as an engine dependency.",
        ],
    )
    add_page_break(doc)

    add_heading(doc, "4. Functional Requirements", level=1)
    add_body(
        doc,
        "Each requirement records the required behavior, inspected implementation evidence, acceptance boundary, and current limitation. Status is current state, not revision history.",
    )
    for index, requirement in enumerate(REQUIREMENTS):
        add_requirement(doc, requirement)
        if index in {3, 7, 11}:
            add_page_break(doc)

    add_page_break(doc)
    add_heading(doc, "5. Workflows and Lifecycle Rules", level=1)
    add_heading(doc, "5.1 Request Telemetry", level=2)
    add_table(
        doc,
        ["Step", "What happens", "Safety rule"],
        [
            ["1", "Middleware starts a monotonic timer before the application view.", "Measurement never changes request routing or authorization."],
            ["2", "After the response, it resolves the route pattern and optional tenant context.", "Unmatched paths and /v1/health/ routes are skipped; raw paths and payloads are never stored."],
            ["3", "The process-local collector adds response family, latency bucket, and throttling to one minute key.", "A thread lock prevents in-process lost updates; recording errors are swallowed."],
            ["4", "A background flush swaps the buffer and merges each key into the database under a row lock.", "Concurrent workers add to one canonical aggregate row."],
            ["5", "Services calculate cards, series, endpoint, and tenant views from selected rows.", "Ranges are bounded; small samples cannot claim a status."],
        ],
        [0.6, 3.5, 3.17],
        font_size=8.0,
    )
    add_heading(doc, "5.2 Probe, Rollup, and Retention Schedule", level=2)
    add_table(
        doc,
        ["Cadence", "Task", "Outcome"],
        [
            ["Every minute", "Capture queue snapshot", "Queue depth, task outcomes, worker capacity, and trend row"],
            ["Every minute", "Evaluate alert rules", "Breach state, alert and incident lifecycle, and first-fire delivery"],
            ["Every five minutes", "Run uptime checks", "Stored probe result and refreshed service or route-group status"],
            ["Hourly at :15", "Roll up uptime", "Daily uptime percentage, failures, worst status, and response average"],
            ["Daily at 03:00", "Prune metrics", "Retention of bounded raw rollups while preserving daily uptime history"],
        ],
        [1.35, 2.45, 3.47],
        font_size=8.1,
    )
    add_heading(doc, "5.3 Sustained Alert and Delivery", level=2)
    add_table(
        doc,
        ["State", "Evaluation", "Effect"],
        [
            ["Clear", "Current value does not breach or cannot be evaluated.", "Reset breach_started_at; resolve an open alert and eligible automatic incident."],
            ["First breach", "Current value crosses the comparator and threshold.", "Persist breach_started_at; do not fire before duration_sec elapses."],
            ["Sustained breach", "The rule remains breaching for the full duration.", "Under the locked rule row, open one alert and one automatic incident."],
            ["First-fire delivery", "Resolve active platform.health.manage holders.", "Write one SENT in-app row and one PENDING email row per recipient; queue email after commit."],
            ["Email retry", "SMTP delivery fails after the row exists.", "Use configured retry count and delay; end in SENT or FAILED with delivery history."],
            ["Repeated breach", "The firing alert still exists.", "Create no duplicate incident or notification."],
        ],
        [1.35, 2.6, 3.32],
        font_size=8.0,
    )

    add_page_break(doc)
    add_heading(doc, "6. Data Model and Relationships", level=1)
    add_table(
        doc,
        ["Model", "Purpose", "Key contract"],
        [
            ["MonitoredService", "Stable operational registry and current state", "Unique key; active ordering; optional status-change timestamp"],
            ["RequestMetric", "Minute route/method/tenant request aggregate", "Unique bucket tuple; bounded histogram; nullable tenant dimension"],
            ["UptimeCheck", "Recurring probe configuration", "Service, check type, target, interval, region, expected thresholds"],
            ["UptimeCheckResult", "One normalized probe outcome", "Status, time, code, bounded error text, and probe metadata"],
            ["UptimeDailyRollup", "Long-window availability evidence", "Unique service/day row with total and failed checks"],
            ["QueueSnapshot", "Point-in-time broker and worker evidence", "Queue/time index and retry-storm state"],
            ["Incident", "Manual or automatic operational incident", "UUID primary key and UUID-backed unique human code"],
            ["IncidentEvent", "Append-only incident timeline", "Kind, actor label, text, and timestamp"],
            ["AlertRule", "Threshold, duration, target, severity, and delivery config", "Persisted breach start and email_and_in_app destination"],
            ["Alert", "Firing or resolved instance of one rule breach", "At most one firing row per rule; optional incident and service"],
            ["Deployment", "Timeline annotation", "Version, environment, kind, actor, text, and occurrence time"],
            ["SLO", "Availability objective", "Service, target percentage, rolling days, and active state"],
        ],
        [1.55, 2.55, 3.17],
        font_size=7.9,
    )
    add_heading(doc, "6.1 Data Safety Decisions", level=2)
    add_callout(
        doc,
        "Operational evidence rules",
        [
            "Telemetry stores aggregates and bounded route patterns, not raw request bodies, credentials, or unbounded URLs.",
            "Incident and notification correlation metadata is internal and does not enter the recipient-facing serializer.",
            "Seed data is configuration only. Empty telemetry remains empty or unknown, except the per-service uptime defect recorded in FR-006.",
            "Retention removes old request metrics, check results, and queue snapshots through scheduled policy while daily rollups preserve longer history.",
        ],
        kind="info",
    )

    add_page_break(doc)
    add_heading(doc, "7. API and Validation Contracts", level=1)
    add_heading(doc, "7.1 Read Surfaces", level=2)
    add_table(
        doc,
        ["Endpoint", "Purpose"],
        [
            ["GET /v1/health/overview/", "Posture, uptime, golden signals, services, request series, deployments, queues, and active incidents"],
            ["GET /v1/health/services/", "Worst-first active service grid"],
            ["GET /v1/health/services/{key}/", "One service, uptime summary, and ten recent alerts"],
            ["GET /v1/health/uptime/monitors/", "All service uptime windows, segments, response series, and SSL data"],
            ["GET /v1/health/uptime/monitors/{key}/", "One uptime monitor or 404"],
            ["GET /v1/health/api-endpoints/", "Endpoint rows, top slow/error routes, and status-code series"],
            ["GET /v1/health/api-endpoints/detail/?route=", "One route's totals, percentiles, histogram, series, and affected tenants"],
            ["GET /v1/health/queues/", "Queue snapshots, trends, and workers"],
            ["GET /v1/health/tasks/", "Tracked jobs filtered by status, queue, kind, or for_tenant"],
            ["GET /v1/health/incidents/", "Incident list filtered by state or severity"],
            ["GET /v1/health/incidents/reliability/", "MTTA, MTTR, incident count, and active count"],
            ["GET /v1/health/incidents/{id}/", "Incident detail with timeline"],
            ["GET /v1/health/alerts/", "Firing alerts by default or resolved history"],
            ["GET /v1/health/alert-rules/", "Configured alert rules"],
            ["GET /v1/health/alert-rules/{id}/", "One alert rule"],
            ["GET /v1/health/tenants/", "Tenant health rows and noisy-neighbour indicators"],
            ["GET /v1/health/tenants/{tenant_id}/", "One tenant's golden signals, series, and endpoints"],
            ["GET /v1/health/deployments/", "Deployment annotations"],
            ["GET /v1/health/slos/", "SLO attainment and remaining error budget"],
        ],
        [3.0, 4.27],
        font_size=7.7,
    )
    add_heading(doc, "7.2 Write Surfaces", level=2)
    add_table(
        doc,
        ["Method and endpoint", "Contract"],
        [
            ["POST /v1/health/incidents/", "Open a manual incident; code optional; selected service keys validated"],
            ["PUT/PATCH /v1/health/incidents/{id}/", "Update incident state, ownership labels, services, narrative, and timestamps"],
            ["POST /v1/health/incidents/{id}/events/", "Append a validated incident timeline event"],
            ["POST /v1/health/alert-rules/", "Create a validated rule with the mandatory email and in-app destination"],
            ["PUT/PATCH /v1/health/alert-rules/{id}/", "Update thresholds, target, duration, severity, or enabled state"],
            ["POST /v1/health/deployments/", "Create a deployment, feature-flag, or configuration annotation"],
        ],
        [3.15, 4.12],
        font_size=8.0,
    )
    add_page_break(doc)
    add_heading(doc, "7.3 Query and Validation Rules", level=2)
    add_table(
        doc,
        ["Input", "Rule", "Failure"],
        [
            ["range", "Supported named window; unknown value falls back to 1h", "No unbounded scan"],
            ["start and end", "Both valid, start before end, maximum 90 days", "Fall back to named range policy"],
            ["for_tenant", "Resolve slug or numeric id separately from authentication tenant", "400 when no tenant matches"],
            ["route", "Required for endpoint detail and matched exactly after collection", "400 when absent"],
            ["service key", "Stable monitored-service key", "404 when absent"],
            ["alert request metric target", "Must exist in explicit request-route mapping", "400 on unsupported service"],
            ["channel", "email_and_in_app", "400 on any other value"],
            ["write permission", "platform.health.manage on unsafe methods", "403 when the effective grant is absent"],
        ],
        [1.45, 4.62, 1.2],
        font_size=7.9,
    )

    add_heading(doc, "8. Dependencies and Operational Evidence", level=1)
    add_table(
        doc,
        ["Dependency", "Contract"],
        [
            ["Django middleware and resolver", "Provide bounded route patterns, response codes, method, latency, and tenant context after authentication scoping."],
            ["PostgreSQL", "Persist telemetry and configuration, merge metric buckets, lock alert rules, and enforce one firing alert per rule."],
            ["Celery beat and workers", "Run five scheduled Health tasks and deliver queued email notifications."],
            ["Redis broker", "Supply queue depths, Celery inspect data, task transport, and Redis probe evidence."],
            ["Module 4 Roles & Permissions", "Publish platform.health.view/manage and resolve active alarm recipients from effective grants."],
            ["Module 8 Notifications & Delivery", "Seed health.alert_fired templates, create recipient-owned in-app and email rows, and retry SMTP delivery."],
            ["Module 3 Identity & Team", "Supply active platform user records and email addresses."],
            ["Module 6 Configuration", "Supply notification retry settings and runtime environment values."],
            ["External HTTP, TCP, SMTP, payment, and TLS targets", "Answer configured probes within timeout and expose only normalized result metadata."],
            ["core.BackgroundJob", "Provide common job status, actor-independent tenant scope, start/finish times, and failure state."],
        ],
        [2.05, 5.22],
        font_size=8.0,
    )
    add_heading(doc, "8.1 Operational Preconditions", level=2)
    add_real_bullets(
        doc,
        [
            "Run seed_health and notification template seeding before relying on probes, alert rules, permissions, SLOs, or alarm templates.",
            "Assign platform.health.manage to at least one active platform operator with a valid email address.",
            "Run Celery beat and workers, keep the broker reachable, and configure SMTP plus probe targets.",
            "Treat PENDING email as queued, SENT as provider acceptance, and FAILED as terminal exhaustion of the configured retry budget.",
            "Monitor the incident timeline for recipient, event, template, routing, or record-count failures that occur before SMTP delivery begins.",
        ],
    )
    add_heading(doc, "8.2 Verification Evidence", level=2)
    add_callout(
        doc,
        "Inspected evidence",
        [
            "Module 30 suite: 37 tests passed, including collection, analytics, small-sample guards, sustained service-scoped alerting, delivery records, incident references, seeding, authentication, and tenant filters.",
            "Module 8 suite: 107 tests passed, including dispatch, templates, recipient ownership, in-app state, email retry, terminal outcomes, settings, history, and privacy boundaries.",
            "The suites are backend evidence only. They do not prove live broker, worker, SMTP, external probe, or recipient inbox behavior in a deployment.",
        ],
        kind="success",
    )

    add_page_break(doc)
    add_heading(doc, "9. Needs Attention", level=1)
    add_table(
        doc,
        ["Priority", "Current gap", "Required completion", "FR"],
        [
            [
                "P1",
                "Unsupported 100 percent uptime claim",
                "Return unknown or no value from each per-service uptime window when no daily rollup exists, align the monitor list/detail and SLO empty-history contracts, and add regression coverage for a newly seeded service with no check result.",
                "FR-006",
            ],
        ],
        [0.7, 1.75, 4.0, 0.82],
        font_size=7.8,
    )
    add_callout(
        doc,
        "Removal rule",
        [
            "Remove this item only after the implementation returns an honest no-data state and the relevant monitor and SLO tests pass.",
            "Do not infer frontend completion, deployed probes, worker uptime, SMTP delivery, or inbox arrival from backend code evidence.",
        ],
        kind="attention",
    )

    add_page_break(doc)
    add_heading(doc, "10. MRD Traceability", level=1)
    add_body(
        doc,
        f"Module 30 carries 13 capability entries in MRD v{MRD_TARGET_VERSION}. Each maps to the controlling requirements below.",
    )
    add_table(
        doc,
        ["MRD Module 30 capability", "Requirements", "Current state"],
        [
            ["System-health overview", "FR-001, FR-003, FR-005, FR-007, FR-008, FR-014", "Implemented"],
            ["Service status and history", "FR-001, FR-004, FR-005, FR-006", "Implemented with FR-006 limit"],
            ["Uptime indicators", "FR-005, FR-006", "Implemented with FR-006 limit"],
            ["SSL status", "FR-005, FR-006", "Implemented"],
            ["Tenant-filterable API request analytics", "FR-002, FR-003, FR-004, FR-015", "Implemented"],
            ["Queue and task health", "FR-007", "Implemented"],
            ["Sustained, service-scoped alert rules, active alerts, and email and in-app operator delivery", "FR-009, FR-010, FR-011, FR-012, FR-013", "Implemented"],
            ["Incident records and events", "FR-008, FR-011, FR-013", "Implemented"],
            ["Reliability summaries", "FR-008, FR-013, FR-014", "Implemented"],
            ["Tenant health", "FR-002, FR-003, FR-004, FR-015", "Implemented"],
            ["Deployment visibility", "FR-014", "Implemented"],
            ["Service-level objective tracking", "FR-005, FR-006", "Implemented with FR-006 limit"],
            ["Unauthenticated provider-webhook operations", "FR-002", "Relocated operation; request health observed without payload capture"],
        ],
        [3.35, 2.15, 1.77],
        font_size=7.8,
    )

    add_page_break(doc)
    add_heading(doc, "11. Change Log", level=1)
    change_log = add_table(
        doc,
        ["Version", "Date", "Summary"],
        [
            [
                "1.0",
                "29 Aug 2026",
                f"First code-aligned baseline for Module 30. Records the service registry, request aggregation, tenant filtering, small-sample status policy, probes, uptime rollups, queues, tasks, incidents, deployments, SLOs, sustained service-scoped alerting, email and in-app delivery, recovery, operational dependencies, one current uptime evidence gap, and traceability to MRD v{MRD_TARGET_VERSION}. Backend evidence only; nothing here is deployed.",
            ],
        ],
        [0.75, 1.15, 5.37],
        font_size=8.0,
    )
    row = change_log.add_row()
    merged = row.cells[0].merge(row.cells[1]).merge(row.cells[2])
    write_cell(
        merged,
        "MAINTENANCE RULE | Version this FRD independently. Update it when Module 30 behavior, acceptance, permissions, data contracts, APIs, alert or incident lifecycle, notification effects, dependencies, status, or current gaps change. Preserve every prior version.",
        size=8,
        bold=True,
        color=BLUE,
        fill=PALE_BLUE,
    )
    set_table_widths(change_log, [0.75, 1.15, 5.37])

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    update_extended_title(
        output_path,
        f"XVS M30 System Health and Monitoring Functional Requirements Document v{FRD_VERSION}",
    )
    assert_no_em_dash(output_path)
    assert_vocabulary(output_path)


def patch_mrd(source_path: Path, output_path: Path) -> None:
    doc = Document(str(source_path))
    doc.core_properties.title = f"XVS Module Requirements Document v{MRD_TARGET_VERSION}"
    doc.core_properties.version = MRD_TARGET_VERSION

    for paragraph in doc.tables[0].rows[0].cells[0].paragraphs:
        for run in paragraph.runs:
            if MRD_SOURCE_VERSION in run.text:
                run.text = run.text.replace(MRD_SOURCE_VERSION, MRD_TARGET_VERSION)

    for row in doc.tables[1].rows:
        label = row.cells[0].text.strip()
        if label == "Version":
            replace_cell(row.cells[1], MRD_TARGET_VERSION, size=9)
        elif label == "Review date":
            replace_cell(row.cells[1], REVIEW_DATE, size=9)
        elif label == "Source scope":
            replace_cell(
                row.cells[1],
                "Backend worktree plus the first Module 30 FRD, one reconciled no-data uptime gap, and the existing 37 passing Module 30 plus 107 passing Module 8 tests (29 August 2026)",
                size=9,
            )

    for row in doc.tables[2].rows:
        if row.cells[0].text.strip().startswith("5."):
            replace_cell(
                row.cells[0],
                f"5. v{MRD_TARGET_VERSION} Documentation Delta",
                size=9,
                bold=True,
                color=BLUE,
            )
            replace_cell(
                row.cells[1],
                "Module 30 gains its first FRD and an honest no-data uptime gap",
                size=9,
            )

    module_thirty = doc.tables[72]
    write_decision_cell(
        module_thirty.rows[6].cells[0],
        "Current decision",
        [
            "Module 30 remains Backend Complete and Integration Complete with 13 capability entries.",
            "Request error-rate and p95 rules use only the selected service's route group. A null target remains platform-wide, and a service without request metrics is refused at configuration.",
            "duration_sec is a real sustained-breach clock. The first breaching evaluation stores its start, a clearing evaluation resets it, and an incident opens only after the full duration.",
            "A firing alert resolves active platform.health.manage holders and creates both email and in-app notification records through Module 8. Email uses its existing worker retries and terminal delivery history.",
            "Rule-row locking and one-firing-alert uniqueness prevent overlapping evaluators from duplicating incidents. Human incident references are UUID-backed.",
            "One material evidence gap remains: per-service uptime windows return 100 percent when no daily rollup exists. FR-006 in Module 30 FRD v1.0 requires unknown or absent instead.",
            "Verification completed with 37 Module 30 tests and 107 Module 8 tests.",
        ],
        size=8.2,
    )

    priority_table = doc.tables[75]
    row = priority_table.add_row()
    values = [
        "P1",
        "Unsupported uptime certainty",
        "Return unknown or no value for every per-service uptime window without daily rollup evidence and add empty-history monitor and SLO regression coverage.",
        "M30",
    ]
    for index, value in enumerate(values):
        replace_cell(row.cells[index], value, size=7.8)

    delta = doc.tables[76]
    while len(delta.rows) > 1:
        delta._tbl.remove(delta.rows[-1]._tr)
    headers = [f"v{MRD_TARGET_VERSION} documentation delta", "Decision", "Evidence"]
    for index, header in enumerate(headers):
        replace_cell(delta.rows[0].cells[index], header, size=8.2, bold=True, color="FFFFFF", fill=BLUE)
    rows = [
        ["Module 30 FRD", "Created", "First code-aligned Module 30 baseline with 15 functional requirements and all 13 capability mappings."],
        ["No-data uptime", "Gap recorded", "Per-service uptime windows currently substitute 100 percent without a daily rollup; FR-006 requires unknown or absent."],
        ["Alert delivery", "Reconciled", "Sustained service-scoped evaluation, one firing alert, email and in-app records, worker retry, and incident correlation agree across MRD and FRDs."],
    ]
    for row_index, values in enumerate(rows):
        new_row = delta.add_row()
        for index, value in enumerate(values):
            replace_cell(new_row.cells[index], value, size=8.2)
    set_table_widths(delta, [1.75, 0.95, 4.57])

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == f"5. v{MRD_SOURCE_VERSION} Documentation Delta":
            write_paragraph(
                paragraph,
                f"5. v{MRD_TARGET_VERSION} Documentation Delta",
                size=17,
                bold=True,
                space_before=15,
                space_after=8,
            )
        elif text.startswith("This revision turns a stored Health alert"):
            write_paragraph(
                paragraph,
                "This revision creates the first detailed Module 30 contract and records the one no-data uptime claim that remains inconsistent with the platform's evidence rule.",
                size=9,
                space_after=5,
            )

    prepend_change_log(
        doc.tables[78],
        MRD_TARGET_VERSION,
        "29 Aug 2026",
        MRD_CHANGE_SUMMARY,
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    update_extended_title(output_path, f"XVS Module Requirements Document v{MRD_TARGET_VERSION}")
    assert_no_em_dash(output_path)
    assert_vocabulary(output_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--reference-frd", type=Path, required=True)
    parser.add_argument("--frd-output", type=Path, required=True)
    parser.add_argument("--mrd-source", type=Path, required=True)
    parser.add_argument("--mrd-output", type=Path, required=True)
    args = parser.parse_args()

    build_health_frd(args.reference_frd, args.frd_output)
    patch_mrd(args.mrd_source, args.mrd_output)
    print(args.frd_output)
    print(args.mrd_output)


if __name__ == "__main__":
    main()

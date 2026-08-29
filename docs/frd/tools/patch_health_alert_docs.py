#!/usr/bin/env python3
"""Version the MRD and Notifications FRD for operational health alert delivery."""

from __future__ import annotations

import argparse
import copy
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table

from generate_requirements_documents import (
    BLUE,
    assert_no_em_dash,
    rebuild_table,
    shrink_inherited_media,
    update_extended_title,
    write_cell,
    write_paragraph,
)


REVIEW_DATE = "29 August 2026"
MRD_SOURCE_VERSION = "2.41"
MRD_TARGET_VERSION = "2.42"
FRD_SOURCE_VERSION = "1.3"
FRD_TARGET_VERSION = "1.4"

MRD_SOURCE_SCOPE = (
    "Backend worktree with sustained, service-scoped Health alert evaluation, "
    "email and in-app operator delivery, UUID-backed incident references, and "
    "37 passing Module 30 plus 107 passing Module 8 tests (29 August 2026)"
)
FRD_CODE_BASELINE = (
    "Backend worktree registering transactional health.alert_fired delivery to "
    "email and in-app notification records (29 August 2026)"
)

MRD_CHANGE_SUMMARY = (
    "Made Module 30 alerts operational rather than dashboard-only. A sustained "
    "rule breach now opens one incident and dispatches the transactional "
    "health.alert_fired event to active platform operators holding "
    "platform.health.manage, creating both an immediate in-app record and a "
    "queued email record with the existing retry and delivery-history lifecycle. "
    "Request error-rate and p95 rules now restrict their metrics to the selected "
    "service routes, unsupported service and metric combinations are refused, "
    "and duration_sec is backed by persisted breach-start state rather than an "
    "aggregate window. Rule-row locking and a conditional unique constraint "
    "prevent duplicate firing alerts, while incident references are UUID-backed. "
    "Modules 8 and 30 remain Backend Complete and Integration Complete with 22 "
    "and 13 capability entries. Reconciled to Notifications FRD v1.4. Backend "
    "evidence and passing 37-test Module 30 and 107-test Module 8 runs only; "
    "nothing here is deployed."
)

FRD_CHANGE_SUMMARY = (
    "Registers health.alert_fired as a transactional Module 30 event with email "
    "and in-app templates. A sustained Health breach resolves active platform "
    "operators through platform.health.manage and uses the shared dispatcher, so "
    "in-app delivery is immediately SENT while email enters the existing PENDING, "
    "retry, SENT or FAILED lifecycle. Notification metadata correlates each row "
    "to its alert and incident without exposing those internal identifiers through "
    "the feed serializer. Adds FR-015 and updates the event, transactional, "
    "delivery-history, dependency and traceability contracts. Module 8 remains "
    "Backend Complete and Integration Complete with 22 capability entries. "
    "Backend evidence and passing 107-test Module 8 plus 37-test Module 30 runs "
    "only; nothing here is deployed."
)


def replace_cell(cell, text: str, **kwargs) -> None:
    while len(cell.paragraphs) > 1:
        paragraph = cell.paragraphs[-1]
        paragraph._p.getparent().remove(paragraph._p)
    write_cell(cell, text, **kwargs)


def replace_paragraph(paragraph, text: str, **kwargs) -> None:
    write_paragraph(paragraph, text, **kwargs)


def replace_cover_version(table, source: str, target: str) -> None:
    for paragraph in table.rows[0].cells[0].paragraphs:
        for run in paragraph.runs:
            if source in run.text:
                run.text = run.text.replace(source, target)


def prepend_change_log(table, version: str, date: str, summary: str) -> None:
    template = table.rows[1]
    new_tr = copy.deepcopy(template._tr)
    template._tr.addprevious(new_tr)
    new_row = table.rows[1]
    replace_cell(new_row.cells[0], version, size=8)
    replace_cell(new_row.cells[1], date, size=8)
    replace_cell(new_row.cells[2], summary, size=8)


def normalize_branch_vocabulary(doc: Document) -> None:
    replacements = (
        ("Cam" + "puses", "Branches"),
        ("cam" + "puses", "branches"),
        ("Cam" + "pus", "Branch"),
        ("cam" + "pus", "branch"),
    )
    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    for paragraph in paragraphs:
        for run in paragraph.runs:
            text = run.text
            for source, target in replacements:
                text = text.replace(source, target)
            if text != run.text:
                run.text = text


def patch_mrd(source: Path, output: Path) -> None:
    doc = Document(str(source))
    doc.core_properties.title = f"XVS Module Requirements Document v{MRD_TARGET_VERSION}"
    doc.core_properties.version = MRD_TARGET_VERSION

    replace_cover_version(doc.tables[0], MRD_SOURCE_VERSION, MRD_TARGET_VERSION)
    for row in doc.tables[1].rows:
        label = row.cells[0].text.strip()
        if label == "Version":
            replace_cell(row.cells[1], MRD_TARGET_VERSION, size=9)
        elif label == "Review date":
            replace_cell(row.cells[1], REVIEW_DATE, size=9)
        elif label == "Source scope":
            replace_cell(row.cells[1], MRD_SOURCE_SCOPE, size=9)

    for row in doc.tables[2].rows:
        if row.cells[0].text.strip().startswith("5."):
            replace_cell(
                row.cells[0],
                f"5. v{MRD_TARGET_VERSION} Documentation Delta",
                size=9,
                bold=True,
                color=BLUE,
            )
            replace_cell(
                row.cells[1],
                "Health alerts now reach operators after a sustained service breach",
                size=9,
            )

    module_eight = doc.tables[23]
    replace_cell(
        module_eight.rows[6].cells[0],
        "▸  Background delivery, retries, and transactional operational alerts",
        size=8.2,
    )

    module_thirty = doc.tables[72]
    replace_cell(
        module_thirty.rows[3].cells[0],
        "▸  Sustained, service-scoped alert rules, active alerts, and email and in-app operator delivery",
        size=8.2,
    )
    replace_cell(
        module_thirty.rows[6].cells[0],
        "CURRENT DECISION\n"
        "• No material capability gap was identified within this module's stated backend scope.\n"
        "• Request error-rate and p95 rules use only the selected service's route group. A null target remains platform-wide, and a service without request metrics is refused at configuration.\n"
        "• duration_sec is a real sustained-breach clock. The first breaching evaluation stores its start, a clearing evaluation resets it, and an incident opens only after the full duration.\n"
        "• A firing alert resolves active platform.health.manage holders and creates both email and in-app notification records through Module 8. Email uses its existing worker retries and terminal delivery history.\n"
        "• Rule-row locking and one-firing-alert uniqueness prevent overlapping evaluators from duplicating incidents. Human incident references no longer depend on lexical allocation and are UUID-backed.\n"
        "• Verification completed with 37 Module 30 tests and 107 Module 8 tests.",
        size=8.2,
    )

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "Module 8: Notifications & Delivery":
            continue
        if text.startswith("Event-driven in-app and email delivery."):
            replace_paragraph(
                paragraph,
                "Event-driven in-app and email delivery. A record is owned by the "
                "tenant of the recipient who reads it, and the tenant an event is "
                "about is kept separately as its origin. Transactional operational "
                "alarms, including Module 30 health alerts, bypass delivery preferences "
                "and use both channels. One-time invitation and reset credentials are "
                "replaced only in the delivery worker, so notification records retain "
                "a marker rather than the raw secret. This is a notification engine, "
                "not a person-to-person chat product.",
                size=9,
                space_after=5,
            )
        elif text.startswith("Operational health, reliability, alert, incident"):
            replace_paragraph(
                paragraph,
                "Operational health, reliability, sustained service-scoped alerts, "
                "email and in-app operator delivery, incident, deployment, queue, and "
                "tenant-status endpoints.",
                size=9,
                space_after=5,
            )
        elif text == f"5. v{MRD_SOURCE_VERSION} Documentation Delta":
            replace_paragraph(
                paragraph,
                f"5. v{MRD_TARGET_VERSION} Documentation Delta",
                size=17,
                bold=True,
                space_before=15,
                space_after=8,
            )
        elif text.startswith("This revision keeps the mandatory tenant authentication"):
            replace_paragraph(
                paragraph,
                "This revision turns a stored Health alert into a delivered operational "
                "alarm and corrects the service, duration, concurrency, and incident "
                "identity contracts that decide when it fires.",
                size=9,
                space_after=5,
            )

    rebuild_table(
        doc.tables[76],
        [f"v{MRD_TARGET_VERSION} documentation delta", "Decision", "Evidence"],
        [
            [
                "Operator delivery",
                "Added",
                "A firing alert creates in-app and email records for active platform.health.manage holders through the transactional health.alert_fired event.",
            ],
            [
                "Service isolation",
                "Corrected",
                "Error-rate and p95 evaluation filters request metrics by the selected service route group and refuses unsupported targets.",
            ],
            [
                "Sustained duration",
                "Corrected",
                "Persisted breach-start state proves the rule remained breaching for duration_sec; a clear evaluation resets the clock.",
            ],
            [
                "Concurrent evaluation",
                "Hardened",
                "A locked rule row and one-firing-alert constraint prevent overlapping evaluators from opening duplicates.",
            ],
            [
                "Incident reference",
                "Corrected",
                "UUID-backed references replace lexical read-and-increment allocation.",
            ],
        ],
        [1.75, 0.95, 4.57],
        font_size=8.2,
    )

    prepend_change_log(
        doc.tables[78], MRD_TARGET_VERSION, "29 Aug 2026", MRD_CHANGE_SUMMARY,
    )

    normalize_branch_vocabulary(doc)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    update_extended_title(output, f"XVS Module Requirements Document v{MRD_TARGET_VERSION}")
    shrink_inherited_media(output)
    assert_no_em_dash(output)


def patch_frd(source: Path, output: Path) -> None:
    doc = Document(str(source))
    doc.core_properties.title = (
        "XVS M08 Notifications and Delivery Functional Requirements "
        f"Document v{FRD_TARGET_VERSION}"
    )
    doc.core_properties.version = FRD_TARGET_VERSION

    replace_cover_version(doc.tables[0], FRD_SOURCE_VERSION, FRD_TARGET_VERSION)
    for row in doc.tables[1].rows:
        label = row.cells[0].text.strip()
        if label == "Version":
            replace_cell(row.cells[1], FRD_TARGET_VERSION, size=9)
        elif label == "Review date":
            replace_cell(row.cells[1], REVIEW_DATE, size=9)
        elif label == "Code baseline":
            replace_cell(row.cells[1], FRD_CODE_BASELINE, size=9)
        elif label == "Source MRD":
            replace_cell(
                row.cells[1],
                f"XVS Module Requirements Document v{MRD_TARGET_VERSION} | Module 8",
                size=9,
            )
        elif label == "Supporting apps":
            replace_cell(
                row.cells[1],
                "vs_tenants, vs_rbac, vs_user, vs_config, vs_health, core",
                size=9,
            )

    replace_paragraph(
        doc.paragraphs[9],
        "Module 8 is the platform's one way of telling somebody something "
        "happened. A domain module raises a named event with a context; this "
        "module decides which channels may carry it for that tenant, renders the "
        "message for each recipient, writes the in-app record, queues the email, "
        "and keeps the outcome. Transactional operational alarms such as a Module "
        "30 health incident bypass preferences and use both email and in-app. "
        "Templates and ordinary rendered content remain visible to administrators. "
        "A one-time invitation or reset credential is the deliberate exception: "
        "history stores an inert marker and only the delivery worker sees the raw "
        "value. Queue tracking remains complete, while the bell is reserved for "
        "outcomes a person is waiting to use or act on.",
        size=9,
        space_after=5,
    )

    replace_cell(
        doc.tables[3].rows[5].cells[1],
        "In-app records written immediately, email queued through Celery with "
        "retries and terminal outcomes, including transactional operational alarms.",
        size=8.5,
    )
    replace_cell(
        doc.tables[6].rows[5].cells[1],
        "Raise registered events and dispatch them. Module 30 resolves active "
        "platform.health.manage holders for a firing Health alert. No HTTP surface; "
        "services call the dispatcher directly.",
        size=8.5,
    )

    fr001 = doc.tables[7]
    replace_cell(
        fr001.rows[2].cells[1],
        "Events are seeded from one registry and resolved by key at dispatch. An "
        "unknown or inactive key raises UnknownEventTypeError rather than sending "
        "anything. Each entry declares the channels it supports, its default, and "
        "whether it is transactional. health.alert_fired is registered with email "
        "and in-app templates.",
        size=8.5,
    )
    replace_cell(
        fr001.rows[3].cells[1],
        "An unregistered key is refused. An event marked inactive dispatches on no "
        "channel regardless of any tenant setting. The settings matrix and the "
        "catalogue endpoint both list exactly the active registry. Template seeding "
        "materializes both health.alert_fired channel templates.",
        size=8.5,
    )

    fr004 = doc.tables[10]
    replace_cell(
        fr004.rows[1].cells[1],
        "A password reset, invitation, or operational alarm that a tenant preference "
        "can suppress can lock somebody out or leave an outage unattended.",
        size=8.5,
    )
    replace_cell(
        fr004.rows[2].cells[1],
        "An event marked transactional bypasses the settings matrix entirely and "
        "dispatches on its supported channels. The platform kill switch still "
        "applies. health.alert_fired is transactional so both operator destinations "
        "remain mandatory.",
        size=8.5,
    )

    fr012 = doc.tables[18]
    replace_cell(
        fr012.rows[2].cells[1],
        "One record is written per recipient per channel with the rendered content "
        "at dispatch time. Email moves from PENDING to SENT or FAILED through Celery "
        "with retries, while in-app is immediately SENT. Health alert rows carry "
        "internal alert and incident correlation metadata. For one-time credentials, "
        "the record holds an inert marker and the delivery task receives a transient "
        "replacement that is not written back.",
        size=8.5,
    )

    fr014 = doc.tables[20]
    heading_xml = copy.deepcopy(doc.paragraphs[45]._p)
    fr014._tbl.addnext(heading_xml)
    heading = Paragraph(heading_xml, fr014._parent)
    replace_paragraph(
        heading,
        "FR-015 Deliver a Health Alarm on Both Operator Channels",
        size=13.5,
        bold=True,
        space_before=11,
        space_after=6,
    )
    table_xml = copy.deepcopy(fr014._tbl)
    heading_xml.addnext(table_xml)
    fr015 = Table(table_xml, fr014._parent)
    replace_cell(fr015.rows[0].cells[0], "FR-015 | Implemented", size=9, bold=True)
    replace_cell(
        fr015.rows[1].cells[1],
        "When a Health rule proves a sustained breach and opens an incident, the "
        "operators authorized to manage that incident must be told through both "
        "email and the in-app feed.",
        size=8.5,
    )
    replace_cell(
        fr015.rows[2].cells[1],
        "Module 30 resolves active users holding platform.health.manage in the "
        "platform tenant and raises transactional health.alert_fired. Dispatch writes "
        "one row per operator per channel and stores internal alert and incident "
        "correlation metadata. A routing, recipient, event, or template failure does "
        "not roll back the incident; Alertmanager records the delivery problem on its "
        "timeline.",
        size=8.5,
    )
    replace_cell(
        fr015.rows[3].cells[1],
        "One firing alert for one eligible operator creates exactly one SENT in-app "
        "record and one PENDING email record. The email task is queued only after "
        "commit and uses the configured retry lifecycle. Re-evaluation of the same "
        "open alert creates no duplicate dispatch.",
        size=8.5,
    )
    replace_cell(
        fr015.rows[4].cells[1],
        "Email still depends on the broker, worker, SMTP provider, and seeded event "
        "templates. In-app delivery is the immediate independent record. SENT email "
        "means provider acceptance, not confirmed inbox arrival.",
        size=8.5,
    )

    dependencies = doc.tables[28]
    template_row = dependencies.rows[1]
    new_tr = copy.deepcopy(template_row._tr)
    dependencies.rows[6]._tr.addprevious(new_tr)
    module_thirty_row = dependencies.rows[6]
    replace_cell(module_thirty_row.cells[0], "Module 30, System Health & Monitoring", size=8.2)
    replace_cell(
        module_thirty_row.cells[1],
        "Raises health.alert_fired only after a sustained, service-scoped breach, "
        "resolves platform.health.manage recipients, and supplies alert and incident "
        "correlation metadata.",
        size=8.2,
    )

    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("Module 8 carries 22 capability entries"):
            replace_paragraph(
                paragraph,
                f"Module 8 carries 22 capability entries in MRD v{MRD_TARGET_VERSION}. "
                "Each maps to the requirements below.",
                size=9,
                space_after=5,
            )

    traceability = doc.tables[32]
    mappings = {
        "In-app notification creation and inbox": "FR-003, FR-010, FR-015",
        "Secret-safe email notification delivery": "FR-005, FR-012, FR-015",
        "Event-type catalogue": "FR-001, FR-015",
        "Background delivery tasks and retries": "FR-012, FR-015",
        "Delivery history and outcomes": "FR-012, FR-014, FR-015",
    }
    for row in traceability.rows[1:]:
        capability = row.cells[0].text.strip()
        if capability in mappings:
            replace_cell(row.cells[1], mappings[capability], size=8.2)

    change_log = doc.tables[33]
    prepend_change_log(
        change_log, FRD_TARGET_VERSION, "29 Aug 2026", FRD_CHANGE_SUMMARY,
    )

    normalize_branch_vocabulary(doc)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    update_extended_title(
        output,
        "XVS M08 Notifications and Delivery Functional Requirements "
        f"Document v{FRD_TARGET_VERSION}",
    )
    shrink_inherited_media(output)
    assert_no_em_dash(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--mrd-source", type=Path, required=True)
    parser.add_argument("--mrd-output", type=Path, required=True)
    parser.add_argument("--frd-source", type=Path, required=True)
    parser.add_argument("--frd-output", type=Path, required=True)
    args = parser.parse_args()
    patch_mrd(args.mrd_source, args.mrd_output)
    patch_frd(args.frd_source, args.frd_output)
    print(args.mrd_output)
    print(args.frd_output)


if __name__ == "__main__":
    main()

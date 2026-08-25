#!/usr/bin/env python3
"""Generate the XVS MRD patch and the School and Branch FRD.

The source DOCX supplies the page furniture, media, typography, and color system.
The generated documents are separate semantic-versioned artifacts.
"""

from __future__ import annotations

import argparse
import copy
import re
import tempfile
import zipfile
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from shrink_cover_art import rewrite_document


BLUE = "355B9D"
PALE_BLUE = "EAF0FA"
VERY_PALE_BLUE = "F4F7FC"
PALE_GREEN = "E4F3EA"
PALE_AMBER = "FFF3D6"
PALE_RED = "FCE8E6"
PALE_GREY = "EEF1F5"
WHITE = "FFFFFF"
BLACK = "000000"
GREY = "5E6B7A"
GREEN = "1F7A4D"
AMBER = "A86800"
RED = "B42318"
PURPLE = "6F42C1"

MRD_VERSION = "2.8"
FRD_VERSION = "1.2"
REVIEW_DATE = "12 August 2026"
CODE_BASELINE = "355c59b665fd27d4c47152fa9638093b2e95dc49 plus uncommitted Branch lifecycle work (12 August 2026)"


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=BLUE, size="8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_cell_width(cell, width_inches: float) -> None:
    width = int(width_inches * 1440)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")
    cell.width = Inches(width_inches)


def set_table_widths(table, widths: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    total = int(sum(widths) * 1440)
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths):
                set_cell_width(cell, widths[idx])


def set_paragraph_border_bottom(paragraph, color=BLUE, size="14", space="5") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def format_run(run, *, size=None, bold=None, color=None, name="Aptos Narrow", italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def clear_paragraph(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def write_paragraph(paragraph, text: str, *, size=9, bold=False, color=BLACK,
                    alignment=None, space_before=0, space_after=3, name="Aptos Narrow"):
    clear_paragraph(paragraph)
    if alignment is not None:
        paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    format_run(run, size=size, bold=bold, color=color, name=name)
    return run


def write_cell(cell, text: str, *, size=8.5, bold=False, color=BLACK,
               fill=None, alignment=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    set_cell_border(cell, color="C7D5EC", size="6")
    if fill:
        set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    write_paragraph(
        paragraph,
        text,
        size=size,
        bold=bold,
        color=color,
        alignment=alignment,
        space_after=0,
    )


def format_header_row(row) -> None:
    for cell in row.cells:
        set_cell_shading(cell, BLUE)
        set_cell_border(cell, color=BLUE, size="8")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                format_run(run, size=8.5, bold=True, color=WHITE)
    set_repeat_table_header(row)


def add_heading(doc, text: str, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    write_paragraph(
        paragraph,
        text,
        size=17 if level == 1 else 13.5,
        bold=True,
        color=BLACK,
        space_before=15 if level == 1 else 11,
        space_after=8 if level == 1 else 6,
        name="Aptos Narrow",
    )
    set_paragraph_border_bottom(paragraph, size="14" if level == 1 else "8")
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_body(doc, text: str, *, bold=False, color=BLACK, size=9, after=5,
             alignment=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph = doc.add_paragraph()
    write_paragraph(
        paragraph,
        text,
        size=size,
        bold=bold,
        color=color,
        alignment=alignment,
        space_after=after,
    )
    return paragraph


def add_bullets(doc, items: list[str], *, size=9, color=BLACK) -> None:
    for item in items:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.2)
        paragraph.paragraph_format.first_line_indent = Inches(-0.14)
        write_paragraph(paragraph, f"• {item}", size=size, color=color, space_after=2)


def add_page_break(doc) -> None:
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[float],
              *, row_fills=True, font_size=8.2):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_widths(table, widths)
    for idx, header in enumerate(headers):
        write_cell(table.rows[0].cells[idx], header, size=font_size, bold=True, color=WHITE, fill=BLUE)
    format_header_row(table.rows[0])
    for row_index, values in enumerate(rows):
        row = table.add_row()
        fill = VERY_PALE_BLUE if row_fills and row_index % 2 else WHITE
        for idx, value in enumerate(values):
            write_cell(row.cells[idx], value, size=font_size, fill=fill)
    table.rows[-1]._tr.addnext(OxmlElement("w:bookmarkEnd")) if False else None
    return table


def add_metadata_table(doc, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    set_table_widths(table, [1.75, 5.52])
    for label, value in rows:
        row = table.add_row()
        write_cell(row.cells[0], label, size=8.8, bold=True, color=WHITE, fill=BLUE)
        write_cell(row.cells[1], value, size=8.8, fill=WHITE)


def add_callout(doc, title: str, lines: list[str], *, kind="info"):
    fill = {
        "info": PALE_BLUE,
        "attention": PALE_AMBER,
        "critical": PALE_RED,
        "success": PALE_GREEN,
    }[kind]
    title_color = {
        "info": BLUE,
        "attention": AMBER,
        "critical": RED,
        "success": GREEN,
    }[kind]
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [7.27])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color=title_color, size="9")
    set_cell_margins(cell, top=100, start=120, bottom=100, end=120)
    write_paragraph(cell.paragraphs[0], title.upper(), size=9, bold=True, color=title_color, space_after=3)
    for line in lines:
        paragraph = cell.add_paragraph()
        write_paragraph(paragraph, f"• {line}", size=8.6, color=BLACK, space_after=2)
    return table


def add_status_key(doc) -> None:
    rows = [
        ["Implemented", "The code path exists and supports the stated core behavior."],
        ["Implemented with limits", "The main path exists, but a material edge case or operational control remains."],
        ["Partial", "Useful behavior exists, but an essential part of the requirement is missing."],
        ["Defective", "An endpoint or service exists, but the current implementation cannot safely complete the operation."],
        ["Not implemented", "No supported code path exists for the stated requirement."],
        ["Relocated", "The behavior belongs to another module and is traced there instead of being duplicated."],
    ]
    table = add_table(doc, ["State", "Meaning"], rows, [1.75, 5.52], font_size=8.5)
    fills = [PALE_GREEN, PALE_GREEN, PALE_AMBER, PALE_RED, PALE_GREY, PALE_BLUE]
    colors = [GREEN, GREEN, AMBER, RED, GREY, BLUE]
    for idx, (fill, color) in enumerate(zip(fills, colors), start=1):
        set_cell_shading(table.rows[idx].cells[0], fill)
        for run in table.rows[idx].cells[0].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(color)


def add_requirement(doc, requirement: dict) -> None:
    status = requirement["status"]
    status_fill = {
        "Implemented": PALE_GREEN,
        "Implemented with limits": PALE_GREEN,
        "Partial": PALE_AMBER,
        "Defective": PALE_RED,
        "Not implemented": PALE_GREY,
        "Relocated": PALE_BLUE,
    }[status]
    status_color = {
        "Implemented": GREEN,
        "Implemented with limits": GREEN,
        "Partial": AMBER,
        "Defective": RED,
        "Not implemented": GREY,
        "Relocated": BLUE,
    }[status]
    add_heading(doc, f"{requirement['id']}  {requirement['title']}", level=2)
    table = doc.add_table(rows=1, cols=2)
    set_table_widths(table, [1.48, 5.79])
    merged = table.rows[0].cells[0].merge(table.rows[0].cells[1])
    write_cell(
        merged,
        f"{requirement['id']} | {status}",
        size=9,
        bold=True,
        color=status_color,
        fill=status_fill,
    )
    for label, value in (
        ("Requirement", requirement["requirement"]),
        ("Current evidence", requirement["evidence"]),
        ("Acceptance", requirement["acceptance"]),
        ("Current limit", requirement["limit"]),
    ):
        row = table.add_row()
        write_cell(row.cells[0], label, size=8.3, bold=True, color=BLUE, fill=VERY_PALE_BLUE)
        write_cell(row.cells[1], value, size=8.3, fill=WHITE)


def add_cover(doc, *, family: str, title: str, subtitle: str, version: str) -> None:
    with tempfile.TemporaryDirectory() as tmp:
        reference = Path(doc._part.package.part_related_by.__self__.partname) if False else None
    logo = doc.add_paragraph()
    logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo.paragraph_format.space_before = Pt(72)
    logo.paragraph_format.space_after = Pt(64)
    logo_run = logo.add_run("X")
    format_run(logo_run, size=36, bold=True, color=BLUE, name="Aptos Display")

    panel = doc.add_table(rows=1, cols=1)
    set_table_widths(panel, [5.8])
    panel.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = panel.cell(0, 0)
    set_cell_shading(cell, BLUE)
    set_cell_border(cell, color=BLACK, size="10")
    set_cell_margins(cell, top=190, start=180, bottom=190, end=180)
    write_paragraph(cell.paragraphs[0], family.upper(), size=20, bold=True, color=WHITE, space_after=7, name="Aptos Display")
    p = cell.add_paragraph()
    write_paragraph(p, title, size=14, bold=True, color=WHITE, space_after=5, name="Aptos Display")
    p = cell.add_paragraph()
    write_paragraph(p, subtitle, size=10, color=WHITE, space_after=9)
    p = cell.add_paragraph()
    write_paragraph(p, f"Prepared by: CodeX Team     Version: {version}", size=9.5, bold=True, color=WHITE, space_after=0)

    year = doc.add_paragraph()
    year.paragraph_format.space_before = Pt(110)
    write_paragraph(year, "2026", size=26, bold=True, color=BLUE, space_after=8, name="Aptos Display")
    add_body(doc, "Internal Documentation", size=9, color=GREY)
    add_page_break(doc)


def set_headers(doc, title: str) -> None:
    section = doc.sections[0]
    for header in (section.header, section.even_page_header):
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        text_runs = [run for run in paragraph.runs if run.text and run.text != "\t"]
        if text_runs:
            text_runs[0].text = title
            format_run(text_runs[0], size=8.5, color=BLUE, name="Aptos Narrow")
            for run in text_runs[1:]:
                run.text = ""
        else:
            run = paragraph.add_run(title)
            format_run(run, size=8.5, color=BLUE, name="Aptos Narrow")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        set_paragraph_border_bottom(paragraph, size="12", space="5")


def remove_body_content(doc) -> None:
    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested in cell.tables:
                yield from iter_table_paragraphs(nested)


def iter_all_paragraphs(doc):
    yield from doc.paragraphs
    for table in doc.tables:
        yield from iter_table_paragraphs(table)
    for section in doc.sections:
        for part in (
            section.header,
            section.even_page_header,
            section.first_page_header,
            section.footer,
            section.even_page_footer,
            section.first_page_footer,
        ):
            yield from part.paragraphs
            for table in part.tables:
                yield from iter_table_paragraphs(table)


def replace_text(doc, old: str, new: str) -> None:
    for paragraph in iter_all_paragraphs(doc):
        if old not in paragraph.text:
            continue
        for run in paragraph.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)
        if old in paragraph.text:
            text = paragraph.text.replace(old, new)
            write_paragraph(paragraph, text, size=9, color=BLACK, space_after=paragraph.paragraph_format.space_after.pt if paragraph.paragraph_format.space_after else 0)


def find_table_by_header(doc, header_text: str):
    for table in doc.tables:
        if table.rows and header_text in " | ".join(cell.text for cell in table.rows[0].cells):
            return table
    raise ValueError(f"Table not found: {header_text}")


def rebuild_table(table, headers: list[str], rows: list[list[str]], widths: list[float], *, font_size=8.2) -> None:
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)
    while len(table.columns) < len(headers):
        table.add_column(Inches(1))
    for idx, header in enumerate(headers):
        write_cell(table.rows[0].cells[idx], header, size=font_size, bold=True, color=WHITE, fill=BLUE)
    format_header_row(table.rows[0])
    for row_index, values in enumerate(rows):
        row = table.add_row()
        fill = VERY_PALE_BLUE if row_index % 2 else WHITE
        for idx, value in enumerate(values):
            write_cell(row.cells[idx], value, size=font_size, fill=fill)
    set_table_widths(table, widths)


def replace_table(doc, table, headers: list[str], rows: list[list[str]], widths: list[float], *, font_size=8.2):
    replacement = add_table(doc, headers, rows, widths, font_size=font_size)
    table._tbl.addprevious(replacement._tbl)
    table._tbl.getparent().remove(table._tbl)
    return replacement


def update_extended_title(path: Path, title: str) -> None:
    with tempfile.NamedTemporaryFile(dir=path.parent, suffix=".docx", delete=False) as handle:
        temporary_path = Path(handle.name)
    try:
        with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(
            temporary_path, "w", zipfile.ZIP_DEFLATED
        ) as target:
            for item in source.infolist():
                data = source.read(item.filename)
                if item.filename == "docProps/app.xml":
                    text = data.decode("utf-8")
                    text = re.sub(
                        r"(<TitlesOfParts>.*?<vt:lpstr>).*?(</vt:lpstr>.*?</TitlesOfParts>)",
                        lambda match: match.group(1) + escape(title) + match.group(2),
                        text,
                        count=1,
                        flags=re.DOTALL,
                    )
                    data = text.encode("utf-8")
                target.writestr(item, data)
        temporary_path.replace(path)
    finally:
        if temporary_path.exists():
            temporary_path.unlink()


def build_mrd(reference_path: Path, output_path: Path) -> None:
    doc = Document(reference_path)
    source_version = doc.core_properties.version or "2.5"
    doc.core_properties.title = f"XVS Module Requirements Document v{MRD_VERSION}"
    doc.core_properties.subject = "Code-backed cross-module requirements and progress baseline"
    doc.core_properties.author = "CodeX Team"
    doc.core_properties.version = MRD_VERSION

    set_headers(doc, "CodeX | Module Requirements Document (MRD)")
    replace_text(doc, "MODULE & FEATURES", "MODULE REQUIREMENTS")
    replace_text(doc, "REQUIREMENTS BREAKDOWN", "DOCUMENT")
    replace_text(doc, "Module & Functional Requirements Breakdown", "Module Requirements Document (MRD)")
    replace_text(doc, "CodeX – Module & Features Requirements Breakdown", "CodeX | Module Requirements Document (MRD)")
    replace_text(doc, "CodeX – Module & Functional Requirements Breakdown", "CodeX | Module Requirements Document (MRD)")
    replace_text(doc, "current feature breakdown", "current module requirements")

    cover = doc.tables[0]
    for paragraph in iter_table_paragraphs(cover):
        for run in paragraph.runs:
            if run.text.strip() == source_version:
                run.text = MRD_VERSION

    control = doc.tables[1]
    control.cell(0, 1).text = "Module Requirements Document (MRD)"
    control.cell(2, 1).text = MRD_VERSION
    control.cell(3, 1).text = REVIEW_DATE
    control.cell(4, 1).text = "Current backend codebase at " + CODE_BASELINE
    for row_index in range(8):
        write_cell(control.rows[row_index].cells[0], control.rows[row_index].cells[0].text, size=8.8, bold=True, color=WHITE, fill=BLUE)
        write_cell(control.rows[row_index].cells[1], control.rows[row_index].cells[1].text, size=8.8, fill=WHITE)
    set_table_widths(control, [1.75, 2.76, 2.76])
    evidence_row = control.rows[8]
    evidence_cell = evidence_row.cells[0]
    grid_span = evidence_cell._tc.get_or_add_tcPr().find(qn("w:gridSpan"))
    if grid_span is None:
        grid_span = OxmlElement("w:gridSpan")
        evidence_cell._tc.get_or_add_tcPr().append(grid_span)
    grid_span.set(qn("w:val"), "3")
    tr_pr = evidence_row._tr.get_or_add_trPr()
    for tag in ("w:gridAfter", "w:wAfter"):
        node = tr_pr.find(qn(tag))
        if node is not None:
            tr_pr.remove(node)
    set_cell_width(evidence_cell, 7.27)

    contents = doc.tables[2]
    contents.cell(3, 1).text = "Current module requirements and module-level needs"
    contents.cell(5, 0).text = f"5. v{MRD_VERSION} Documentation Delta"
    contents.cell(5, 1).text = "Terminology, organization, metadata, and presentation corrections"
    for idx, row in enumerate(contents.rows):
        fill = VERY_PALE_BLUE if idx % 2 else WHITE
        write_cell(row.cells[0], row.cells[0].text, size=8.6, bold=True, color=BLUE, fill=fill)
        write_cell(row.cells[1], row.cells[1].text, size=8.6, fill=fill)
    set_table_widths(contents, [2.65, 4.62])

    replace_text(doc, f"5. v{source_version} Capability Delta", f"5. v{MRD_VERSION} Documentation Delta")
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("This section records the backend capabilities added or materially expanded"):
            write_paragraph(
                paragraph,
                f"This patch changes documentation terminology, organization, metadata, and layout only. It does not change module statuses, capability counts, priority gaps, or the recommended build order established in v{source_version}.",
                size=9,
                color=BLACK,
                space_after=5,
            )
            break
    delta = find_table_by_header(doc, f"v{source_version} capability delta")
    rebuild_table(
        delta,
        [f"v{MRD_VERSION} documentation delta", "Decision", "Evidence"],
        [
            ["Document family terminology", "Corrected", "The cross-module tracker is now the Module Requirements Document (MRD); detailed module documents remain Functional Requirements Documents (FRDs)."],
            ["Requirements folder structure", "Separated", "MRD versions and module-level FRD versions are stored in distinct document families while historical files remain preserved."],
            ["Metadata and headers", "Corrected", "Cover, document control, core properties, odd and even headers, version labels, and source baseline now agree."],
            ["Change-log layout", "Repaired", "The table is constrained to the A4 content width and no longer clips outside the page."],
        ],
        [2.55, 1.2, 3.52],
        font_size=8.1,
    )

    change_log = replace_table(
        doc,
        doc.tables[-1],
        ["Version", "Date", "Summary"],
        [
            [MRD_VERSION, "11 Aug 2026", "Renamed the cross-module artifact as the Module Requirements Document (MRD), separated MRD and FRD document families, corrected metadata and header inconsistencies, and repaired the overflowing change-log layout. No roadmap meaning changed."],
            ["2.5", "11 Aug 2026", "Added vendor-invoice duplicate controls: same-vendor hard blocking, cross-vendor warnings with confirmation, early and save-time validation, inline correction, and retry-safe creation."],
            ["2.4", "9 Aug 2026", "Added guarded runtime security and integration settings, scoped compliance overrides, safe connection tests, entitlement reset and scheduling, renewal warnings and bulk scheduling, personal audit views, asynchronous audit exports, fixed-query capability evaluation, and runtime-consumer evidence."],
            ["2.3", "9 Aug 2026", "Revamped against the current backend worktree. Uncommitted code counted as complete. Corrected app mappings and status, added missing feature tables and capabilities, separated backend from integration state, removed unsupported claims, and added priority gaps and build order."],
            ["2.2", "2026", "Previous planning and module-tracking baseline."],
        ],
        [0.75, 1.15, 5.37],
        font_size=7.8,
    )
    rule_row = change_log.add_row()
    merged = rule_row.cells[0].merge(rule_row.cells[1]).merge(rule_row.cells[2])
    write_cell(
        merged,
        "MAINTENANCE RULE | Update backend and integration states independently. Remove a gap only when the implementation and affected integration are verified. Keep module numbers stable and version this MRD independently from module FRDs.",
        size=8,
        bold=True,
        color=BLUE,
        fill=PALE_BLUE,
    )
    set_table_widths(change_log, [0.75, 1.15, 5.37])

    # Only the intended version label in the cover and document-control table changes.
    for paragraph in iter_all_paragraphs(doc):
        if paragraph.text.strip() == f"Version:  {source_version}":
            replace_text(doc, f"Version:  {source_version}", f"Version:  {MRD_VERSION}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    update_extended_title(output_path, f"XVS Module Requirements Document v{MRD_VERSION}")
    shrink_inherited_media(output_path)


FR_REQUIREMENTS = [
    {
        "id": "FR-001",
        "title": "Create a School and Tenant Boundary",
        "status": "Implemented",
        "requirement": "An authorized platform caller can create a School as the business identity and a paired Tenant as the durable security boundary.",
        "evidence": "SchoolCreateSerializer and School.save() create the School and Tenant atomically. New Schools begin PENDING unless an explicit supported status path changes them.",
        "acceptance": "A successful request persists one School, one protected Tenant, and any valid nested onboarding records. Hard database failures roll back the transaction.",
        "limit": "Invitation provisioning uses an inner savepoint and may fail without rolling back the School. The queued admin link then requires operational reconciliation.",
    },
    {
        "id": "FR-002",
        "title": "Generate Stable School Identifiers",
        "status": "Implemented with limits",
        "requirement": "Each School has a stable internal ID, unique slug, and unique uppercase school code suitable for routing, scoping, and operator selection.",
        "evidence": "The slug is generated from the name when omitted, reserved values are rejected, and the code is allocated through the tenant document-number service with the SC prefix.",
        "acceptance": "Slug and code collisions are rejected. List responses expose the internal School ID used by scoped configuration endpoints.",
        "limit": "School names are not unique. The public update serializer does not allow name, slug, or code changes.",
    },
    {
        "id": "FR-003",
        "title": "Apply Platform Onboarding Defaults",
        "status": "Implemented",
        "requirement": "Omitted onboarding fields use centrally managed platform defaults without overwriting explicit caller choices.",
        "evidence": "School creation resolves default ownership type, term structure, currency, and branch country from vs_config. Standalone branch creation uses the same branch-country default.",
        "acceptance": "Changing a supported platform default affects later School or Branch creation only when the corresponding request field is omitted. Explicit valid input wins.",
        "limit": "These defaults do not provide an onboarding checklist, readiness state, go-live decision, or rollback workflow.",
    },
    {
        "id": "FR-004",
        "title": "Provision the Primary School Administrator",
        "status": "Implemented with limits",
        "requirement": "School onboarding can create a primary administrator contact, user, tenant role assignment, invitation, and asynchronous email delivery request.",
        "evidence": "SchoolPrimaryAdmin, ContactInfo, User, TenantUserRoleAssignment, InvitationService, and the Celery invitation task form the current path.",
        "acceptance": "A valid unused email produces a School-scoped inactive user with the prebuilt School administrator role and an invitation record.",
        "limit": "Provisioning exceptions are logged and swallowed. The School remains created and the admin link can remain QUEUED without an exposed retry queue.",
    },
    {
        "id": "FR-005",
        "title": "Create a Branch Under an Active School",
        "status": "Implemented with limits",
        "requirement": "An authorized caller can create a Branch under an ACTIVE School, receive a per-School code, and record the initial PENDING lifecycle event.",
        "evidence": "BranchCreateSerializer validates the School, applies the platform branch-country default when omitted, allocates max(code)+1 while locking rows, exposes the database ID in list/detail responses, and creates BranchLifecycle history.",
        "acceptance": "The created Branch belongs to the route School, has PENDING status, an opened timestamp, a code, and an initial lifecycle entry.",
        "limit": "The database has no unique constraint for (school, code) and no conditional constraint for one main Branch. Serializer checks do not protect direct ORM or all concurrent paths.",
    },
    {
        "id": "FR-006",
        "title": "Provision the Primary Branch Administrator",
        "status": "Implemented with limits",
        "requirement": "Branch onboarding can create a Branch-scoped administrator role assignment, contact, user, and invitation.",
        "evidence": "BranchPrimaryAdmin and the shared admin-provisioning service create the current link. Nested School onboarding can reuse the School admin email without creating a duplicate user.",
        "acceptance": "A valid administrator is linked to the target Branch and its School tenant, and the invitation state is visible.",
        "limit": "Standalone Branch creation declares primary admin data as optional but the creation path effectively requires it. Failed provisioning has the same reconciliation gap as School admins.",
    },
    {
        "id": "FR-007",
        "title": "Assign Package Capacity and Capabilities",
        "status": "Implemented with limits",
        "requirement": "School creation can select an active PackagePlan, validate capacity values, set subscription expiry, and grant dependency-closed capabilities.",
        "evidence": "SchoolPackageSetup validates student, teacher, and administrator capacities and vs_config stores entitlements, dependencies, schedules, overrides, and audit evidence.",
        "acceptance": "Invalid capacity or expiry input fails validation. Effective capability evaluation honors entitlement, dependency state, time window, and scoped override precedence.",
        "limit": "PackagePlan.max_branch is stored but not enforced during School or Branch creation.",
    },
    {
        "id": "FR-008",
        "title": "List, Search, View, and Summarize Schools",
        "status": "Implemented",
        "requirement": "Authorized platform users can page, filter, search, order, retrieve, and summarize Schools.",
        "evidence": "SchoolListView supports status flags, text search across School and Branch fields, allow-listed ordering, nested main Branch data, and stable School IDs. Detail returns branding, admins, package setup, and Branches.",
        "acceptance": "The list is paginated and restricted by platform.schools.view. Statistics return all, active, pending, and inactive counts in one aggregate query.",
        "limit": "SchoolDetailView currently converts unexpected exceptions into a response containing the exception type, message, and traceback. This is unsafe for production.",
    },
    {
        "id": "FR-009",
        "title": "Update School Metadata and Branding",
        "status": "Partial",
        "requirement": "Authorized operators can update supported School metadata and the School logo without changing identity or lifecycle fields through the general update endpoint.",
        "evidence": "SchoolUpdateSerializer allows ownership type, address, website, motto, term structure, currency, registration ID, and optional branding logo.",
        "acceptance": "A material supported change is saved and the endpoint returns the refreshed School detail payload.",
        "limit": "A branding-only request is rejected as 'No changes detected' before branding is upserted. School update does not emit the expected platform audit event. Favicon and broader theme fields do not exist.",
    },
    {
        "id": "FR-010",
        "title": "List, View, Update, and Summarize Branches",
        "status": "Implemented",
        "requirement": "Authorized operators can operate on Branches within the School identified by the route.",
        "evidence": "Branch endpoints support status/main filters, search, ordering, database ID plus business code, detail, metadata update, statistics, and before/after update audit evidence.",
        "acceptance": "A Branch from another route School is not returned. Statistics include all five Branch states.",
        "limit": "Branch-bound staff authorization still requires explicit Branch filtering in every dependent domain; tenant filtering alone is only a School boundary.",
    },
    {
        "id": "FR-011",
        "title": "Transition Branch Lifecycle",
        "status": "Implemented with limits",
        "requirement": "Privileged staff can move a Branch through an allowed lifecycle matrix and write a lifecycle event only after a successful state change.",
        "evidence": "The transition endpoint accepts ACTIVE, SUSPENDED, INACTIVE, and CLOSED targets, resolves the Branch within the route School, and requires platform.branches.manage in addition to XVision staff. Branch.ALLOWED_TRANSITIONS publishes the permitted edges, CLOSED is terminal, and PENDING is not a target. Branch.transition() runs in one transaction, persists status with activated_at, deactivated_at, and closed_at, and appends the BranchLifecycle row.",
        "acceptance": "The implementation must validate allowed transitions, set lifecycle timestamps consistently, persist the state, and append actor/reason history atomically.",
        "limit": "Closing a Branch has no defined effect on dependent records, sessions, or Branch-bound users. Lifecycle control is platform-side only: a School administrator cannot transition a Branch of their own School.",
    },
    {
        "id": "FR-012",
        "title": "Enforce Tenant and Branch Access Boundaries",
        "status": "Partial",
        "requirement": "School-owned records must be isolated by Tenant, and Branch-bound roles must be limited to their assigned Branch where the operation is Branch-specific.",
        "evidence": "School owns a protected Tenant, Branch uses TenantAwareManager, route School slugs scope Branch views, and endpoints declare platform permission keys.",
        "acceptance": "Changing a School slug, Branch code, tenant selector, or entity selector cannot expose another tenant's records. Branch-bound roles receive only authorized Branch data.",
        "limit": "Platform jobs require explicit scoping and Branch-level filtering is not universal across dependent applications.",
    },
    {
        "id": "FR-013",
        "title": "Record School and Branch Audit Evidence",
        "status": "Partial",
        "requirement": "Material School, Branch, configuration, and lifecycle changes must record the actor, target, outcome, before state, and difference where applicable.",
        "evidence": "School create, Branch create, and Branch update emit platform audit evidence. Configuration and entitlement changes use a separate immutable configuration stream.",
        "acceptance": "Successful and failed privileged operations are traceable without exposing secrets or raw traceback data.",
        "limit": "School update and reset configuration do not emit the expected platform events. Branch lifecycle history is now written on every successful transition.",
    },
    {
        "id": "FR-014",
        "title": "Reset School-Owned Configuration",
        "status": "Partial",
        "requirement": "A super administrator can deliberately reset defined School-owned settings to inherited platform defaults using School-bound confirmation and a recorded reason.",
        "evidence": "POST /v1/i/{slug}/reset-config/ exists and currently deletes SchoolBranding inside a transaction.",
        "acceptance": "The confirmation must be bound to the target School, the exact reset scope must be declared, before/after state must be audited, and unrelated operational data must remain intact.",
        "limit": "Any non-empty token passes. Entitlements, configuration values, overrides, localization, package setup, and audit evidence are not reset.",
    },
    {
        "id": "FR-015",
        "title": "Manage School Lifecycle and Safe Deletion",
        "status": "Not implemented",
        "requirement": "The platform requires explicit activation, suspension, reactivation, recoverable soft deletion, and guarded hard deletion with defined session and downstream-record behavior.",
        "evidence": "School currently exposes PENDING, ACTIVE, and INACTIVE fields, and School.save() synchronizes the paired Tenant status.",
        "acceptance": "Lifecycle operations require dedicated permissions, synchronize Tenant state, revoke or restrict access as defined, protect dependent records, and emit audit evidence.",
        "limit": "There is no SUSPENDED or DELETED School state, no lifecycle endpoint, no soft-delete field, and no guarded delete service. Direct School deletion can cascade Branch data while leaving the protected Tenant relationship as a separate concern.",
    },
    {
        "id": "FR-016",
        "title": "Handle Branch Overrides, Transfer, Clone, and Delete Safety",
        "status": "Partial",
        "requirement": "Branch capability overrides must resolve safely, while configuration cloning, staff/student transfer, and Branch deletion must respect ownership and dependent records.",
        "evidence": "vs_config supports Branch-scoped overrides and effective-state evaluation.",
        "acceptance": "Clone, transfer, and delete operations must be explicit, permissioned, tenant-safe, audited, and blocked or migrated when active dependent records exist.",
        "limit": "No supported Branch clone, staff/student transfer, or Branch delete endpoint and dependency-check service exists.",
    },
]


def build_frd(reference_path: Path, output_path: Path) -> None:
    doc = Document(reference_path)
    remove_body_content(doc)
    doc.core_properties.title = f"School and Branch Management Functional Requirements Document v{FRD_VERSION}"
    doc.core_properties.subject = "Code-aligned functional requirements for XVS Module 1"
    doc.core_properties.author = "CodeX Team"
    doc.core_properties.version = FRD_VERSION
    set_headers(doc, "CodeX | School & Branch Management | Functional Requirements Document (FRD)")

    add_cover(
        doc,
        family="Functional Requirements Document",
        title="Module 1: School & Branch Management",
        subtitle="XVision Systems | Code-aligned functional baseline",
        version=FRD_VERSION,
    )

    add_heading(doc, "Document Control", level=1)
    add_metadata_table(
        doc,
        [
            ("Document", "Functional Requirements Document (FRD)"),
            ("Module", "M01 | School & Branch Management"),
            ("Version", FRD_VERSION),
            ("Review date", REVIEW_DATE),
            ("Code baseline", CODE_BASELINE),
            ("Source MRD", f"XVS Module Requirements Document v{MRD_VERSION} | Module 1"),
            ("Primary app", "vs_schools"),
            ("Supporting apps", "vs_tenants, vs_config, vs_rbac, vs_user, vs_audit, vs_import_data"),
            ("Status", "Code-aligned baseline for Product and Engineering review"),
            ("Owner", "CodeX Team"),
        ],
    )
    add_callout(
        doc,
        "Evidence boundary",
        [
            "Implemented means the stated backend path is present in the inspected code. It does not prove frontend completion, deployment, production adoption, or data migration.",
            "Partial, Defective, and Not implemented items are current requirements and risks, not historical notes.",
            "The MRD and this FRD have independent versions. A change can require one or both documents to advance.",
        ],
        kind="info",
    )
    add_page_break(doc)

    add_heading(doc, "Table of Contents", level=1)
    add_table(
        doc,
        ["Section", "Purpose"],
        [
            ["Document Control", "Version, ownership, source baseline, and evidence rules"],
            ["1. Purpose and Scope", "Module responsibility, boundaries, and terminology"],
            ["2. Context and Status Model", "Architecture position and requirement-state definitions"],
            ["3. Actors, Permissions, and Ownership", "Who can act and how School, Tenant, and Branch scope work"],
            ["4. Functional Requirements", "Testable behavior, current evidence, acceptance, and limits"],
            ["5. Workflows and Lifecycle Rules", "School creation, Branch transition, and reset behavior"],
            ["6. Data Model and Relationships", "Persisted ownership and safety contracts"],
            ["7. API and Validation Contracts", "Routes, permissions, inputs, and response behavior"],
            ["8. Dependencies and Operational Evidence", "Configuration, audit, identity, import, and delivery contracts"],
            ["9. Needs Attention", "Current defects and required completion"],
            ["10. MRD Traceability", "Module 1 capability mapping"],
            ["11. Change Log", "Independent FRD revision history"],
        ],
        [2.35, 4.92],
        font_size=8.5,
    )
    add_body(doc, "Use the Word Navigation pane to jump between headings. This contents page is intentionally static for reliable headless rendering.", size=8.5, color=GREY)
    add_page_break(doc)

    add_heading(doc, "1. Purpose and Scope", level=1)
    add_body(
        doc,
        "This FRD defines the functional behavior of Module 1 as required and as currently evidenced in the XVS backend. It turns the Module 1 MRD capability list into testable rules for School identity, Tenant ownership, Branch administration, onboarding, access control, configuration, lifecycle operations, and operational evidence.",
    )
    add_heading(doc, "1.1 In Scope", level=2)
    add_bullets(
        doc,
        [
            "School and paired Tenant creation, identifiers, metadata, status, and branding.",
            "Branch creation, identity, main-Branch designation, administrators, statistics, and lifecycle records.",
            "Package capacity, module entitlements, dependency resolution, and scoped overrides used during onboarding.",
            "Platform permissions, tenant isolation, Branch scoping, audit evidence, validation, and API contracts.",
            "Required lifecycle, reset, transfer, clone, and deletion behavior that is incomplete or defective today.",
        ],
    )
    add_heading(doc, "1.2 Out of Scope and Relocated Responsibilities", level=2)
    add_table(
        doc,
        ["Responsibility", "Owner / Decision"],
        [
            ["Student and staff records", "Future Student Management and Staff Management modules. Module 1 only supplies School and Branch ownership boundaries."],
            ["General identity and authentication", "Identity, Team & Organogram. This FRD covers only primary admin provisioning and access effects of School lifecycle."],
            ["Capability catalogue and configuration storage", "Configuration & Capability Management. This FRD records the School and Branch contracts that consume it."],
            ["Generic audit search and export", "Audit & Activity Logging. This FRD specifies which School and Branch events must be emitted."],
            ["Bulk import orchestration", "Bulk Data Import. It must reuse the validated School and Branch creation contracts."],
        ],
        [2.45, 4.82],
        font_size=8.4,
    )

    add_heading(doc, "2. Context and Status Model", level=1)
    add_body(
        doc,
        "A School is the operator-facing business identity. Its one-to-one Tenant is the durable security boundary used by authentication, RBAC, configuration, workflow, finance, procurement, notifications, support, and every tenant-aware domain. Branches represent campuses or operating locations inside that School boundary.",
    )
    add_heading(doc, "2.1 Status Model", level=2)
    add_status_key(doc)
    add_callout(
        doc,
        "Current module decision",
        [
            f"Module 1 remains Backend Partial and Integration Partial in MRD v{MRD_VERSION}.",
            "The defining blockers are School lifecycle controls and safe handling of closed or deleted Schools and Branches. The Branch transition defect is resolved in this revision.",
        ],
        kind="attention",
    )
    add_page_break(doc)

    add_heading(doc, "3. Actors, Permissions, and Ownership", level=1)
    add_heading(doc, "3.1 Actors", level=2)
    add_table(
        doc,
        ["Actor", "Responsibility"],
        [
            ["Platform School viewer", "List, search, summarize, and retrieve Schools and Branches."],
            ["Platform School administrator", "Create and update Schools using platform.schools.create or platform.schools.update."],
            ["Platform Branch administrator", "Create and update Branches using platform.branches.create or platform.branches.update."],
            ["XVision staff", "Run Branch lifecycle transitions with platform.branches.manage, and view package/capability catalogues."],
            ["XVision super administrator", "Run the current reset configuration operation."],
            ["School administrator", "Operate within the paired Tenant through permissions provisioned by the School admin role."],
            ["Branch administrator", "Operate within the School tenant and assigned Branch where dependent views enforce Branch scope."],
            ["Background job or import process", "Use explicit Tenant and School scope and preserve the same validation and audit contracts."],
        ],
        [2.1, 5.17],
        font_size=8.4,
    )
    add_heading(doc, "3.2 Endpoint Gates", level=2)
    add_table(
        doc,
        ["Operation", "Permission / Gate", "Scope"],
        [
            ["School list, detail, statistics", "platform.schools.view", "Platform School collection"],
            ["Create School", "platform.schools.create", "Creates School, Tenant, and optional nested onboarding records"],
            ["Update School", "platform.schools.update", "Supported metadata and logo only"],
            ["Branch list, detail, statistics", "platform.branches.view", "Route School plus ambient tenant filtering"],
            ["Create Branch", "platform.branches.create", "ACTIVE route School"],
            ["Update Branch", "platform.branches.update", "Route School and Branch business code"],
            ["Transition Branch", "IsVisionStaff and platform.branches.manage", "Resolved within the route School"],
            ["Reset School configuration", "IsVisionSuperAdmin", "Current behavior removes branding only"],
            ["Platform onboarding defaults", "config.value.view / config.value.update", "Platform-only curated settings"],
            ["Entitlements", "config.entitlement.view / manage", "Platform or School Tenant scope"],
            ["Overrides", "config.override.view / manage", "Platform, School, or Branch scope"],
        ],
        [2.1, 2.3, 2.87],
        font_size=8.0,
    )
    add_heading(doc, "3.3 Ownership Rules", level=2)
    add_bullets(
        doc,
        [
            "School has one protected Tenant. Deleting the Tenant cannot silently delete the School.",
            "Branch belongs to exactly one School and uses TenantAwareManager for ambient School-tenant filtering.",
            "Platform actors and background jobs can have no ambient School tenant, so they must scope explicit targets.",
            "Branch-bound authorization is narrower than tenant authorization and must be enforced separately in every Branch-sensitive consumer.",
        ],
    )
    add_page_break(doc)

    add_heading(doc, "4. Functional Requirements", level=1)
    add_body(doc, "Each requirement records the required behavior, inspected implementation evidence, acceptance boundary, and current limitation. Status is current state, not revision history.")
    for index, requirement in enumerate(FR_REQUIREMENTS):
        add_requirement(doc, requirement)
        if index in {3, 7, 11}:
            add_page_break(doc)

    add_page_break(doc)
    add_heading(doc, "5. Workflows and Lifecycle Rules", level=1)
    add_heading(doc, "5.1 School Creation Transaction", level=2)
    add_table(
        doc,
        ["Step", "Current behavior", "Required outcome"],
        [
            ["1. Validate", "Validate School identifiers, nested Branch rules, admins, package, and explicit values; fill omitted onboarding defaults.", "Reject invalid payloads before durable side effects."],
            ["2. Create boundary", "Create Tenant and School in one transaction; allocate the SC document code.", "Persist exactly one protected Tenant per School."],
            ["3. Optional School records", "Create branding and primary School admin link.", "Keep identity records consistent and traceable."],
            ["4. Branches", "Create Branches, initial lifecycle rows, primary admin links, roles, and invitations.", "Keep all database writes atomic except the deliberate invitation savepoint."],
            ["5. Package and capabilities", "Create SchoolPackageSetup and dependency-closed entitlements.", "Reject invalid capacity, expiry, or capability input."],
            ["6. Evidence", "Emit School and Branch creation audit events; dispatch invitations asynchronously.", "Expose failed invitation reconciliation without rolling back valid School creation."],
        ],
        [0.9, 3.15, 3.22],
        font_size=8.0,
    )
    add_heading(doc, "5.2 Branch Lifecycle", level=2)
    add_table(
        doc,
        ["Area", "Current state", "Required rule"],
        [
            ["States", "PENDING, ACTIVE, SUSPENDED, INACTIVE, CLOSED", "Branch.ALLOWED_TRANSITIONS publishes the matrix and refuses every other move with INVALID_BRANCH_TRANSITION."],
            ["Persistence", "The save call persists status, activated_at, deactivated_at, and closed_at.", "BranchLifecycle is written in the same transaction as the status change."],
            ["Timestamps", "activated_at records first activation only; deactivated_at is stamped on SUSPENDED, INACTIVE, and CLOSED and cleared on return to ACTIVE; CLOSED stamps closed_at.", "Retain this policy for every new state."],
            ["Access", "Closed-Branch access effects are not defined.", "Define authentication, session, write, read-only, transfer, and dependent-record rules."],
            ["Authorization", "IsVisionStaff and platform.branches.manage", "Retain both gates so a mis-granted key cannot reach a platform commercial action."],
        ],
        [1.25, 2.85, 3.17],
        font_size=8.1,
    )
    add_heading(doc, "5.3 School Reset", level=2)
    add_callout(
        doc,
        "Required reset contract",
        [
            "Declare whether branding, configuration values, School overrides, Branch overrides, entitlements, and package setup are included.",
            "Bind confirmation to the School identity and reject reusable arbitrary tokens.",
            "Record before and after snapshots and the operator's reason.",
            "Never delete operational School or Branch data as a side effect of configuration reset.",
        ],
        kind="attention",
    )

    add_page_break(doc)
    add_heading(doc, "6. Data Model and Relationships", level=1)
    add_table(
        doc,
        ["Model", "Purpose", "Key safety contract"],
        [
            ["School", "Business identity and School metadata", "Unique slug/code; protected one-to-one Tenant; only PENDING, ACTIVE, INACTIVE today"],
            ["Tenant", "Security and ownership boundary", "One School profile; status synchronized by School.save()"],
            ["Branch", "Campus or operating location", "Belongs to School; tenant-aware manager; no database uniqueness for code/main flag"],
            ["SchoolBranding", "School logo container", "One-to-one School; logo only"],
            ["PackagePlan", "Subscription plan catalogue", "Unique code/name; capacity maxima; active catalogue flag"],
            ["SchoolPackageSetup", "Selected plan, capacities, expiry", "One-to-one School; protected plan relationship"],
            ["ContactInfo", "Primary admin contact card", "Email indexed but not database-unique"],
            ["SchoolPrimaryAdmin", "School admin contact and invitation state", "One-to-one School"],
            ["BranchPrimaryAdmin", "Branch admin contact and invitation state", "One-to-one Branch"],
            ["BranchLifecycle", "Branch state-change history", "Deleted with Branch today; written on every successful transition"],
        ],
        [1.55, 2.35, 3.37],
        font_size=8.0,
    )
    add_heading(doc, "6.1 Relationship Decisions", level=2)
    add_callout(
        doc,
        "Database enforcement gaps",
        [
            "Add a unique constraint for (school, code) and a conditional unique constraint for one is_main Branch per School.",
            "Define deletion retention for BranchLifecycle before adding Branch deletion.",
            "Keep Tenant protection and explicit School lifecycle behavior separate from raw ORM cascade behavior.",
        ],
        kind="critical",
    )

    add_page_break(doc)
    add_heading(doc, "7. API and Validation Contracts", level=1)
    add_heading(doc, "7.1 School Endpoints", level=2)
    add_table(
        doc,
        ["Method", "Endpoint", "Permission", "Purpose"],
        [
            ["GET", "/v1/i/", "platform.schools.view", "Paginated School list, search, filters, ordering, and stable IDs"],
            ["POST", "/v1/i/create/", "platform.schools.create", "Create School, Tenant, and optional onboarding records"],
            ["GET", "/v1/i/stats/", "platform.schools.view", "Counts by School status"],
            ["GET", "/v1/i/package-plans/", "IsVisionStaff", "Active package catalogue"],
            ["GET", "/v1/i/modules/", "IsVisionStaff", "Active capability catalogue and dependencies"],
            ["GET", "/v1/i/{slug}/", "platform.schools.view", "School detail and nested records"],
            ["PUT/PATCH", "/v1/i/{slug}/update/", "platform.schools.update", "Supported metadata and logo"],
            ["POST", "/v1/i/{slug}/reset-config/", "IsVisionSuperAdmin", "Current branding-only reset"],
        ],
        [0.65, 2.25, 1.65, 2.72],
        font_size=7.7,
    )
    add_heading(doc, "7.2 Branch Endpoints", level=2)
    add_table(
        doc,
        ["Method", "Endpoint", "Permission", "Purpose"],
        [
            ["GET", "/v1/i/{slug}/branches/", "platform.branches.view", "School-scoped Branch list with IDs and filters"],
            ["POST", "/v1/i/{slug}/branches/create/", "platform.branches.create", "Create PENDING Branch under ACTIVE School"],
            ["GET", "/v1/i/{slug}/branches/stats/", "platform.branches.view", "Counts by Branch status"],
            ["GET", "/v1/i/{slug}/branches/{code}/detail/", "platform.branches.view", "Branch detail and primary admin"],
            ["PUT/PATCH", "/v1/i/{slug}/branches/{code}/update/", "platform.branches.update", "Update Branch metadata"],
            ["POST", "/v1/i/{slug}/branches/{code}/transition/", "IsVisionStaff and platform.branches.manage", "Lifecycle transition; 409 on a refused edge or repeated state"],
        ],
        [0.65, 2.55, 1.65, 2.42],
        font_size=7.6,
    )
    add_heading(doc, "7.3 Related Configuration Endpoints", level=2)
    add_table(
        doc,
        ["Endpoint", "Contract"],
        [
            ["/v1/config/platform-settings/", "Read/update platform identity and School onboarding defaults; explicit values remain caller-owned."],
            ["/v1/config/entitlements/", "List or schedule platform/School capability grants."],
            ["/v1/config/entitlements/{capability}/", "Remove a scoped grant and reveal inherited entitlement state."],
            ["/v1/config/entitlements/calendar/", "Inspect expiry, renewal, and scheduled activation windows."],
            ["/v1/config/entitlements/bulk-schedule/", "Atomically schedule up to 100 distinct entitlement targets."],
            ["/v1/config/overrides/", "Manage platform, School, or Branch capability overrides."],
            ["/v1/config/effective-capabilities/", "Return final capability state after grant, dependency, schedule, and override resolution."],
        ],
        [2.55, 4.72],
        font_size=8.0,
    )

    add_page_break(doc)
    add_heading(doc, "7.4 Validation and Response Rules", level=2)
    add_table(
        doc,
        ["Area", "Required / Current rule", "Response"],
        [
            ["School slug", "Lowercase URL-safe, unique, not reserved; collision suggestions", "400 on invalid input"],
            ["School code", "Unique and allocated when omitted", "400 or database error on collision"],
            ["Nested Branches", "Exactly one main Branch when Branches are supplied; unique names in payload", "400"],
            ["Branch target", "Route School must exist and be ACTIVE for creation", "404 or 400"],
            ["Admin email", "Must not already belong to a User except the deliberate nested School-admin reuse path", "400"],
            ["Package capacity", "At least 1 and no more than plan maximum", "400"],
            ["Subscription expiry", "Cannot be in the past", "400"],
            ["No-op update", "At least one direct School or Branch field must change", "400"],
            ["Cross-School target", "Tenant and route scope must hide foreign records", "404 or empty result"],
            ["Unexpected School detail failure", "Must log server-side and return a generic response", "Current code leaks debug details"],
        ],
        [1.5, 4.57, 1.2],
        font_size=7.9,
    )

    add_heading(doc, "8. Dependencies and Operational Evidence", level=1)
    add_table(
        doc,
        ["Dependency", "Direction", "Module 1 contract"],
        [
            ["vs_tenants", "Required write and scope", "Create one Tenant per School; provide ambient tenant context and SC document numbering."],
            ["vs_rbac", "Authorization and provisioning", "Evaluate platform permissions; provision School/Branch admin roles and Branch assignments."],
            ["vs_user", "Outbound write and async delivery", "Create pending users and invitations; send activation email."],
            ["vs_config", "Bidirectional", "Supply onboarding defaults, entitlements, dependencies, schedules, overrides, and audit evidence."],
            ["vs_audit", "Outbound evidence", "Store actor, target, outcome, before state, and differences for material actions."],
            ["vs_import_data", "Inbound reuse", "Reuse School/Branch validation and creation behavior and provide batch rollback."],
            ["vs_admin_console", "Read consumer", "Consume School counts/details and expose controlled operator navigation."],
            ["Tenant-aware domain apps", "Ownership contract", "Resolve through Tenant, School, and Branch without cross-School leakage."],
            ["Database-backed media", "Storage contract", "Store School logos and serve them through authenticated media routes."],
        ],
        [1.6, 1.7, 3.97],
        font_size=8.0,
    )
    add_heading(doc, "8.1 Minimum Verification", level=2)
    add_bullets(
        doc,
        [
            "Permission denied and cross-Tenant isolation for every endpoint.",
            "School creation happy path, nested rollback, admin provisioning failure, and explicit-vs-default onboarding values.",
            "Concurrent Branch code allocation plus database constraint coverage once constraints are added.",
            "Every allowed and rejected School/Branch lifecycle transition with session and access effects.",
            "School update, Branch update, reset, lifecycle, and delete audit evidence.",
            "Capability dependency, schedule, override precedence, reset, and expiry behavior.",
            "Generic production error responses without exception class, message, traceback, secrets, or PII.",
            "Empty-list response shape through the shared success-response and paginator layers.",
        ],
    )

    add_page_break(doc)
    add_heading(doc, "9. Needs Attention", level=1)
    add_table(
        doc,
        ["Priority", "Current gap", "Required completion", "FR"],
        [
            ["P0", "School lifecycle controls", "Add activation, suspension, reactivation, recoverable soft deletion, guarded hard deletion, access effects, and audit evidence.", "FR-015"],
            ["P0", "Unsafe School detail errors", "Return generic production responses and keep traceback data in server logs only.", "FR-008"],
            ["P1", "Database Branch constraints", "Enforce unique (school, code) and one main Branch per School at the database boundary.", "FR-005"],
            ["P1", "School update and reset audit", "Emit before/after evidence and allow valid branding-only updates.", "FR-009, FR-013, FR-014"],
            ["P1", "Admin provisioning reconciliation", "Expose queued/failed invitations, retry action, and operator visibility.", "FR-004, FR-006"],
            ["P1", "Reset semantics", "Declare reset scope, bind confirmation to School, and restore inherited state safely.", "FR-014"],
            ["P1", "Branch delete and closed access", "Define retention, dependency checks, transfer rules, read/write restrictions, and safe deletion. CLOSED is now terminal, so its effect on dependent records must be stated.", "FR-011, FR-016"],
            ["P2", "Package Branch capacity", "Enforce max_branch or remove it from the active product contract.", "FR-007"],
        ],
        [0.7, 1.65, 4.1, 0.82],
        font_size=7.7,
    )
    add_callout(
        doc,
        "Removal rule",
        [
            "An item leaves Needs Attention only when implementation and relevant verification resolve it fully.",
            "Partly resolved items must be rewritten to describe the remaining risk rather than deleted.",
            "Optional ideas and speculative features do not belong here.",
        ],
        kind="attention",
    )

    add_page_break(doc)
    add_heading(doc, "10. MRD Traceability", level=1)
    add_body(doc, f"This table reconciles every Module 1 capability entry in MRD v{MRD_VERSION} to the controlling functional requirement in this FRD.")
    add_table(
        doc,
        ["MRD Module 1 capability", "FRD coverage", "Current state"],
        [
            ["Create, list, retrieve, and update Schools", "FR-001, FR-008, FR-009", "Implemented with update limits"],
            ["Unique School slug and code allocation", "FR-002", "Implemented with name-uniqueness limit"],
            ["School status, contact, location, and timezone data", "FR-001, FR-009, FR-015", "Partial; no School timezone field or full lifecycle"],
            ["School branding and logo record", "FR-009", "Partial; logo only and branding-only update defect"],
            ["Package plan and enabled-module setup", "FR-007", "Implemented with max_branch limit"],
            ["Create first School and Branch administrators", "FR-004, FR-006", "Implemented with reconciliation gap"],
            ["Create, list, retrieve, and update Branches", "FR-005, FR-010", "Implemented"],
            ["Per-School Branch code allocation with locking", "FR-005", "Implemented without database uniqueness"],
            ["Main-Branch designation and Branch statistics", "FR-005, FR-010", "Implemented without database main-Branch constraint"],
            ["Branch status transitions and lifecycle history", "FR-011", "Implemented with closed-Branch access limit"],
            ["Tenant-aware School and Branch query scoping", "FR-012", "Partial at Branch-bound authorization layer"],
            ["Reset School-owned configuration to platform defaults", "FR-014", "Partial; branding-only"],
        ],
        [3.15, 1.75, 2.37],
        font_size=8.0,
    )

    add_page_break(doc)
    add_heading(doc, "11. Change Log", level=1)
    change_log = add_table(
        doc,
        ["Version", "Date", "Summary"],
        [
            ["1.2", "12 Aug 2026", f"FR-011 moves from Defective to Implemented with limits. Branch.transition() persists the real lifecycle timestamps, publishes an allowed-transition matrix with CLOSED terminal and PENDING unreachable, and writes BranchLifecycle in the same transaction. The endpoint resolves the Branch within its route School and now requires platform.branches.manage. Reconciled to MRD v{MRD_VERSION}."],
            ["1.1", "11 Aug 2026", "Redesigned using the XVS MRD visual system; reconciled to current code and MRD v2.5.1; added platform onboarding defaults, Branch IDs, explicit acceptance criteria, workflow rules, current Needs Attention, and MRD traceability."],
            ["1.0", "1 Aug 2026", "Initial code-aligned engineering FRD based on MRD predecessor v2.2 and commit 192f9782e3e9812a215ad4e3f322777ac12d9461."],
        ],
        [0.75, 1.15, 5.37],
        font_size=8.0,
    )
    row = change_log.add_row()
    merged = row.cells[0].merge(row.cells[1]).merge(row.cells[2])
    write_cell(
        merged,
        "MAINTENANCE RULE | Version this FRD independently. Update it only when Module 1 behavior, acceptance criteria, permissions, data contracts, APIs, dependencies, status, or current gaps change. Preserve every prior version.",
        size=8,
        bold=True,
        color=BLUE,
        fill=PALE_BLUE,
    )
    set_table_widths(change_log, [0.75, 1.15, 5.37])

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    update_extended_title(
        output_path,
        f"School and Branch Management Functional Requirements Document v{FRD_VERSION}",
    )
    shrink_inherited_media(output_path)


def shrink_inherited_media(path: Path) -> None:
    """Re-encode the cover art a freshly built document inherited from its reference.

    Every document here is built from an earlier one, so its media comes along
    for the ride. The cover photo used to arrive as a 1.28 MB JPEG, 95% of the
    finished file, and it would keep arriving that way in every future version
    unless the reference happened to be one that had already been shrunk. Doing
    it on the way out means it cannot depend on which reference was passed.
    """
    for name, (before, after) in rewrite_document(path, dry_run=False).items():
        print(f"{path}: {name} {before / 1024:.0f} KB -> {after / 1024:.0f} KB")


def assert_no_em_dash(path: Path) -> None:
    doc = Document(path)
    text = "\n".join(paragraph.text for paragraph in iter_all_paragraphs(doc))
    if "\u2014" in text:
        raise ValueError(f"Em dash found in {path}")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--reference", type=Path, required=True)
    parser.add_argument(
        "--mrd-output",
        type=Path,
        help=(
            "Optional. build_mrd() patches the v2.5.1 rename delta onto its "
            "reference and is kept for history only. The MRD has moved on past "
            "that revision, so pass this only with a matching reference - "
            "later MRD versions are produced by patch_mrd.py."
        ),
    )
    parser.add_argument("--frd-output", type=Path, required=True)
    args = parser.parse_args()

    if args.mrd_output is not None:
        build_mrd(args.reference, args.mrd_output)
        assert_no_em_dash(args.mrd_output)
        print(args.mrd_output)
    build_frd(args.reference, args.frd_output)
    assert_no_em_dash(args.frd_output)
    print(args.frd_output)


if __name__ == "__main__":
    main()
